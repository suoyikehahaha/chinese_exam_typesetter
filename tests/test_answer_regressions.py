from __future__ import annotations

from pathlib import Path
import tempfile
import unittest

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.notice_material_formatting import apply_notice_material_formatting
from app.flexible_importers import (
    _annotate_material_roles,
    _lift_embedded_sections,
    _normalize_header_notices,
    _normalize_prose_title_authors,
    _normalize_single_poetry,
)


class AnswerRegressionTests(unittest.TestCase):
    def test_candidate_notice_is_supported(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            source = Path(directory) / "notice.docx"
            document = Document()
            for text in (
                "某校五月检测",
                "语  文",
                "考生须知：",
                "1.本卷满分150分，考试时间150分钟；",
                "2.答案写在答题卷上。",
                "一、阅读（72分）",
            ):
                document.add_paragraph(text)
            document.save(source)
            raw = {
                "metadata": {
                    "exam_name": "某校五月检测",
                    "subject_name": "语　文",
                    "meta_text": "考生须知：",
                    "notices": [],
                    "exam_info_text": "1.本卷满分150分，考试时间150分钟；",
                },
                "blocks": [],
            }
            _normalize_header_notices(source, raw)
            self.assertEqual(raw["metadata"]["notice_title"], "考生须知：")
            self.assertEqual(
                raw["metadata"]["notices"],
                [
                    "本卷满分150分，考试时间150分钟；",
                    "答案写在答题卷上。",
                ],
            )
            self.assertEqual(raw["metadata"]["meta_text"], "")
            self.assertNotIn("exam_info_text", raw["metadata"])

    def test_lifts_major_section_from_question_payload(self) -> None:
        raw = {
            "blocks": [
                {
                    "type": "question",
                    "question": {
                        "number": 17,
                        "embedded_segments": [
                            [
                                {
                                    "text": "二 、语言文字运用 （本题共5小题，18分）",
                                    "role": "body",
                                }
                            ]
                        ],
                    },
                },
                {"type": "instruction", "text": "阅读下面的文字。"},
            ]
        }
        _lift_embedded_sections(raw)
        self.assertEqual(
            [block["type"] for block in raw["blocks"]],
            ["question", "section_title", "instruction"],
        )
        self.assertEqual(
            raw["blocks"][1]["text"],
            "二 、语言文字运用 （本题共5小题，18分）",
        )
        self.assertEqual(
            raw["blocks"][0]["question"]["embedded_segments"],
            [],
        )

    def test_prose_title_multi_author_and_source(self) -> None:
        raw = {
            "blocks": [
                {"type": "instruction", "text": "阅读下面的文字，完成各题。"},
                {
                    "type": "material",
                    "paragraphs": [
                        "我在“天路”守护你",
                        "贺勇 、阿尔达克",
                        "3月，大地回春。",
                        "（摘选自《人民日报》副刊，有删减）",
                    ],
                },
            ]
        }
        _normalize_prose_title_authors(raw)
        material = raw["blocks"][1]
        self.assertEqual(material["title"], "我在“天路”守护你")
        self.assertEqual(material["author"], "贺勇 、阿尔达克")
        self.assertEqual(material["paragraphs"], ["3月，大地回春。"])
        self.assertEqual(
            material["source"],
            "（摘选自《人民日报》副刊，有删减）",
        )

    def test_dynasty_prefixed_single_poem(self) -> None:
        raw = {
            "blocks": [
                {"type": "instruction", "text": "阅读下面这首清诗，完成各题。"},
                {
                    "type": "material",
                    "paragraphs": [
                        "春晚书山家屋壁二首（其二）",
                        "【五代】贯休",
                        "水香塘黑蒲森森，鸳鸯鸂鶒如家禽。",
                    ],
                },
            ]
        }
        _normalize_single_poetry(raw)
        self.assertEqual(raw["blocks"][1]["type"], "poetry")
        self.assertEqual(raw["blocks"][1]["title"], "春晚书山家屋壁二首（其二）")
        self.assertEqual(raw["blocks"][1]["author"], "【五代】贯休")

    def test_centered_bold_material_heading_is_annotated(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            source = Path(directory) / "subheading.docx"
            document = Document()
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.add_run("在这儿工作，就是要耐得住寂寞").bold = True
            document.save(source)
            raw = {
                "blocks": [
                    {
                        "type": "material",
                        "paragraphs": [
                            "正文。",
                            "在这儿工作，就是要耐得住寂寞",
                            "后文。",
                        ],
                    }
                ]
            }
            _annotate_material_roles(source, raw)
            block = raw["blocks"][0]
            self.assertEqual(
                block["paragraph_roles"],
                ["body", "subheading", "body"],
            )
            spec = block["paragraph_formats"][0]
            self.assertEqual(spec["font"], "黑体")
            self.assertEqual(spec["alignment"], "居中")

    def test_notice_alias_replaces_rendered_heading(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            source = Path(directory) / "rendered.docx"
            document = Document()
            document.styles.add_style("Exam_notice_title", 1)
            paragraph = document.add_paragraph(
                "注意事项：",
                style="Exam_notice_title",
            )
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            document.save(source)
            apply_notice_material_formatting(
                source,
                {
                    "metadata": {"notice_title": "考生须知："},
                    "blocks": [],
                },
            )
            paragraph = Document(source).paragraphs[0]
            self.assertEqual(paragraph.text, "考生须知：")
            self.assertEqual(paragraph.runs[0].font.name, "SimHei")


if __name__ == "__main__":
    unittest.main()
