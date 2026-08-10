from __future__ import annotations

from pathlib import Path
import tempfile
import unittest

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

from app.answer_typesetting import (
    parse_answer_docx,
    render_answer_docx,
    standalone_answer_model,
)
from app.flexible_importers import _trim_answer_from_exam, import_exam


class AnswerTypesettingTests(unittest.TestCase):
    def test_standalone_answer_is_semantic_and_keeps_table(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            source = Path(directory) / "answer.docx"
            output = Path(directory) / "answer-output.docx"
            document = Document()
            title = document.add_paragraph("某市高中毕业班综合测试")
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title.runs[0].bold = True
            subtitle = document.add_paragraph("语文试题参考答案及评分建议")
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
            subtitle.runs[0].bold = True
            document.add_paragraph("一、阅读（72分）")
            document.add_paragraph("（一）阅读Ⅰ（19分）")
            document.add_paragraph("1．（3分）")
            document.add_paragraph("B（A项错误。）")
            document.add_paragraph("4．（4分）")
            document.add_paragraph("①第一点；②第二点。")
            document.add_paragraph("评分参考：")
            document.add_paragraph("每答出一点给2分。")
            document.add_table(rows=1, cols=2)
            document.save(source)

            raw = import_exam(source)
            self.assertEqual(raw["document_kind"], "answer")
            self.assertTrue(
                any(block["type"] == "answer_table" for block in raw["blocks"])
            )
            render_answer_docx(raw, output)
            rendered = Document(output)
            self.assertEqual(len(rendered.tables), 1)
            self.assertEqual(rendered.paragraphs[0].alignment, WD_ALIGN_PARAGRAPH.CENTER)
            self.assertEqual(
                rendered.paragraphs[0].runs[0]
                ._element.get_or_add_rPr()
                .get_or_add_rFonts()
                .get(qn("w:eastAsia")),
                "SimSun",
            )
            scoring = next(
                paragraph
                for paragraph in rendered.paragraphs
                if paragraph.text == "评分参考："
            )
            self.assertEqual(
                scoring.runs[0]
                ._element.get_or_add_rPr()
                .get_or_add_rFonts()
                .get(qn("w:eastAsia")),
                "SimHei",
            )

    def test_answer_parser_assigns_guangzhou_roles(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            source = Path(directory) / "roles.docx"
            document = Document()
            for text in (
                "测试卷",
                "语文试题参考答案及评分建议",
                "一、阅读（72分）",
                "（一）阅读Ⅰ（19分）",
                "1．（3分）",
                "B（解释。）",
                "4．（4分）",
                "①结论一；②结论二。",
                "评分参考：",
                "每答出一点给2分。",
            ):
                paragraph = document.add_paragraph(text)
                if text in {"测试卷", "语文试题参考答案及评分建议"}:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            document.save(source)
            answer = parse_answer_docx(source)
            model = standalone_answer_model(answer)
            questions = [
                block for block in model["blocks"] if block["type"] == "answer_question"
            ]
            self.assertEqual(
                questions[0]["paragraphs"][0]["role"],
                "objective_answer",
            )
            self.assertEqual(
                [entry["role"] for entry in questions[1]["paragraphs"]],
                ["subjective_answer", "scoring_label", "scoring_rule"],
            )

    def test_q23_prompt_and_requirements_survive_answer_cutoff(self) -> None:
        raw = {
            "blocks": [
                {
                    "type": "question",
                    "question": {
                        "number": 23,
                        "composition_material": ["核心材料", "参考答案", "1．答案：B"],
                        "composition_prompt": ["以上材料引发了怎样的思考？"],
                        "composition_requirements": ["要求：不少于800字。"],
                    },
                }
            ]
        }
        _trim_answer_from_exam(raw)
        question = raw["blocks"][0]["question"]
        self.assertEqual(question["composition_material"], ["核心材料"])
        self.assertEqual(
            question["composition_prompt"],
            ["以上材料引发了怎样的思考？"],
        )
        self.assertEqual(
            question["composition_requirements"],
            ["要求：不少于800字。"],
        )


if __name__ == "__main__":
    unittest.main()
