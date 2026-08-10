"""Run-level formatting helpers that preserve native character properties."""

from __future__ import annotations

from pathlib import Path
from typing import Any, Iterable

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


def set_run_font_preserving_properties(
    run: Any,
    font_name: str,
    size_pt: float,
    *,
    bold: bool | None = None,
) -> None:
    """Set font metadata while retaining underline, emphasis, border and shading."""

    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    for key in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn(f"w:{key}"), font_name)
    _set_single_value(rpr, "w:sz", str(round(float(size_pt) * 2)))
    _set_single_value(rpr, "w:szCs", str(round(float(size_pt) * 2)))
    if bold is not None:
        _set_boolean(rpr, "w:b", bool(bold))
        _set_boolean(rpr, "w:bCs", bool(bold))
    run.font.name = font_name
    run.font.size = Pt(float(size_pt))
    if bold is not None:
        run.bold = bool(bold)


def protect_inline_properties(docx_path: str | Path) -> None:
    """Normalize boxed A-H markers without flattening their run properties."""

    target = Path(docx_path)
    document = Document(target)
    for paragraph in _all_paragraphs(document):
        for run in paragraph.runs:
            text = run.text.strip()
            if len(text) == 1 and text in "ABCDEFGH" and _has_border(run):
                set_run_font_preserving_properties(run, "SimSun", 10.5)
    document.save(target)


def _has_border(run: Any) -> bool:
    rpr = run._element.get_or_add_rPr()
    return rpr.find(qn("w:bdr")) is not None


def _set_single_value(parent: Any, tag: str, value: str) -> None:
    element = parent.find(qn(tag))
    if element is None:
        element = OxmlElement(tag)
        parent.append(element)
    element.set(qn("w:val"), value)


def _set_boolean(parent: Any, tag: str, enabled: bool) -> None:
    element = parent.find(qn(tag))
    if enabled:
        if element is None:
            element = OxmlElement(tag)
            parent.append(element)
        element.set(qn("w:val"), "1")
    elif element is not None:
        parent.remove(element)


def _all_paragraphs(document: Any) -> Iterable[Any]:
    yield from document.paragraphs
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs


__all__ = ["protect_inline_properties", "set_run_font_preserving_properties"]
