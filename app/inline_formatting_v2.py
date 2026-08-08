"""按编辑器结构定位应用字符级格式。"""

from __future__ import annotations

from copy import deepcopy
from pathlib import Path
import re
from typing import Any

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt


FONT_NAMES = {
    "宋体": "SimSun",
    "黑体": "SimHei",
    "楷体": "KaiTi",
    "仿宋": "FangSong",
}
QUESTION_RE = re.compile(r"^(\d{1,2})．")


def apply_inline_formats_v2(docx_path: str | Path, raw_exam: dict[str, Any]) -> None:
    """应用题干、选项、题内材料、作文和普通内容块的局部格式。"""

    target_path = Path(docx_path)
    document = Document(target_path)
    for block in raw_exam.get("blocks", []):
        entries = (
            block.get("question", {}).get("inline_formats", [])
            if block.get("type") == "question"
            else block.get("inline_formats", [])
        )
        for entry in entries:
            paragraph, prefix = _locate(document, block, entry)
            if paragraph is None:
                continue
            _format_range(
                paragraph,
                int(entry.get("start", 0)) + prefix,
                int(entry.get("end", 0)) + prefix,
                FONT_NAMES.get(
                    str(entry.get("font", "")),
                    str(entry.get("font", "")),
                ),
                float(entry.get("size_pt", 10.5)),
            )
    document.save(target_path)


def _locate(
    document: Any,
    block: dict[str, Any],
    entry: dict[str, Any],
) -> tuple[Any | None, int]:
    target = str(entry.get("target", "stem"))
    index = int(entry.get("target_index", 0))
    if block.get("type") != "question":
        lines = _block_lines(block)
        if index >= len(lines):
            return None, 0
        return _find_exact(document, lines[index]), 0

    question = block["question"]
    if target == "stem":
        number = int(question["number"])
        for paragraph in document.paragraphs:
            match = QUESTION_RE.match(paragraph.text)
            if match and int(match.group(1)) == number:
                return paragraph, len(f"{number}．")
        return None, 0
    if target == "option":
        values = question.get("options", [])
        return (
            _find_exact(document, str(values[index])) if index < len(values) else None,
            0,
        )
    if target == "embedded":
        values = question.get("embedded_segments", [])
        text = (
            "".join(str(segment.get("text", "")) for segment in values[index])
            if index < len(values)
            else ""
        )
        return _find_exact(document, text), 0
    if target == "segmentation":
        return _find_exact(document, str(question.get("segmentation_text", ""))), 0
    if target == "subquestion":
        values = question.get("subquestions", [])
        text = str(values[index]) if index < len(values) else ""
        marker = f"（{index + 1}）"
        return _find_exact(document, marker + text), len(marker)
    values = question.get(target, [])
    text = str(values[index]) if index < len(values) else ""
    return _find_exact(document, text), 0


def _block_lines(block: dict[str, Any]) -> list[str]:
    if block.get("type") in {"section_title", "instruction"}:
        return [str(block.get("text", ""))]
    if block.get("type") == "subsection":
        return [str(block.get("name", "")) + str(block.get("meta", ""))]
    values: list[str] = []
    for key in ("title", "author"):
        if block.get(key):
            values.append(str(block[key]))
    values.extend(str(value) for value in block.get("paragraphs", []))
    for key in ("note", "source"):
        if block.get(key):
            values.append(str(block[key]))
    return values


def _find_exact(document: Any, text: str) -> Any | None:
    for paragraph in document.paragraphs:
        if paragraph.text.strip() == text.strip():
            return paragraph
    return None


def _format_range(
    paragraph: Any,
    start: int,
    end: int,
    font_name: str,
    size: float,
) -> None:
    if start >= end:
        return
    old_runs = list(paragraph.runs)
    fragments: list[tuple[str, Any, bool]] = []
    cursor = 0
    for run in old_runs:
        text = run.text
        run_start = cursor
        run_end = cursor + len(text)
        points = {run_start, run_end}
        if run_start < start < run_end:
            points.add(start)
        if run_start < end < run_end:
            points.add(end)
        boundaries = sorted(points)
        for left, right in zip(boundaries, boundaries[1:]):
            fragments.append(
                (
                    text[left - run_start : right - run_start],
                    run,
                    start <= left and right <= end,
                )
            )
        cursor = run_end
    for run in old_runs:
        paragraph._p.remove(run._r)
    for text, source_run, selected in fragments:
        run = paragraph.add_run(text)
        if source_run._r.rPr is not None:
            run._r.insert(0, deepcopy(source_run._r.rPr))
        if selected:
            run.font.name = font_name
            run.font.size = Pt(size)
            rfonts = run._element.get_or_add_rPr().get_or_add_rFonts()
            for key in ("ascii", "hAnsi", "eastAsia", "cs"):
                rfonts.set(qn(f"w:{key}"), font_name)
