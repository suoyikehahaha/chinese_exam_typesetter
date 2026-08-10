from __future__ import annotations

from pathlib import Path
import tempfile
import unittest

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from app.exam_format_rules import apply_exam_format_rules
from app.flexible_importers import import_exam
from app.semantic_formatting import apply_semantic_formatting
from app.source_decorations import restore_source_decorations


class SourcePreservationTests(unittest.TestCase):
    def test_qingdao_header_poetry_and_decorations(self) -> None:
        source = Path(
            r"D:\Desktop\青岛市2026年高三年级第三次适应性检测语文试题.docx"
        )
        if not source.exists():
            self.skipTest("Qingdao acceptance document is unavailable")
        exam = import_exam(source)
        self.assertEqual(
            exam["metadata"]["exam_name"],
            "青岛市2026年高三年级第三次适应性检测",
        )
        self.assertEqual(exam["metadata"]["meta_text"], "2026.05")
        poetry = next(block for block in exam["blocks"] if block["type"] == "poetry")
        self.assertEqual(poetry["author"], "方岳")
        self.assertTrue(poetry["note"].startswith("[注]"))
        self.assertGreater(len(exam["metadata"]["source_decorations"]), 0)

    def test_note_and_visible_option_indent(self) -> None:
        with tempfile.TemporaryDirectory() as folder:
            path = Path(folder) / "rules.docx"
            document = Document()
            document.add_paragraph("[注]方岳，南宋词人。")
            document.add_paragraph("A．一个较长的选择项。")
            document.save(path)
            apply_semantic_formatting(path)
            apply_exam_format_rules(path)
            result = Document(path)
            note, option = result.paragraphs
            self.assertEqual(
                note.runs[0]._r.xpath("./w:rPr/w:rFonts/@w:eastAsia")[0],
                "FangSong",
            )
            self.assertAlmostEqual(option.paragraph_format.left_indent.pt, 33.6)
            self.assertAlmostEqual(option.paragraph_format.first_line_indent.pt, -17.85)

    def test_source_underline_and_emphasis_are_restored(self) -> None:
        with tempfile.TemporaryDirectory() as folder:
            path = Path(folder) / "marks.docx"
            document = Document()
            document.add_paragraph("一段带标记的文字")
            document.save(path)
            raw = {
                "metadata": {
                    "source_decorations": [
                        {
                            "text": "一段带标记的文字",
                            "occurrence": 0,
                            "ranges": [
                                {"start": 2, "end": 5, "underline": "single"},
                                {"start": 5, "end": 7, "emphasis": "underDot"},
                            ],
                        }
                    ]
                }
            }
            restore_source_decorations(path, raw)
            paragraph = Document(path).paragraphs[0]
            underline_text = "".join(
                run.text
                for run in paragraph.runs
                if run._r.xpath("./w:rPr/w:u/@w:val")
            )
            emphasis_text = "".join(
                run.text
                for run in paragraph.runs
                if run._r.xpath("./w:rPr/w:em/@w:val")
            )
            self.assertEqual(underline_text, "带标记")
            self.assertEqual(emphasis_text, "的文")


if __name__ == "__main__":
    unittest.main()
