"""Final formatting pass for notice aliases and semantic material roles."""

from __future__ import annotations

from pathlib import Path
import re
from typing import Any, Iterator

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt

from .contextual_formatting_v6 import apply_contextual_formatting_v6


SOURCE_RE = re.compile(
    r"^\s*[（(]\s*(?:"
    r"(?:摘自|摘编自|摘选自|选自|节选自|改编自|据).+|"
    r"本文.+|有删改|有改动"
    r")[）)]\s*$"
)


def apply_contextual_formatting_v7(
    docx_path: str | Path,
    raw_exam: dict[str, Any],
) -> None:
    """Apply the v6 rules plus flexible notice and material-role formatting."""

    target = Path(docx_path)
    apply_contextual_formatting_v6(target, raw_exam)
    document = Document(target)
    metadata = raw_exam.get("metadata", {})
    _replace_notice_title(document, str(metadata.get("notice_title", "")))
    _remove_empty_exam_meta(document)
    _format_material_roles(document, raw_exam)
    for paragraph in _all_paragraphs(document):
        if SOURCE_RE.fullmatch(paragraph.text.strip()):
            _paragraph_style(paragraph, "FangSong", 10.5, "right")
    document.save(target)


def _replace_notice_title(document: Any, notice_title: str) -> None:
    if not notice_title:
        return
    for paragraph in document.paragraphs:
        if paragraph.style.name != "Exam_notice_title":
            continue
        paragraph.clear()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.left_indent = Pt(0)
        paragraph.paragraph_format.first_line_indent = Pt(0)
        run = paragraph.add_run(notice_title)
        _font(run, "SimHei", 10.5)
        return


def _remove_empty_exam_meta(document: Any) -> None:
    for paragraph in list(document.paragraphs):
        if (
            paragraph.style.name == "Exam_exam_meta"
            and not paragraph.text.strip()
        ):
            paragraph._element.getparent().remove(paragraph._element)


def _format_material_roles(document: Any, raw_exam: dict[str, Any]) -> None:
    role_texts: dict[str, set[str]] = {
        "subheading": set(),
        "source": set(),
        "label": set(),
    }
    for block in raw_exam.get("blocks", []):
        if block.get("type") != "material":
            continue
        paragraphs = [str(value).strip() for value in block.get("paragraphs", [])]
        roles = [str(value) for value in block.get("paragraph_roles", [])]
        for index, text in enumerate(paragraphs):
            if index < len(roles) and roles[index] in role_texts:
                role_texts[roles[index]].add(text)

    for paragraph in _all_paragraphs(document):
        text = paragraph.text.strip()
        if text in role_texts["subheading"]:
            _paragraph_style(paragraph, "SimHei", 10.5, "center")
        elif text in role_texts["source"]:
            _paragraph_style(paragraph, "FangSong", 10.5, "right")
        elif text in role_texts["label"]:
            _paragraph_style(paragraph, "SimHei", 10.5, "left")


def _paragraph_style(
    paragraph: Any,
    font_name: str,
    size: float,
    alignment: str,
) -> None:
    paragraph.alignment = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
    }[alignment]
    paragraph.paragraph_format.left_indent = Pt(0)
    paragraph.paragraph_format.first_line_indent = Pt(0)
    for run in paragraph.runs:
        _font(run, font_name, size)


def _font(run: Any, name: str, size: float) -> None:
    run.font.name = name
    run.font.size = Pt(size)
    rfonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    for key in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn(f"w:{key}"), name)


def _all_paragraphs(document: Any) -> Iterator[Any]:
    yield from document.paragraphs
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs


__all__ = ["apply_contextual_formatting_v7"]
