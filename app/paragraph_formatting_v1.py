"""Apply whole-block and selected-paragraph formatting overrides."""

from __future__ import annotations

from pathlib import Path
import re
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt


QUESTION_RE = re.compile(r"^(\d{1,2})．")
FONT_NAMES = {
    "宋体": "SimSun",
    "黑体": "SimHei",
    "楷体": "KaiTi",
    "仿宋": "FangSong",
}
ALIGNMENTS = {
    "左对齐": WD_ALIGN_PARAGRAPH.LEFT,
    "居中": WD_ALIGN_PARAGRAPH.CENTER,
    "右对齐": WD_ALIGN_PARAGRAPH.RIGHT,
    "两端对齐": WD_ALIGN_PARAGRAPH.JUSTIFY,
}


def apply_paragraph_formats_v1(
    docx_path: str | Path,
    raw_exam: dict[str, Any],
) -> None:
    """Apply structure-wide formats followed by selected-line formats."""

    target = Path(docx_path)
    document = Document(target)
    for block in raw_exam.get("blocks", []):
        if block.get("type") == "question":
            question = block.get("question", {})
            stem_paragraph = _find_question(document, int(question.get("number", 0)))
            spec = question.get("format")
            if stem_paragraph is not None and spec:
                _apply(stem_paragraph, spec)
            option_spec = _option_spec(spec or {})
            if option_spec:
                for option in question.get("options", []):
                    paragraph = _find_exact(document, str(option))
                    if paragraph is not None:
                        _apply(paragraph, option_spec)
            for entry in question.get("paragraph_formats", []):
                paragraph, _prefix = _locate(document, block, entry)
                if paragraph is not None:
                    _apply(paragraph, entry)
        else:
            spec = block.get("format")
            if spec:
                for text in _block_lines(block):
                    paragraph = _find_exact(document, text)
                    if paragraph is not None:
                        _apply(paragraph, spec)
            for entry in block.get("paragraph_formats", []):
                paragraph, _prefix = _locate(document, block, entry)
                if paragraph is not None:
                    _apply(paragraph, entry)
    document.save(target)


def _option_spec(spec: dict[str, Any]) -> dict[str, Any]:
    if not spec:
        return {}
    text_before = _number(spec.get("option_left_indent_chars"), 1.5)
    hanging = _number(spec.get("option_hanging_indent_chars"), 1.7)
    return {
        "font": spec.get("option_font", "宋体"),
        "size_pt": spec.get("option_size_pt", 10.5),
        "bold": spec.get("option_bold", False),
        "left_indent_chars": text_before,
        "special_indent": "悬挂",
        "special_indent_chars": hanging,
        "alignment": spec.get("option_alignment", "左对齐"),
        "line_spacing": spec.get("option_line_spacing", 1.25),
        "space_before_pt": spec.get("option_space_before_pt", 0),
        "space_after_pt": spec.get("option_space_after_pt", 0),
    }


def _apply(paragraph: Any, spec: dict[str, Any]) -> None:
    font_name = FONT_NAMES.get(str(spec.get("font", "")), str(spec.get("font", "")))
    size = _number(spec.get("size_pt"), 10.5)
    bold_value = spec.get("bold")
    if font_name:
        for run in paragraph.runs:
            _font(run, font_name, size, bold_value)

    left = _number(
        spec.get("left_indent_chars"),
        _number(spec.get("text_before_chars"), 0),
    )
    right = _number(spec.get("right_indent_chars"), 0)
    special = str(spec.get("special_indent", ""))
    amount = _number(
        spec.get("special_indent_chars"),
        abs(_number(spec.get("first_line_indent_chars"), 0)),
    )
    if not special:
        first = _number(spec.get("first_line_indent_chars"), 0)
        special = "首行" if first > 0 else ("悬挂" if first < 0 else "无")
        amount = abs(first)

    if special == "悬挂":
        paragraph.paragraph_format.left_indent = Pt(size * (left + amount))
        paragraph.paragraph_format.first_line_indent = Pt(-size * amount)
    else:
        paragraph.paragraph_format.left_indent = Pt(size * left)
        paragraph.paragraph_format.first_line_indent = Pt(
            size * amount if special == "首行" else 0
        )
    paragraph.paragraph_format.right_indent = Pt(size * right)
    paragraph.paragraph_format.line_spacing = _number(
        spec.get("line_spacing"),
        1.25,
    )
    paragraph.paragraph_format.space_before = Pt(
        _number(spec.get("space_before_pt"), 0)
    )
    paragraph.paragraph_format.space_after = Pt(
        _number(spec.get("space_after_pt"), 0)
    )
    paragraph.paragraph_format.keep_with_next = bool(
        spec.get("keep_with_next", False)
    )
    paragraph.paragraph_format.page_break_before = bool(
        spec.get("page_break_before", False)
    )
    alignment = str(spec.get("alignment", "左对齐"))
    if alignment in ALIGNMENTS:
        paragraph.alignment = ALIGNMENTS[alignment]


def _font(run: Any, name: str, size: float, bold: Any) -> None:
    run.font.name = name
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bool(bold)
    rfonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    for key in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn(f"w:{key}"), name)


def _locate(
    document: Any,
    block: dict[str, Any],
    entry: dict[str, Any],
) -> tuple[Any | None, int]:
    target = str(entry.get("target", "stem"))
    index = int(entry.get("target_index", 0))
    if block.get("type") != "question":
        lines = _block_lines(block)
        return (
            (_find_exact(document, lines[index]), 0)
            if index < len(lines)
            else (None, 0)
        )
    question = block["question"]
    if target == "stem":
        number = int(question.get("number", 0))
        return _find_question(document, number), len(f"{number}．")
    if target == "option":
        values = question.get("options", [])
        return (
            (_find_exact(document, str(values[index])), 0)
            if index < len(values)
            else (None, 0)
        )
    if target == "embedded":
        values = question.get("embedded_segments", [])
        text = (
            "".join(str(segment.get("text", "")) for segment in values[index])
            if index < len(values)
            else ""
        )
        return _find_exact(document, text), 0
    if target == "segmentation":
        return _find_exact(document, str(question.get("segmentation_text", ""))), 0
    if target == "subquestion":
        values = question.get("subquestions", [])
        text = str(values[index]) if index < len(values) else ""
        marker = f"（{index + 1}）"
        return _find_exact(document, marker + text), len(marker)
    values = question.get(target, [])
    text = str(values[index]) if index < len(values) else ""
    return _find_exact(document, text), 0


def _block_lines(block: dict[str, Any]) -> list[str]:
    if block.get("type") in {"section_title", "instruction"}:
        return [str(block.get("text", ""))]
    if block.get("type") == "subsection":
        return [str(block.get("name", "")) + str(block.get("meta", ""))]
    values: list[str] = []
    for key in ("title", "author"):
        if block.get(key):
            values.append(str(block[key]))
    values.extend(str(value) for value in block.get("paragraphs", []))
    for key in ("note", "source"):
        if block.get(key):
            values.append(str(block[key]))
    return values


def _find_question(document: Any, number: int) -> Any | None:
    for paragraph in document.paragraphs:
        match = QUESTION_RE.match(paragraph.text.strip())
        if match and int(match.group(1)) == number:
            return paragraph
    return None


def _find_exact(document: Any, text: str) -> Any | None:
    for paragraph in document.paragraphs:
        if paragraph.text.strip() == text.strip():
            return paragraph
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if paragraph.text.strip() == text.strip():
                        return paragraph
    return None


def _number(value: Any, default: float) -> float:
    try:
        return float(value)
    except (TypeError, ValueError):
        return default
