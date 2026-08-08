"""Format variable-length classical Chinese segmentation markers."""

from __future__ import annotations

from pathlib import Path
import re
from typing import Any, Iterator

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


def apply_segmentation_formatting_v2(docx_path: str | Path) -> None:
    """Box every uppercase marker present and preserve the source sequence."""

    target = Path(docx_path)
    document = Document(target)
    for paragraph in _all_paragraphs(document):
        if paragraph.style.name != "Exam_segmentation_text":
            continue
        text = paragraph.text
        for run in list(paragraph.runs):
            paragraph._p.remove(run._r)
        for token in re.split(r"([A-Z])", text):
            if not token:
                continue
            marker = len(token) == 1 and token.isascii() and token.isupper()
            run = paragraph.add_run(token)
            _font(run, "SimSun" if marker else "KaiTi", 10.5)
            if marker:
                border = OxmlElement("w:bdr")
                border.set(qn("w:val"), "single")
                border.set(qn("w:sz"), "4")
                border.set(qn("w:space"), "1")
                border.set(qn("w:color"), "auto")
                run._element.get_or_add_rPr().append(border)
    document.save(target)


def _font(run: Any, name: str, size: float) -> None:
    run.font.name = name
    run.font.size = Pt(size)
    rfonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    for key in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn(f"w:{key}"), name)


def _all_paragraphs(document: Any) -> Iterator[Any]:
    yield from document.paragraphs
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs
