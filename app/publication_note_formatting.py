"""Formatting pass for parenthetical article-end publication notes."""

from __future__ import annotations

from pathlib import Path
from typing import Any, Iterator

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt

from .composition_formatting import apply_composition_formatting


def apply_publication_note_formatting(
    docx_path: str | Path,
    raw_exam: dict[str, Any],
) -> None:
    """Apply SimSun 10.5 pt and right alignment to publication notes."""

    target = Path(docx_path)
    apply_composition_formatting(target, raw_exam)
    expected: set[str] = set()
    for block in raw_exam.get("blocks", []):
        if block.get("type") != "material":
            continue
        paragraphs = [str(value).strip() for value in block.get("paragraphs", [])]
        roles = [str(value) for value in block.get("paragraph_roles", [])]
        for index, text in enumerate(paragraphs):
            if index < len(roles) and roles[index] == "publication_note":
                expected.add(text)
    if not expected:
        return

    document = Document(target)
    for paragraph in _all_paragraphs(document):
        if paragraph.text.strip() not in expected:
            continue
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        paragraph.paragraph_format.left_indent = Pt(0)
        paragraph.paragraph_format.right_indent = Pt(0)
        paragraph.paragraph_format.first_line_indent = Pt(0)
        for run in paragraph.runs:
            run.font.name = "SimSun"
            run.font.size = Pt(10.5)
            fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
            for key in ("ascii", "hAnsi", "eastAsia", "cs"):
                fonts.set(qn(f"w:{key}"), "SimSun")
    document.save(target)


def _all_paragraphs(document: Any) -> Iterator[Any]:
    yield from document.paragraphs
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs


__all__ = ["apply_publication_note_formatting"]
