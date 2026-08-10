from __future__ import annotations

import tempfile
import unittest
from pathlib import Path

from docx import Document
from PIL import Image

from app.flexible_importers import _collect_source_decorations
from app.flexible_importers import _refine_exam_name
from app.internal_preview_core import _material, render_internal_preview
from desktop_app import CurrentDesktopApp


class _Value:
    def __init__(self, value: object = "") -> None:
        self.value = value

    def set(self, value: object) -> None:
        self.value = value

    def get(self) -> object:
        return self.value


class PreviewRegressionTests(unittest.TestCase):
    def test_option_template_reset_updates_question_and_controls(self) -> None:
        question = {"option_layout": "two_column", "format": {}}
        app = object.__new__(CurrentDesktopApp)
        app.selected_block_index = 0
        app.raw_exam = {"blocks": [{"type": "question", "question": question}]}
        app.option_layout_var = _Value()
        app.option_font_var = _Value()
        app.option_size_var = _Value()
        app.option_left_var = _Value()
        app.option_hanging_var = _Value()
        app.status_var = _Value()
        app.loading_fields = False
        app._push_direct_history = lambda: None
        app._schedule_canvas_preview = lambda: None

        app._restore_options_group()

        self.assertEqual(question["option_layout"], "vertical")
        self.assertEqual(question["format"]["option_font"], "宋体")
        self.assertEqual(question["format"]["option_left_indent_chars"], 1.5)
        self.assertEqual(question["format"]["option_hanging_indent_chars"], 1.7)
        self.assertEqual(app.option_layout_var.get(), "四行单列")

    def test_global_template_reset_preserves_imported_metadata_text(self) -> None:
        app = object.__new__(CurrentDesktopApp)
        app.raw_exam = {
            "metadata": {
                "exam_name": "导入试卷名称",
                "subject_name": "语　文",
                "meta_text": "导入说明",
                "target_pages": 12,
            }
        }
        app.layout_path = Path("templates/layout.yaml")
        app._global_commit_after = None
        app.target_pages_var = _Value()
        app.margin_top_var = _Value()
        app.margin_bottom_var = _Value()
        app.margin_left_var = _Value()
        app.margin_right_var = _Value()
        app.status_var = _Value()
        app.loading_fields = False
        app._push_direct_history = lambda: None
        app._schedule_canvas_preview = lambda: None

        app._restore_global_group()

        metadata = app.raw_exam["metadata"]
        self.assertEqual(metadata["exam_name"], "导入试卷名称")
        self.assertEqual(metadata["meta_text"], "导入说明")
        self.assertEqual(metadata["target_pages"], 8)
        self.assertEqual(metadata["page_overrides"]["margin_top_mm"], 20.0)

    def test_notice_item_is_not_used_as_exam_title(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            source = Path(directory) / "header.docx"
            document = Document()
            for value in (
                "高三语文学科练习",
                "考试时间：150分钟 试卷满分：150分",
                "注意事项：",
                "1.考试结束后，将答题卡交回。",
                "语文",
                "一、阅读（72分）",
            ):
                document.add_paragraph(value)
            document.save(source)

            result: dict[str, object] = {}
            _refine_exam_name(source, result)

        self.assertEqual(result["metadata"]["exam_name"], "高三语文学科练习")


    def test_source_collector_keeps_bold_and_underline(self) -> None:
        document = Document()
        paragraph = document.add_paragraph()
        run = paragraph.add_run("正文标记")
        run.bold = True
        run.underline = True

        entries = _collect_source_decorations(document)

        self.assertEqual(len(entries), 1)
        self.assertTrue(entries[0]["ranges"][0]["bold"])
        self.assertEqual(entries[0]["ranges"][0]["underline"], "single")

    def test_internal_preview_renders_native_drawing_marker(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            folder = Path(directory)
            image_path = folder / "source.png"
            Image.new("RGB", (80, 40), (220, 30, 30)).save(image_path)
            source = folder / "source.docx"
            document = Document()
            paragraph = document.add_paragraph()
            paragraph.add_run().add_picture(str(image_path))
            document.save(source)

            raw = {
                "metadata": {
                    "source_docx_path": str(source),
                    "native_objects": [
                        {
                            "kind": "drawing",
                            "source_paragraph_index": 0,
                            "marker": "[[NATIVE_DRAWING:0]]",
                        }
                    ],
                },
                "blocks": [
                    {
                        "type": "material",
                        "paragraphs": ["[[NATIVE_DRAWING:0]]"],
                    }
                ],
            }
            output = folder / "preview"
            result = render_internal_preview(raw, Path("templates/layout.yaml"), output)

            self.assertEqual(result.actual_pages, 1)
            with Image.open(result.pages[0]).convert("RGB") as rendered:
                red_pixels = sum(
                    1
                    for red, green, blue in rendered.get_flattened_data()
                    if red > 150 and green < 80 and blue < 80
                )
            self.assertGreater(red_pixels, 100)

    def test_material_label_is_left_aligned_in_internal_preview(self) -> None:
        calls: list[tuple[tuple[object, ...], dict[str, object]]] = []

        def capture(*args: object, **kwargs: object) -> None:
            calls.append((args, kwargs))

        _material(
            {"paragraphs": ["材料一：", "正文内容。"]},
            0,
            capture,
        )

        self.assertEqual(calls[0][0][1], "材料一：")
        self.assertEqual(calls[0][1]["align"], "left")


if __name__ == "__main__":
    unittest.main()
