from __future__ import annotations

from pathlib import Path

from docx import Document

from app.regional_header_formatting import apply_regional_header_formatting
from app.flexible_importers import (
    _normalize_composition_prompts,
    _normalize_multi_poetry,
    _normalize_subsections,
)


def test_lifts_ascii_parenthesized_subsection_from_previous_question() -> None:
    raw = {
        "blocks": [
            {
                "type": "question",
                "question": {
                    "number": 9,
                    "embedded_segments": [
                        [{"text": "(三)阅读Ⅲ(本题共5小题，22分)", "role": "body"}]
                    ],
                },
            },
            {"type": "instruction", "text": "阅读下面的文言文。"},
        ]
    }
    _normalize_subsections(raw)
    assert [block["type"] for block in raw["blocks"]] == [
        "question",
        "subsection",
        "instruction",
    ]
    assert raw["blocks"][1]["name"] == "(三)阅读Ⅲ"
    assert raw["blocks"][1]["meta"] == "(本题共5小题，22分)"
    assert raw["blocks"][0]["question"]["embedded_segments"] == []


def test_splits_two_poems_with_dynasty_prefixed_authors() -> None:
    raw = {
        "blocks": [
            {"type": "instruction", "text": "阅读下面两首诗，完成各题。"},
            {
                "type": "material",
                "paragraphs": [
                    "[甲]第一首",
                    "(金)甲乙丙",
                    "甲诗正文。",
                    "[乙]第二首",
                    "（宋）丁戊己",
                    "乙诗正文。",
                ],
            },
        ]
    }
    _normalize_multi_poetry(raw)
    assert [block["type"] for block in raw["blocks"]] == [
        "instruction",
        "poetry",
        "poetry",
    ]
    assert raw["blocks"][1]["title"] == "[甲]第一首"
    assert raw["blocks"][1]["author"] == "(金)甲乙丙"
    assert raw["blocks"][2]["title"] == "[乙]第二首"


def test_moves_task_sentence_into_composition_prompt() -> None:
    task = "请你以指定主题发表演讲。写一篇演讲稿，谈谈你的感悟与思考。"
    raw = {
        "blocks": [
            {
                "type": "question",
                "question": {
                    "number": 23,
                    "composition_material": ["核心材料。", task],
                    "composition_prompt": [],
                },
            }
        ]
    }
    _normalize_composition_prompts(raw)
    question = raw["blocks"][0]["question"]
    assert question["composition_material"] == ["核心材料。"]
    assert question["composition_prompt"] == [task]


def test_contextual_formatter_keeps_subsection_meta_unbolded(
    tmp_path: Path,
) -> None:
    source = tmp_path / "input.docx"
    document = Document()
    document.styles.add_style("Exam_subsection", 1)
    paragraph = document.add_paragraph(style="Exam_subsection")
    paragraph.add_run("(二)阅读Ⅱ")
    paragraph.add_run("(本题共4小题，16分)")
    document.save(source)

    apply_regional_header_formatting(source, {"metadata": {}})
    result = Document(source)
    paragraph = result.paragraphs[0]
    assert paragraph.runs[0].bold is True
    assert paragraph.runs[1].bold is False
    assert paragraph.runs[0].font.name == "SimSun"
    assert paragraph.runs[1].font.name == "SimSun"
