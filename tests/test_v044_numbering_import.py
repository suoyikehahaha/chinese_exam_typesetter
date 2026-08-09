from __future__ import annotations

from pathlib import Path
import tempfile
import unittest

from docx import Document

from app.flexible_importers_v13 import (
    NumberingLevel,
    _format_prefix,
    _numbering_definitions,
    materialize_automatic_numbering,
)


class NumberingImportV044Tests(unittest.TestCase):
    def test_list_number_style_is_materialized(self) -> None:
        document = Document()
        first = document.add_paragraph("下列说法正确的一项是（3分）", style="List Number")
        second = document.add_paragraph("请概括材料内容。（4分）", style="List Number")

        inserted = materialize_automatic_numbering(document)

        self.assertEqual(inserted, 2)
        self.assertTrue(first.text.startswith("1．"))
        self.assertTrue(second.text.startswith("2．"))

    def test_prefix_formats_match_exam_parser(self) -> None:
        self.assertEqual(_format_prefix(NumberingLevel("decimal", "%1.", 10), 10), "10．")
        self.assertEqual(_format_prefix(NumberingLevel("upperLetter", "%1.", 1), 2), "B．")
        self.assertEqual(_format_prefix(NumberingLevel("decimal", "(%1)", 1), 3), "（3）")

    def test_saved_original_remains_unchanged(self) -> None:
        with tempfile.TemporaryDirectory() as folder:
            source = Path(folder) / "source.docx"
            document = Document()
            document.add_paragraph("题干文字（3分）", style="List Number")
            document.save(source)
            loaded = Document(source)
            materialize_automatic_numbering(loaded)
            self.assertEqual(Document(source).paragraphs[0].text, "题干文字（3分）")

    def test_missing_numbering_relationship_is_safe(self) -> None:
        class MissingNumberingDocument:
            class Part:
                @property
                def numbering_part(self) -> object:
                    raise KeyError("numbering")

            part = Part()

        self.assertEqual(_numbering_definitions(MissingNumberingDocument()), {})


if __name__ == "__main__":
    unittest.main()
