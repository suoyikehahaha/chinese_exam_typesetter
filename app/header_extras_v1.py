"""Render optional header lines that are outside the standard title fields."""

from __future__ import annotations

from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.text.paragraph import Paragraph


def apply_header_extras_v1(
    docx_path: str | Path,
    raw_exam: dict[str, Any],
) -> None:
    """Place confidentiality text before the formal exam title."""

    text = str(
        raw_exam.get("metadata", {}).get("confidentiality_text", "")
    ).strip()
    if not text:
        return
    target = Path(docx_path)
    document = Document(target)
    if not document.paragraphs:
        return
    element = OxmlElement("w:p")
    document.paragraphs[0]._p.addprevious(element)
    paragraph = Paragraph(element, document._body)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    run.font.name = "SimSun"
    run.font.size = Pt(10.5)
    rfonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    for key in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn(f"w:{key}"), "SimSun")
    document.save(target)
