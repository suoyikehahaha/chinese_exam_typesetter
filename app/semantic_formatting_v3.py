"""材料、文本标签及广义出处的语义格式。"""

from __future__ import annotations

from pathlib import Path
import re
from typing import Any, Iterator

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt


LABEL_RE = re.compile(
    r"^((?:材料|文本)(?:[一二三四五六七八九十]+|\d+)\s*[：:])\s*(.*)$"
)
SOURCE_RE = re.compile(
    r"^\s*[（(]\s*(?:(?:摘编自|选自|节选自|改编自|据).+|本文.+)[）)]\s*$"
)


def apply_semantic_formatting_v3(docx_path: str | Path) -> None:
    """应用标签黑体五号和出处仿宋五号右对齐。"""

    target = Path(docx_path)
    document = Document(target)
    for paragraph in _all_paragraphs(document):
        text = paragraph.text.strip()
        match = LABEL_RE.match(text)
        if match:
            _format_label(paragraph, match.group(1), match.group(2))
        elif SOURCE_RE.match(text):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            paragraph.paragraph_format.first_line_indent = Pt(0)
            paragraph.paragraph_format.left_indent = Pt(0)
            for run in paragraph.runs:
                _font(run, "FangSong", 10.5)
    document.save(target)


def _format_label(paragraph: Any, label: str, body: str) -> None:
    for run in list(paragraph.runs):
        paragraph._p.remove(run._r)
    label_run = paragraph.add_run(label)
    _font(label_run, "SimHei", 10.5)
    if body:
        body_run = paragraph.add_run(body)
        _font(body_run, "KaiTi", 10.5)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT


def _font(run: Any, name: str, size: float) -> None:
    run.font.name = name
    run.font.size = Pt(size)
    rfonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    for key in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn(f"w:{key}"), name)


def _all_paragraphs(document: Any) -> Iterator[Any]:
    yield from document.paragraphs
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs
