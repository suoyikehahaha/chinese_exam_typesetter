"""面向中文试卷的段落级禁则与标点溢出设置。"""

from __future__ import annotations

from pathlib import Path
from typing import Any, Iterator

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def enable_chinese_typography(docx_path: str | Path) -> None:
    """逐段启用中文禁则，允许行末标点悬挂到边界外。"""

    target = Path(docx_path)
    document = Document(target)
    settings = document.settings.element
    compat = settings.find(qn("w:compat"))
    if compat is None:
        compat = OxmlElement("w:compat")
        settings.append(compat)
    _on(compat, "kinsoku")
    _on(compat, "overflowPunct")
    theme_lang = settings.find(qn("w:themeFontLang"))
    if theme_lang is None:
        theme_lang = OxmlElement("w:themeFontLang")
        settings.append(theme_lang)
    theme_lang.set(qn("w:eastAsia"), "zh-CN")
    spacing = settings.find(qn("w:characterSpacingControl"))
    if spacing is None:
        spacing = OxmlElement("w:characterSpacingControl")
        settings.append(spacing)
    spacing.set(qn("w:val"), "compressPunctuation")

    for paragraph in _all_paragraphs(document):
        ppr = paragraph._p.get_or_add_pPr()
        _on(ppr, "kinsoku")
        _on(ppr, "overflowPunct")
        _on(ppr, "wordWrap")
        for run in paragraph.runs:
            rpr = run._element.get_or_add_rPr()
            lang = rpr.find(qn("w:lang"))
            if lang is None:
                lang = OxmlElement("w:lang")
                rpr.append(lang)
            lang.set(qn("w:eastAsia"), "zh-CN")
    document.save(target)


def _on(parent: Any, name: str) -> None:
    element = parent.find(qn(f"w:{name}"))
    if element is None:
        element = OxmlElement(f"w:{name}")
        parent.append(element)
    element.set(qn("w:val"), "1")


def _all_paragraphs(document: Any) -> Iterator[Any]:
    yield from document.paragraphs
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs
