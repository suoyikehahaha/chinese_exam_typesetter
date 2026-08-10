from __future__ import annotations

from pathlib import Path
import tempfile
import unittest

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.exam_context_formatting import apply_exam_context_formatting
from app.flexible_importers import import_exam
from app.models import ExamDocument
from app.validators import validate_exam


class AdaptiveImportTests(unittest.TestCase):
    def test_marker_gap_is_non_blocking_warning(self) -> None:
        raw = {
            "metadata": {
                "exam_name": "测试",
                "subject_name": "语文",
                "meta_text": "",
                "total_score": 3,
                "notices": [],
            },
            "blocks": [
                {
                    "type": "question",
                    "question": {
                        "number": 1,
                        "kind": "subjective",
                        "stem": "断句。",
                        "score": 3,
                        "options": [],
                        "segmentation_text": "甲A乙B丙D丁",
                    },
                }
            ],
        }
        issues = validate_exam(ExamDocument.from_dict(raw))
        marker_issues = [item for item in issues if item.code.startswith("segmentation")]
        self.assertEqual(marker_issues[0].severity, "warning")
        self.assertIn("缺少C", marker_issues[0].message)

    def test_qingdao_title_author_and_metadata(self) -> None:
        source = Path(
            r"D:\Desktop\青岛市2026年高三年级第三次适应性检测语文试题.docx"
        )
        if not source.exists():
            self.skipTest("Qingdao acceptance document is unavailable")
        exam = import_exam(source)
        material = next(
            block
            for block in exam["blocks"]
            if block.get("type") == "material"
            and block.get("title") == "抢炮蒙"
        )
        self.assertEqual(material["author"], "福森")
        self.assertEqual(exam["metadata"]["meta_text"], "2026.05")

    def test_header_source_and_memorization_formatting(self) -> None:
        with tempfile.TemporaryDirectory() as folder:
            path = Path(folder) / "context.docx"
            document = Document()
            document.add_paragraph("2026.05")
            document.add_paragraph("（摘自杨成武《长征胜利万岁》）")
            document.add_paragraph("班级活动中用（1）“____，____”表达感受。")
            document.save(path)
            raw = {
                "metadata": {"meta_text": "2026.05"},
                "blocks": [
                    {"type": "subsection", "name": "（五）名篇名句默写"},
                    {
                        "type": "question",
                        "question": {
                            "stem": "补写出下列句子中的空缺部分。",
                            "embedded_segments": [
                                [
                                    {
                                        "text": "班级活动中用（1）“____，____”表达感受。",
                                        "role": "body",
                                    }
                                ]
                            ],
                        },
                    },
                ],
            }
            apply_exam_context_formatting(path, raw)
            date, source, memorization = Document(path).paragraphs
            self.assertEqual(date.alignment, WD_ALIGN_PARAGRAPH.RIGHT)
            self.assertEqual(
                source.runs[0]._r.xpath("./w:rPr/w:rFonts/@w:eastAsia")[0],
                "FangSong",
            )
            self.assertEqual(
                memorization.runs[0]._r.xpath("./w:rPr/w:rFonts/@w:eastAsia")[0],
                "SimSun",
            )


if __name__ == "__main__":
    unittest.main()
