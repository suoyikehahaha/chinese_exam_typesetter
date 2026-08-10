from __future__ import annotations

from pathlib import Path
import tempfile
import unittest

from docx import Document

from app.regional_header_formatting import apply_regional_header_formatting
from app.flexible_importers import (
    _normalize_composition_prompts,
    _normalize_multi_poetry,
    _normalize_subsections,
)


class FormattingRegressionTests(unittest.TestCase):
    def test_lifts_ascii_parenthesized_subsection(self) -> None:
        raw = {
            "blocks": [
                {
                    "type": "question",
                    "question": {
                        "number": 9,
                        "embedded_segments": [
                            [
                                {
                                    "text": "(三)阅读Ⅲ(本题共5小题，22分)",
                                    "role": "body",
                                }
                            ]
                        ],
                    },
                },
                {"type": "instruction", "text": "阅读下面的文言文。"},
            ]
        }
        _normalize_subsections(raw)
        self.assertEqual(
            [block["type"] for block in raw["blocks"]],
            ["question", "subsection", "instruction"],
        )
        self.assertEqual(raw["blocks"][1]["name"], "(三)阅读Ⅲ")
        self.assertEqual(raw["blocks"][1]["meta"], "(本题共5小题，22分)")
        self.assertEqual(raw["blocks"][0]["question"]["embedded_segments"], [])

    def test_splits_two_poems_with_dynasty_prefixed_authors(self) -> None:
        raw = {
            "blocks": [
                {"type": "instruction", "text": "阅读下面两首诗，完成各题。"},
                {
                    "type": "material",
                    "paragraphs": [
                        "[甲]第一首",
                        "(金)甲乙丙",
                        "甲诗正文。",
                        "[乙]第二首",
                        "（宋）丁戊己",
                        "乙诗正文。",
                    ],
                },
            ]
        }
        _normalize_multi_poetry(raw)
        self.assertEqual(
            [block["type"] for block in raw["blocks"]],
            ["instruction", "poetry", "poetry"],
        )
        self.assertEqual(raw["blocks"][1]["title"], "[甲]第一首")
        self.assertEqual(raw["blocks"][1]["author"], "(金)甲乙丙")
        self.assertEqual(raw["blocks"][2]["title"], "[乙]第二首")

    def test_moves_task_sentence_into_composition_prompt(self) -> None:
        task = "请你以指定主题发表演讲。写一篇演讲稿，谈谈你的感悟与思考。"
        raw = {
            "blocks": [
                {
                    "type": "question",
                    "question": {
                        "number": 23,
                        "composition_material": ["核心材料。", task],
                        "composition_prompt": [],
                    },
                }
            ]
        }
        _normalize_composition_prompts(raw)
        question = raw["blocks"][0]["question"]
        self.assertEqual(question["composition_material"], ["核心材料。"])
        self.assertEqual(question["composition_prompt"], [task])

    def test_subsection_meta_remains_unbolded(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            source = Path(directory) / "input.docx"
            document = Document()
            document.styles.add_style("Exam_subsection", 1)
            paragraph = document.add_paragraph(style="Exam_subsection")
            paragraph.add_run("(二)阅读Ⅱ")
            paragraph.add_run("(本题共4小题，16分)")
            document.save(source)

            apply_regional_header_formatting(source, {"metadata": {}})
            paragraph = Document(source).paragraphs[0]
            self.assertIs(paragraph.runs[0].bold, True)
            self.assertIs(paragraph.runs[1].bold, False)
            self.assertEqual(paragraph.runs[0].font.name, "SimSun")
            self.assertEqual(paragraph.runs[1].font.name, "SimSun")


if __name__ == "__main__":
    unittest.main()
