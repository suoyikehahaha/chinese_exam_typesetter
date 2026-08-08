"""Flexible importer v10 for notices, section recovery and semantic materials."""

from __future__ import annotations

from pathlib import Path
import re
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from .flexible_importers_v9 import import_exam as import_exam_v9
from .flexible_importers_v9 import parse_plain_lines, save_exam


MAJOR_SECTION_RE = re.compile(
    r"^\s*[一二三四五六七八九十]+\s*[、，,]\s*(?P<body>.+?)\s*$"
)
NOTICE_TITLE_RE = re.compile(r"^\s*(注意事项|考生须知)\s*[：:]?\s*$")
NOTICE_ITEM_RE = re.compile(r"^\s*\d{1,2}\s*[．.、]\s*(?P<body>.+?)\s*$")
LABEL_RE = re.compile(r"^\s*(?:材料|文本)(?:[一二三四五六七八九十]+|\d+)\s*[：:]")
NOTE_RE = re.compile(
    r"^\s*(?:【\s*注\s*】|\[\s*注\s*\]|［\s*注\s*］|注\s*[：:])"
)
SOURCE_RE = re.compile(
    r"^\s*[（(]\s*(?:"
    r"(?:摘自|摘编自|摘选自|选自|节选自|改编自|据).+|"
    r"本文.+|有删改|有改动"
    r")[）)]\s*$"
)
AUTHOR_RE = re.compile(
    r"^\s*(?:(?:【[^】]{1,8}】|\[[^\]]{1,8}\]|［[^］]{1,8}］|"
    r"（[^）]{1,8}）|\([^)]{1,8}\))\s*)?"
    r"[\u3400-\u9fff·・、，,\s]{2,30}\s*$"
)
BODY_LEAD_RE = re.compile(r"^(?:[①-⑳]|\d|[“‘（(《]|\u3400-\u9fff)")
SECTION_CUES = ("阅读", "语言文字运用", "写作")
POETRY_CUES = ("诗", "词", "曲")


def import_exam(path: str | Path) -> dict[str, Any]:
    """Import common exam variants without tying structure to fixed question types."""

    source = Path(path)
    result = import_exam_v9(source)
    if source.suffix.lower() == ".docx":
        _normalize_header_notices(source, result)
    _lift_embedded_sections(result)
    _normalize_single_poetry(result)
    _normalize_prose_title_authors(result)
    if source.suffix.lower() == ".docx":
        _annotate_material_roles(source, result)
    else:
        _annotate_material_roles(None, result)
    return result


def _normalize_header_notices(source: Path, result: dict[str, Any]) -> None:
    """Treat ``考生须知`` and ``注意事项`` as equivalent optional notice headings."""

    document = Document(source)
    lines = [paragraph.text.strip() for paragraph in document.paragraphs]
    lines = [line for line in lines if line]
    first_section = next(
        (
            index
            for index, line in enumerate(lines)
            if _is_major_section(line)
        ),
        min(len(lines), 18),
    )
    header = lines[:first_section]
    notice_index = next(
        (
            index
            for index, line in enumerate(header)
            if NOTICE_TITLE_RE.fullmatch(line)
        ),
        None,
    )
    if notice_index is None:
        return

    match = NOTICE_TITLE_RE.fullmatch(header[notice_index])
    assert match is not None
    raw_notice_lines = header[notice_index + 1 :]
    notices: list[str] = []
    for line in raw_notice_lines:
        item = NOTICE_ITEM_RE.fullmatch(line)
        notices.append(item.group("body").strip() if item else line.strip())

    metadata = result.setdefault("metadata", {})
    metadata["notice_title"] = f"{match.group(1)}："
    metadata["notices"] = notices
    metadata["meta_text"] = _header_meta_before_notice(
        header[:notice_index],
        metadata,
    )
    exam_info = str(metadata.get("exam_info_text", "")).strip()
    if exam_info and exam_info in raw_notice_lines:
        metadata.pop("exam_info_text", None)


def _header_meta_before_notice(
    lines: list[str],
    metadata: dict[str, Any],
) -> str:
    excluded = {
        str(metadata.get("exam_name", "")).strip(),
        str(metadata.get("subject_name", "")).strip(),
        str(metadata.get("confidentiality_text", "")).strip(),
    }
    candidates = [
        line
        for line in lines
        if line not in excluded
        and ("满分" in line or "考试时间" in line or "考试用时" in line)
    ]
    return candidates[0] if candidates else ""


def _lift_embedded_sections(result: dict[str, Any]) -> None:
    """Move major section headings out of question or material payloads."""

    normalized: list[dict[str, Any]] = []
    for block in result.get("blocks", []):
        lifted: list[dict[str, str]] = []
        question = block.get("question")
        if question:
            kept: list[list[dict[str, Any]]] = []
            for segments in question.get("embedded_segments", []):
                text = "".join(
                    str(segment.get("text", "")) for segment in segments
                ).strip()
                if _is_major_section(text):
                    lifted.append({"type": "section_title", "text": text})
                else:
                    kept.append(segments)
            question["embedded_segments"] = kept
            normalized.append(block)
            normalized.extend(lifted)
            continue

        if block.get("type") == "material":
            paragraphs = [str(value).strip() for value in block.get("paragraphs", [])]
            if any(_is_major_section(text) for text in paragraphs):
                current: list[str] = []
                for text in paragraphs:
                    if _is_major_section(text):
                        if current:
                            item = dict(block)
                            item["paragraphs"] = current
                            normalized.append(item)
                            current = []
                        normalized.append({"type": "section_title", "text": text})
                    else:
                        current.append(text)
                if current:
                    item = dict(block)
                    item["paragraphs"] = current
                    normalized.append(item)
                continue
        normalized.append(block)
    result["blocks"] = normalized


def _is_major_section(text: str) -> bool:
    match = MAJOR_SECTION_RE.fullmatch(text)
    return bool(match and any(cue in match.group("body") for cue in SECTION_CUES))


def _normalize_single_poetry(result: dict[str, Any]) -> None:
    """Promote one-poem material blocks with dynasty-prefixed authors."""

    normalized: list[dict[str, Any]] = []
    previous_instruction = ""
    for block in result.get("blocks", []):
        if block.get("type") == "instruction":
            previous_instruction = str(block.get("text", ""))
            normalized.append(block)
            continue
        if block.get("type") != "material" or not any(
            cue in previous_instruction for cue in POETRY_CUES
        ):
            normalized.append(block)
            continue
        paragraphs = [str(value).strip() for value in block.get("paragraphs", [])]
        if len(paragraphs) < 3 or not _is_author_line(paragraphs[1]):
            normalized.append(block)
            continue
        body = paragraphs[2:]
        note = str(block.get("note", ""))
        if not note and body and NOTE_RE.match(body[-1]):
            note = body.pop()
        normalized.append(
            {
                "type": "poetry",
                "title": paragraphs[0],
                "author": paragraphs[1],
                "paragraphs": body,
                "note": note,
            }
        )
    result["blocks"] = normalized


def _normalize_prose_title_authors(result: dict[str, Any]) -> None:
    """Promote prose title and author lines, including multiple authors."""

    previous_instruction = ""
    for block in result.get("blocks", []):
        if block.get("type") == "instruction":
            previous_instruction = str(block.get("text", ""))
            continue
        if block.get("type") != "material":
            continue
        if block.get("title") or block.get("author"):
            _promote_trailing_source(block)
            continue
        if any(cue in previous_instruction for cue in POETRY_CUES):
            continue
        paragraphs = [str(value).strip() for value in block.get("paragraphs", [])]
        if len(paragraphs) < 3:
            continue
        title, author, first_body = paragraphs[:3]
        if (
            not title
            or len(title) > 40
            or LABEL_RE.match(title)
            or not _is_author_line(author)
            or not BODY_LEAD_RE.match(first_body)
        ):
            continue
        block["title"] = title
        block["author"] = author
        block["paragraphs"] = paragraphs[2:]
        _promote_trailing_source(block)


def _is_author_line(text: str) -> bool:
    value = text.strip()
    if not AUTHOR_RE.fullmatch(value):
        return False
    han_count = len(re.findall(r"[\u3400-\u9fff]", value))
    return 2 <= han_count <= 20 and not any(
        mark in value for mark in ("。", "！", "？", "；", "：", "“", "”")
    )


def _promote_trailing_source(block: dict[str, Any]) -> None:
    paragraphs = [str(value).strip() for value in block.get("paragraphs", [])]
    if not block.get("source") and paragraphs and SOURCE_RE.fullmatch(paragraphs[-1]):
        block["source"] = paragraphs.pop()
        block["paragraphs"] = paragraphs


def _annotate_material_roles(
    source: Path | None,
    result: dict[str, Any],
) -> None:
    """Record label, source and centered subheading roles for later rendering."""

    centered_bold: set[str] = set()
    if source is not None:
        for paragraph in Document(source).paragraphs:
            text = paragraph.text.strip()
            visible_runs = [run for run in paragraph.runs if run.text.strip()]
            if (
                text
                and len(text) <= 60
                and paragraph.alignment == WD_ALIGN_PARAGRAPH.CENTER
                and any(run.bold for run in visible_runs)
            ):
                centered_bold.add(text)

    for block in result.get("blocks", []):
        if block.get("type") != "material":
            continue
        paragraphs = [str(value).strip() for value in block.get("paragraphs", [])]
        roles: list[str] = []
        formats = list(block.get("paragraph_formats", []))
        offset = int(bool(block.get("title"))) + int(bool(block.get("author")))
        for index, text in enumerate(paragraphs):
            role = "body"
            spec: dict[str, Any] | None = None
            if LABEL_RE.match(text):
                role = "label"
                spec = _format_spec("黑体", "左对齐")
            elif SOURCE_RE.fullmatch(text):
                role = "source"
                spec = _format_spec("仿宋", "右对齐")
            elif text in centered_bold:
                role = "subheading"
                spec = _format_spec("黑体", "居中")
            roles.append(role)
            if spec is not None:
                spec["target_index"] = offset + index
                spec["semantic_role"] = role
                formats.append(spec)
        block["paragraph_roles"] = roles
        if formats:
            block["paragraph_formats"] = formats


def _format_spec(font: str, alignment: str) -> dict[str, Any]:
    return {
        "font": font,
        "size_pt": 10.5,
        "bold": False,
        "alignment": alignment,
        "left_indent_chars": 0,
        "special_indent": "无",
        "special_indent_chars": 0,
        "line_spacing": 1.25,
        "space_before_pt": 0,
        "space_after_pt": 0,
    }


__all__ = ["import_exam", "parse_plain_lines", "save_exam"]
