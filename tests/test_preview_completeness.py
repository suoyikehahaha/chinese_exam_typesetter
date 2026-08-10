from pathlib import Path
import tempfile
import unittest

from docx import Document

from app.internal_preview_native import (
    _render_table_image,
    _table_rows,
    preview_exam,
    render_internal_preview,
)


ROOT = Path(__file__).resolve().parents[1]


class PreviewCompletenessTests(unittest.TestCase):
    def test_native_table_keeps_every_cell_string(self) -> None:
        document = Document()
        table = document.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "组别"
        table.cell(0, 1).text = "诗歌选段"
        table.cell(1, 0).text = "第一组"
        table.cell(1, 1).text = "春风又绿江南岸"
        self.assertEqual(
            _table_rows(table),
            [["组别", "诗歌选段"], ["第一组", "春风又绿江南岸"]],
        )
        image = _render_table_image(table)
        self.assertGreater(image.width, 300)
        self.assertGreater(image.height, 40)

    def test_answer_blocks_are_converted_without_losing_text(self) -> None:
        raw = {
            "metadata": {},
            "blocks": [
                {
                    "type": "answer_question",
                    "number": 1,
                    "header": "1．答案：A",
                    "paragraphs": [
                        {"text": "解析：保留这一整段文字。", "role": "subjective_answer"}
                    ],
                },
                {"type": "answer_table", "source_table_index": 0},
            ],
        }
        prepared = preview_exam(raw)
        self.assertEqual(prepared["blocks"][0]["type"], "material")
        self.assertEqual(
            prepared["blocks"][0]["paragraphs"],
            ["1．答案：A", "解析：保留这一整段文字。"],
        )
        self.assertEqual(
            prepared["blocks"][1]["paragraphs"],
            ["[[NATIVE_TABLE:0]]"],
        )

    def test_renderer_locates_question_answer_and_table_blocks(self) -> None:
        with tempfile.TemporaryDirectory() as folder:
            root = Path(folder)
            source = root / "source.docx"
            document = Document()
            table = document.add_table(rows=2, cols=2)
            table.cell(0, 0).text = "组别"
            table.cell(0, 1).text = "内容"
            table.cell(1, 0).text = "一"
            table.cell(1, 1).text = "不得从预览中消失"
            document.save(source)
            raw = {
                "metadata": {
                    "exam_name": "预览完整性测试",
                    "subject_name": "语　文",
                    "source_docx_path": str(source),
                    "native_objects": [
                        {
                            "kind": "table",
                            "source_table_index": 0,
                            "marker": "[[NATIVE_TABLE:0]]",
                        }
                    ],
                },
                "target_pages": 8,
                "blocks": [
                    {
                        "type": "question",
                        "question": {
                            "number": 21,
                            "kind": "subjective",
                            "stem": "阅读表格，回答问题。",
                            "score": 4,
                            "options": [],
                            "embedded_segments": [
                                [{"text": "[[NATIVE_TABLE:0]]", "role": "body"}]
                            ],
                        },
                    },
                    {
                        "type": "answer_question",
                        "number": 21,
                        "header": "21．答案：示例",
                        "paragraphs": [{"text": "答案正文完整显示。"}],
                    },
                    {"type": "answer_table", "source_table_index": 0},
                ],
            }
            result = render_internal_preview(
                raw,
                ROOT / "templates" / "layout.yaml",
                root / "preview",
            )
            self.assertTrue({0, 1, 2}.issubset(set(result.locators)))
            self.assertTrue(all(path.is_file() for path in result.pages))


if __name__ == "__main__":
    unittest.main()
