"""Canonical context-aware importer for the supported Word document formats.

The module keeps saved-draft compatibility for JSON, TXT, and Markdown while
the desktop facade handles DOCX and legacy DOC conversion.
"""

from __future__ import annotations

import json
from dataclasses import dataclass
from pathlib import Path
import re
import tempfile
from typing import Any, Iterable, Iterator

from docx import Document
from docx.document import Document as DocumentObject
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph

from .answer_typesetting import (
    ANSWER_MARKER_RE,
    attach_answer_blocks,
    find_answer_start,
    is_standalone_answer_docx,
    parse_answer_docx,
    standalone_answer_model,
)
from .models.identity import ensure_block_ids




# merged from flexible_importers.py


QUESTION_RE = re.compile(r"^\s*(\d{1,2})\s*[\uff0e.\u3001]\s*(.+)$")

OPTION_START_RE = re.compile(r"^\s*[A-D]\s*[\uff0e.\u3001]")

INLINE_OPTION_RE = re.compile(
    r"([A-D])\s*[\uff0e.\u3001]\s*(.*?)(?=(?:[A-D])\s*[\uff0e.\u3001]|$)"
)

SECTION_RE = re.compile(r"^\s*[\u4e00\u4e8c\u4e09\u56db\u4e94\u516d\u4e03\u516b\u4e5d\u5341]+\u3001")

SUBSECTION_RE = re.compile(r"^\s*\uff08[\u4e00\u4e8c\u4e09\u56db\u4e94\u516d\u4e03\u516b\u4e5d\u5341]+\uff09")

SCORE_RE = re.compile(r"[\uff08(](\d+(?:\.\d+)?)\u5206[\uff09)]\s*$")

NOTICE_RE = re.compile(r"^\s*\u6ce8\u610f\u4e8b\u9879\s*[\uff1a:]?\s*$")

NOTICE_ITEM_RE = re.compile(r"^\s*\d{1,2}\s*[\uff0e.\u3001]\s*(.+)$")

QUESTION_CUES = (
    "\u4e0b\u5217",
    "\u8bf7",
    "\u7b80\u8981",
    "\u6982\u62ec",
    "\u5206\u6790",
    "\u6839\u636e",
    "\u5982\u4f55",
    "\u4e3a\u4ec0\u4e48",
    "\u54ea\u4e9b",
    "\u8865\u5199",
    "\u7ffb\u8bd1",
    "\u5199\u51fa",
    "\u9009\u51fa",
    "\u6307\u51fa",
    "\u5b8c\u6210",
    "\u9605\u8bfb\u4e0b\u9762",
    "\u6700\u6070\u5f53",
    "\u4e0d\u6b63\u786e",
    "\u6b63\u786e\u7684\u4e00\u9879",
)

def _docx_lines(document: DocumentObject) -> Iterator[str]:
    """\u6309\u6587\u6863\u987a\u5e8f\u63d0\u53d6\u6b63\u6587\u6bb5\u843d\u548c\u8868\u683c\u884c\u3002"""

    for child in document.element.body.iterchildren():
        if isinstance(child, CT_P):
            text = Paragraph(child, document).text.strip()
            if text:
                yield text
        elif isinstance(child, CT_Tbl):
            table = Table(child, document)
            for row in table.rows:
                cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
                text = "\t".join(item for item in cells if item)
                if text:
                    yield text

def parse_plain_lines(lines: Iterable[str], title: str = "\u9ad8\u4e2d\u8bed\u6587\u8bd5\u5377") -> dict[str, Any]:
    """\u4f7f\u7528\u5377\u9996\u72b6\u6001\u3001\u8fde\u7eed\u9898\u53f7\u548c\u9898\u578b\u8bc1\u636e\u89e3\u6790\u666e\u901a\u6587\u672c\u3002"""

    clean_lines = [str(line).strip() for line in lines if str(line).strip()]
    if not clean_lines:
        raise ValueError("\u6587\u4ef6\u4e2d\u6ca1\u6709\u53ef\u8bc6\u522b\u7684\u6587\u5b57\u3002")

    metadata = _metadata_from_header(clean_lines, title)
    blocks: list[dict[str, Any]] = []
    pending_material: list[str] = []
    current_question: dict[str, Any] | None = None
    questions_started = False
    expected_number = 1
    in_notice = False
    first_section_seen = False

    def flush_material() -> None:
        nonlocal pending_material
        if pending_material:
            blocks.append({"type": "material", "paragraphs": pending_material})
            pending_material = []

    for index, line in enumerate(clean_lines):
        if NOTICE_RE.match(line):
            in_notice = True
            current_question = None
            continue

        if SECTION_RE.match(line):
            flush_material()
            in_notice = False
            first_section_seen = True
            current_question = None
            blocks.append({"type": "section_title", "text": line})
            continue

        if in_notice:
            notice_match = NOTICE_ITEM_RE.match(line)
            if notice_match:
                metadata["notices"].append(notice_match.group(1))
                continue
            if SUBSECTION_RE.match(line) or line.startswith("\u9605\u8bfb\u4e0b\u9762"):
                in_notice = False
            elif not _looks_like_question(line, clean_lines, index):
                metadata["notices"].append(line)
                continue

        if SUBSECTION_RE.match(line):
            flush_material()
            current_question = None
            name, meta = _split_subsection(line)
            blocks.append({"type": "subsection", "name": name, "meta": meta})
            continue

        if line.startswith("\u9605\u8bfb\u4e0b\u9762"):
            flush_material()
            current_question = None
            blocks.append({"type": "instruction", "text": line})
            continue

        question_match = QUESTION_RE.match(line)
        if question_match:
            number = int(question_match.group(1))
            stem = question_match.group(2)
            confidence = _question_confidence(stem, clean_lines, index)
            should_accept = False
            if number == expected_number:
                should_accept = questions_started or confidence >= 3
            elif questions_started and number > expected_number and confidence >= 4:
                should_accept = True
            elif not questions_started and number == 1 and confidence >= 4:
                should_accept = True

            if should_accept:
                flush_material()
                score_match = SCORE_RE.search(stem)
                score = float(score_match.group(1)) if score_match else None
                if score_match:
                    stem = stem[: score_match.start()].rstrip()
                current_question = {
                    "number": number,
                    "kind": "subjective",
                    "stem": stem,
                    "score": score,
                    "options": [],
                }
                blocks.append({"type": "question", "question": current_question})
                questions_started = True
                expected_number = number + 1
                continue

        inline_options = _extract_options(line)
        if current_question is not None and inline_options:
            current_question["kind"] = "objective"
            current_question["options"].extend(inline_options)
            continue

        if current_question is not None:
            segments = current_question.setdefault("embedded_segments", [])
            segments.append([{"text": line, "role": "body"}])
        elif first_section_seen or not _is_header_line(line, metadata):
            pending_material.append(line)

    flush_material()
    question_numbers = [
        block["question"]["number"]
        for block in blocks
        if block.get("type") == "question"
    ]
    if not question_numbers:
        raise ValueError(
            "\u6ca1\u6709\u8bc6\u522b\u5230\u9898\u76ee\u3002\u5efa\u8bae\u68c0\u67e5\u9898\u53f7\u662f\u5426\u4f7f\u7528\u201c1\uff0e\u9898\u5e72\u201d\u5f62\u5f0f\uff0c"
            "\u6216\u5148\u4fdd\u5b58\u4e3a\u53ef\u590d\u5236\u6587\u5b57\u7684 DOCX\u3001PDF\u3002"
        )
    metadata["total_score"] = _score_total_or_default(blocks, metadata["total_score"])
    return {"metadata": metadata, "blocks": blocks}

def _metadata_from_header(lines: list[str], fallback_title: str) -> dict[str, Any]:
    first_section = next(
        (index for index, line in enumerate(lines) if SECTION_RE.match(line)),
        min(len(lines), 12),
    )
    header = lines[:first_section]
    exam_name = next(
        (
            line
            for line in header
            if ("\u8003\u8bd5" in line or "\u8bd5\u5377" in line) and not NOTICE_RE.match(line)
        ),
        fallback_title,
    )
    subject_name = next(
        (line for line in header if re.fullmatch(r"\s*\u8bed\s*\u6587\s*", line)),
        "\u8bed\u3000\u6587",
    )
    meta_text = next(
        (
            line
            for line in header
            if ("\u6ee1\u5206" in line or "\u8003\u8bd5\u65f6\u95f4" in line) and not QUESTION_RE.match(line)
        ),
        "\u8bf7\u5728\u5de6\u4fa7\u68c0\u67e5\u8bd5\u5377\u4fe1\u606f\u3001\u9898\u76ee\u5206\u503c\u548c\u7248\u5f0f\u53c2\u6570\u3002",
    )
    total_match = re.search(r"\u6ee1\u5206\s*(\d+(?:\.\d+)?)\s*\u5206", " ".join(header))
    total_score = float(total_match.group(1)) if total_match else 150
    return {
        "exam_name": exam_name,
        "subject_name": subject_name,
        "meta_text": meta_text,
        "total_score": total_score,
        "notices": [],
    }

def _question_confidence(stem: str, lines: list[str], index: int) -> int:
    score = 0
    if SCORE_RE.search(stem):
        score += 3
    if any(cue in stem for cue in QUESTION_CUES):
        score += 2
    if "\uff1f" in stem or stem.endswith("?"):
        score += 1
    if _extract_options(stem):
        score += 3
    for following in lines[index + 1 : index + 4]:
        if _extract_options(following):
            score += 3
            break
    return score

def _looks_like_question(line: str, lines: list[str], index: int) -> bool:
    match = QUESTION_RE.match(line)
    return bool(match and _question_confidence(match.group(2), lines, index) >= 3)

def _extract_options(line: str) -> list[str]:
    if not OPTION_START_RE.match(line):
        return []
    matches = INLINE_OPTION_RE.findall(line)
    return [f"{letter}\uff0e{text.strip()}" for letter, text in matches if text.strip()]

def _split_subsection(line: str) -> tuple[str, str]:
    marker = line.find("\uff08\u672c\u9898")
    if marker < 0:
        return line, ""
    return line[:marker], line[marker:]

def _is_header_line(line: str, metadata: dict[str, Any]) -> bool:
    return line in {
        metadata.get("exam_name"),
        metadata.get("subject_name"),
        metadata.get("meta_text"),
    } or bool(re.fullmatch(r"\d{4}[.\u5e74]\d{1,2}\u6708?", line))

def _score_total_or_default(blocks: list[dict[str, Any]], default: float) -> float:
    scores = [
        block["question"].get("score")
        for block in blocks
        if block.get("type") == "question"
    ]
    numeric = [float(score) for score in scores if score is not None]
    return sum(numeric) if numeric and len(numeric) == len(scores) else default

def save_exam(data: dict[str, Any], path: str | Path) -> Path:
    """\u4fdd\u5b58\u53ef\u7ee7\u7eed\u7f16\u8f91\u7684\u7ed3\u6784\u5316\u8bd5\u9898\u3002"""

    target = Path(path)
    target.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")
    return target


# merged rule group: native DOCX objects


TITLE_AUTHOR_RE = re.compile(
    r"^(.+?)[\t\u3000 ]{2,}([\u3400-\u9fff\u00b7]{2,12})$"
)

def _docx_lines_with_native(
    document: DocumentObject,
) -> tuple[list[str], list[dict[str, Any]]]:
    """\u6309\u539f\u987a\u5e8f\u63d0\u53d6\u6587\u5b57\uff0c\u5e76\u7528\u5360\u4f4d\u7b26\u8bb0\u5f55\u8868\u683c\u548c\u72ec\u7acb\u56fe\u7247\u3002"""

    lines: list[str] = []
    native: list[dict[str, Any]] = []
    paragraph_index = 0
    table_index = 0
    for child in document.element.body.iterchildren():
        if isinstance(child, CT_P):
            paragraph = Paragraph(child, document)
            text = paragraph.text.strip()
            drawings = paragraph._p.xpath(".//w:drawing | .//w:pict")
            if text:
                lines.append(text)
            if drawings:
                if text:
                    native.append(
                        {
                            "kind": "drawing",
                            "source_paragraph_index": paragraph_index,
                            "target_text": text,
                        }
                    )
                else:
                    marker = f"[[NATIVE_DRAWING:{paragraph_index}]]"
                    lines.append(marker)
                    native.append(
                        {
                            "kind": "drawing",
                            "source_paragraph_index": paragraph_index,
                            "marker": marker,
                        }
                    )
            paragraph_index += 1
        elif isinstance(child, CT_Tbl):
            table = Table(child, document)
            marker = f"[[NATIVE_TABLE:{table_index}]]"
            lines.append(marker)
            native.append(
                {
                    "kind": "table",
                    "source_table_index": table_index,
                    "marker": marker,
                    "rows": len(table.rows),
                    "columns": len(table.columns),
                }
            )
            table_index += 1
    return lines, native

def _normalize_composition(result: dict[str, Any]) -> None:
    """\u628a\u5199\u4f5c\u9898\u9644\u5c5e\u6bb5\u843d\u5206\u6210\u6750\u6599\u3001\u5f15\u5bfc\u8bed\u548c\u8981\u6c42\u3002"""

    for block in result.get("blocks", []):
        question = block.get("question")
        if not question:
            continue
        if question.get("number") != 23 and "\u5199\u4f5c" not in str(question.get("stem", "")):
            continue
        paragraphs = [
            "".join(str(segment.get("text", "")) for segment in segments)
            for segments in question.get("embedded_segments", [])
        ]
        if not paragraphs:
            continue
        material: list[str] = []
        prompt: list[str] = []
        requirements: list[str] = []
        for text in paragraphs:
            stripped = text.strip()
            if stripped.startswith(("\u8981\u6c42\uff1a", "\u8981\u6c42:")):
                requirements.append(stripped)
            elif "\u4ee5\u4e0a\u6750\u6599" in stripped and any(
                cue in stripped
                for cue in ("\u5f15\u53d1", "\u8054\u60f3", "\u601d\u8003", "\u542f\u793a", "\u611f\u609f", "\u8bf7\u5199")
            ):
                prompt.append(stripped)
            else:
                material.append(stripped)
        question["composition_material"] = material
        question["composition_prompt"] = prompt
        question["composition_requirements"] = requirements
        question["embedded_segments"] = []

def _normalize_title_authors(result: dict[str, Any]) -> None:
    """\u8bc6\u522b\u6807\u9898\u548c\u4f5c\u8005\u540c\u6bb5\u7684\u6587\u7ae0\u4e0e\u8bd7\u6b4c\u3002"""

    blocks = result.get("blocks", [])
    previous_instruction = ""
    normalized: list[dict[str, Any]] = []
    for block in blocks:
        if block.get("type") == "instruction":
            previous_instruction = str(block.get("text", ""))
            normalized.append(block)
            continue
        if block.get("type") != "material":
            normalized.append(block)
            continue
        paragraphs = list(block.get("paragraphs", []))
        matched = [
            (index, match)
            for index, text in enumerate(paragraphs)
            if (match := TITLE_AUTHOR_RE.match(str(text)))
        ]
        if not matched:
            normalized.append(block)
            continue
        if "\u8bd7" in previous_instruction:
            for position, (start, match) in enumerate(matched):
                end = matched[position + 1][0] if position + 1 < len(matched) else len(paragraphs)
                body = paragraphs[start + 1 : end]
                note = ""
                if body and str(body[-1]).startswith(("\u3010\u6ce8\u3011", "[\u6ce8]", "\uff3b\u6ce8\uff3d")):
                    note = str(body.pop())
                normalized.append(
                    {
                        "type": "poetry",
                        "title": match.group(1).strip(),
                        "author": match.group(2).strip(),
                        "paragraphs": body,
                        "note": note,
                    }
                )
        else:
            first_index, first_match = matched[0]
            prefix = paragraphs[:first_index]
            if prefix:
                normalized.append({"type": "material", "paragraphs": prefix})
            updated = dict(block)
            updated["title"] = first_match.group(1).strip()
            updated["author"] = first_match.group(2).strip()
            updated["paragraphs"] = paragraphs[first_index + 1 :]
            normalized.append(updated)
    result["blocks"] = normalized


# merged rule group: question details


SUBQUESTION_RE = re.compile(r"^[\uff08(](\d+)[\uff09)]\s*(.*)$")

SEGMENTATION_MARKERS_RE = re.compile(r"A.+B.+C")

def _normalize_question_details(result: dict[str, Any]) -> None:
    for block in result.get("blocks", []):
        question = block.get("question")
        if not question:
            continue
        embedded = question.get("embedded_segments", [])
        remaining: list[list[dict[str, str]]] = []
        subquestions = list(question.get("subquestions", []))
        for segments in embedded:
            text = "".join(str(item.get("text", "")) for item in segments).strip()
            sub_match = SUBQUESTION_RE.match(text)
            if sub_match:
                subquestions.append(sub_match.group(2).strip())
            elif (
                question.get("number") == 10
                and SEGMENTATION_MARKERS_RE.search(text)
            ):
                question["segmentation_text"] = text
            else:
                remaining.append(segments)
        question["embedded_segments"] = remaining
        if subquestions:
            question["subquestions"] = subquestions


# merged rule group: headers and poetry


SECTION_RE = re.compile(r"^[\u4e00\u4e8c\u4e09\u56db\u4e94\u516d\u4e03\u516b\u4e5d\u5341]+\u3001")

SUBJECT_RE = re.compile(
    r"^\s*\u8bed\s*\u6587(?:\s*\u8bd5\s*\u9898)?(?:\s+|\u3000+)?(.*)$"
)

DATE_RE = re.compile(r"\b20\d{2}(?:[.\u5e74/-]\d{1,2}(?:\u6708)?)?\b")

NOTE_RE = re.compile(r"^\s*(?:\u3010\s*\u6ce8\s*\u3011|\[\s*\u6ce8\s*\]|\uff3b\s*\u6ce8\s*\uff3d|\u6ce8\s*[\uff1a:])")

AUTHOR_RE = re.compile(
    r"^[\u3400-\u9fff\u00b7]{2,12}(?:\s*(?:\u3010\s*\u6ce8\s*\u3011|\[\s*\u6ce8\s*\]|\uff3b\s*\u6ce8\s*\uff3d))?$"
)

def _normalize_docx_header(source: Path, result: dict[str, Any]) -> None:
    document = Document(source)
    lines = [paragraph.text.strip() for paragraph in document.paragraphs]
    lines = [line for line in lines if line]
    first_section = next(
        (index for index, line in enumerate(lines) if SECTION_RE.match(line)),
        min(len(lines), 16),
    )
    header = lines[:first_section]
    if not header:
        return

    notice_index = next(
        (
            index
            for index, line in enumerate(header)
            if re.fullmatch(r"\s*\u6ce8\u610f\u4e8b\u9879\s*[\uff1a:]?\s*", line)
        ),
        len(header),
    )
    identity_lines = header[:notice_index]
    subject_index: int | None = None
    subject_tail = ""
    for index, line in enumerate(identity_lines):
        match = SUBJECT_RE.match(line)
        if match:
            subject_index = index
            subject_tail = match.group(1).strip()
            break

    exam_name = ""
    if subject_index is not None:
        candidates = identity_lines[:subject_index]
        exam_name = next(
            (
                line
                for line in candidates
                if not _looks_like_meta(line)
            ),
            "",
        )
    if not exam_name:
        exam_name = next(
            (
                line
                for line in identity_lines
                if not SUBJECT_RE.match(line) and not _looks_like_meta(line)
            ),
            str(result.get("metadata", {}).get("exam_name", source.stem)),
        )

    metadata = result.setdefault("metadata", {})
    metadata["exam_name"] = exam_name
    metadata["subject_name"] = "\u8bed\u3000\u6587"
    meta_candidates = [
        line
        for index, line in enumerate(identity_lines)
        if index != subject_index and line != exam_name and _looks_like_meta(line)
    ]
    if subject_tail:
        meta_candidates.insert(0, subject_tail)
    metadata["meta_text"] = next(
        (
            value
            for value in meta_candidates
            if value and not SUBJECT_RE.match(value)
        ),
        "",
    )

    header_set = set(identity_lines)
    cleaned: list[dict[str, Any]] = []
    for block in result.get("blocks", []):
        if block.get("type") == "material" and not cleaned:
            updated = dict(block)
            paragraphs = [
                str(value)
                for value in updated.get("paragraphs", [])
                if str(value).strip() not in header_set
                and not SUBJECT_RE.match(str(value).strip())
            ]
            updated["paragraphs"] = paragraphs
            if paragraphs or updated.get("title") or updated.get("author"):
                cleaned.append(updated)
            continue
        cleaned.append(block)
    result["blocks"] = cleaned

def _looks_like_meta(text: str) -> bool:
    return bool(
        DATE_RE.search(text)
        or "\u6ee1\u5206" in text
        or "\u8003\u8bd5\u65f6\u95f4" in text
        or "\u7b54\u9898" in text
        or "\u8003\u751f" in text
    )

def _normalize_poetry_blocks(result: dict[str, Any]) -> None:
    blocks = result.get("blocks", [])
    normalized: list[dict[str, Any]] = []
    previous_instruction = ""
    for block in blocks:
        if block.get("type") == "instruction":
            previous_instruction = str(block.get("text", ""))
            normalized.append(block)
            continue
        if block.get("type") != "material" or not any(
            cue in previous_instruction for cue in ("\u8bd7", "\u8bcd", "\u66f2")
        ):
            normalized.append(block)
            continue
        paragraphs = [str(value).strip() for value in block.get("paragraphs", [])]
        if len(paragraphs) < 3 or not AUTHOR_RE.fullmatch(paragraphs[1]):
            normalized.append(block)
            continue
        title = paragraphs[0]
        author = NOTE_RE.sub("", paragraphs[1]).strip()
        author = re.sub(
            r"(?:\u3010\s*\u6ce8\s*\u3011|\[\s*\u6ce8\s*\]|\uff3b\s*\u6ce8\s*\uff3d)\s*$",
            "",
            author,
        ).strip()
        body = paragraphs[2:]
        note = ""
        if body and NOTE_RE.match(body[-1]):
            note = body.pop()
        normalized.append(
            {
                "type": "poetry",
                "title": title,
                "author": author,
                "paragraphs": body,
                "note": note,
            }
        )
    result["blocks"] = normalized

def _normalize_notes(result: dict[str, Any]) -> None:
    for block in result.get("blocks", []):
        if block.get("type") not in {"material", "poetry"}:
            continue
        paragraphs = list(block.get("paragraphs", []))
        if paragraphs and NOTE_RE.match(str(paragraphs[-1])):
            block["note"] = str(paragraphs.pop())
            block["paragraphs"] = paragraphs


# merged rule group: source decorations


def _collect_source_decorations(document: Any) -> list[dict[str, Any]]:
    collected: list[dict[str, Any]] = []
    occurrences: dict[str, int] = {}
    for paragraph in document.paragraphs:
        raw_text = paragraph.text
        text = raw_text.strip()
        if not text:
            continue
        leading = len(raw_text) - len(raw_text.lstrip())
        ranges: list[dict[str, Any]] = []
        cursor = 0
        paragraph_style_bold = bool(
            getattr(getattr(paragraph, "style", None), "font", None)
            and getattr(paragraph.style.font, "bold", False)
        )
        for run in paragraph.runs:
            start = cursor - leading
            end = start + len(run.text)
            cursor += len(run.text)
            if end <= 0 or start >= len(text):
                continue
            underline_nodes = run._r.xpath("./w:rPr/w:u")
            emphasis_nodes = run._r.xpath("./w:rPr/w:em")
            underline = [node.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val") or "single" for node in underline_nodes]
            emphasis = [node.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val") or "underDot" for node in emphasis_nodes]
            bold = bool(run.bold) or (run.bold is None and paragraph_style_bold)
            if not underline and not emphasis and not bold:
                continue
            mark: dict[str, Any] = {
                "start": max(0, start),
                "end": min(len(text), end),
            }
            if underline and underline[0] not in {"none", "0", "false"}:
                mark["underline"] = str(underline[0])
            if emphasis and emphasis[0] not in {"none", "0", "false"}:
                mark["emphasis"] = str(emphasis[0])
            if bold:
                mark["bold"] = True
            if len(mark) > 2:
                ranges.append(mark)
        occurrence = occurrences.get(text, 0)
        occurrences[text] = occurrence + 1
        if ranges:
            collected.append(
                {
                    "text": text,
                    "occurrence": occurrence,
                    "ranges": ranges,
                }
            )
    return collected


# merged rule group: formal title refinement


SECTION_RE = re.compile(r"^[\u4e00\u4e8c\u4e09\u56db\u4e94\u516d\u4e03\u516b\u4e5d\u5341]+\u3001")

SUBJECT_RE = re.compile(
    r"^\s*\u8bed\s*\u6587(?:\s*\u8bd5\s*\u9898)?(?:\s+|\u3000+)?(.*)$"
)

CONFIDENTIAL_RE = re.compile(r"(?:\u4fdd\u5bc6|\u7edd\u5bc6|\u542f\u7528\u524d|\u8bd5\u9898\u7c7b\u578b)")

TITLE_CUES = ("\u8003\u8bd5", "\u68c0\u6d4b", "\u6d4b\u8bd5", "\u8bd5\u5377", "\u6a21\u62df")

HEADER_NOTICE_RE = re.compile(r"^\s*(?:\u6ce8\u610f\u4e8b\u9879|\u8003\u751f\u987b\u77e5)\s*[\uff1a:]?\s*$")

HEADER_ITEM_RE = re.compile(r"^\s*\d{1,2}\s*[.\uff0e\u3001]\s*")

HEADER_INFO_CUES = ("\u8003\u8bd5\u65f6\u95f4", "\u8bd5\u5377\u6ee1\u5206", "\u8003\u8bd5\u7528\u65f6", "\u672c\u8bd5\u5377\u5171")

def _refine_exam_name(source: Path, result: dict[str, Any]) -> None:
    lines = [
        paragraph.text.strip()
        for paragraph in Document(source).paragraphs
        if paragraph.text.strip()
    ]
    first_section = next(
        (index for index, line in enumerate(lines) if SECTION_RE.match(line)),
        min(len(lines), 16),
    )
    header = lines[:first_section]
    subject_index = next(
        (index for index, line in enumerate(header) if SUBJECT_RE.match(line)),
        len(header),
    )
    candidates = [
        line
        for line in header[:subject_index]
        if not CONFIDENTIAL_RE.search(line)
        and not re.fullmatch(r"20\d{2}(?:[.\u5e74/-]\d{1,2}(?:\u6708)?)?", line)
    ]
    # Notice headings, numbered notice items, and exam metadata are header
    # content. Excluding them prevents the last notice item from becoming the
    # formal exam title when a document has no separate title line.
    title_candidates = [
        line
        for line in candidates
        if not HEADER_NOTICE_RE.fullmatch(line)
        and not HEADER_ITEM_RE.match(line)
        and not any(cue in line for cue in HEADER_INFO_CUES)
    ]
    formal = next(
        (
            line
            for line in reversed(title_candidates)
            if any(cue in line for cue in TITLE_CUES)
        ),
        title_candidates[0] if title_candidates else candidates[0] if candidates else "",
    )
    if formal:
        result.setdefault("metadata", {})["exam_name"] = formal
    confidentiality = [
        line for line in header[:subject_index] if CONFIDENTIAL_RE.search(line)
    ]
    if confidentiality:
        result.setdefault("metadata", {})["confidentiality_text"] = confidentiality[0]


# merged rule group: adaptive title and score recognition


AUTHOR_RE = re.compile(
    r"^(?:\[[^\]]{1,8}\]|\uff3b[^\uff3d]{1,8}\uff3d|\u3014[^\u3015]{1,8}\u3015|\uff08[^\uff09]{1,8}\uff09)?"
    r"[\u3400-\u9fff\u00b7]{2,14}$"
)

BODY_LEAD_RE = re.compile(
    r"^(?:[\u2460-\u2473\u3251-\u32bf]|\d{1,3}[\uff0e.]|[\u201c\u2018\uff08(\u300a]|[\u3400-\u9fff])"
)

LABEL_RE = re.compile(r"^(?:\u6750\u6599|\u6587\u672c)(?:[\u4e00\u4e8c\u4e09\u56db\u4e94\u516d\u4e03\u516b\u4e5d\u5341]+|\d+)\s*[\uff1a:]")

SCORE_RE = re.compile(r"[\uff08(]\s*(\d+(?:\.\d+)?)\s*\u5206\s*[\uff09)]\s*$")

def _normalize_separate_title_authors(result: dict[str, Any]) -> None:
    """Promote short title and author lines in prose reading materials."""

    previous_instruction = ""
    for block in result.get("blocks", []):
        if block.get("type") == "instruction":
            previous_instruction = str(block.get("text", ""))
            continue
        if block.get("type") != "material":
            continue
        if block.get("title") or block.get("author"):
            continue
        paragraphs = [str(value).strip() for value in block.get("paragraphs", [])]
        if len(paragraphs) < 3:
            continue
        title, author, first_body = paragraphs[:3]
        if any(cue in previous_instruction for cue in ("\u8bd7", "\u8bcd", "\u66f2")):
            continue
        if (
            not title
            or len(title) > 40
            or len(author) > 20
            or LABEL_RE.match(title)
            or not AUTHOR_RE.fullmatch(author)
            or not BODY_LEAD_RE.match(first_body)
        ):
            continue
        block["title"] = title
        block["author"] = author
        block["paragraphs"] = paragraphs[2:]

def _normalize_spaced_scores(result: dict[str, Any]) -> None:
    """Recognize scores such as ``\uff083 \u5206\uff09`` without altering question text."""

    for block in result.get("blocks", []):
        question = block.get("question")
        if not question or question.get("score") is not None:
            continue
        stem = str(question.get("stem", ""))
        match = SCORE_RE.search(stem)
        if not match:
            continue
        question["score"] = float(match.group(1))
        question["stem"] = stem[: match.start()].rstrip()

def _record_segmentation_markers(result: dict[str, Any]) -> None:
    """Record exactly the uppercase markers found in each segmentation passage."""

    for block in result.get("blocks", []):
        question = block.get("question")
        if not question or not question.get("segmentation_text"):
            continue
        text = str(question["segmentation_text"])
        question["segmentation_markers"] = re.findall(r"[A-Z]", text)


# merged rule group: subsections and composition prompts


SUBSECTION_RE = re.compile(
    r"^\s*(?P<name>[\uff08(]\s*[\u4e00\u4e8c\u4e09\u56db\u4e94\u516d\u4e03\u516b\u4e5d\u5341]+\s*[\uff09)]"
    r"\s*[^\uff08(]+?)"
    r"(?P<meta>[\uff08(]\s*\u672c\u9898\u5171.+?[\uff09)])\s*$"
)

AUTHOR_RE = re.compile(
    r"^\s*(?:[\uff08(]\s*[\u3400-\u9fff]{1,4}\s*[\uff09)]\s*)?"
    r"[\u3400-\u9fff\u00b7]{2,14}\s*$"
)

POETRY_TITLE_RE = re.compile(
    r"^\s*(?:"
    r"[\[\uff3b\u3010]\s*[\u7532\u4e59\u4e19\u4e01\u5176\u4e00\u4e8c\u4e09\u56db\u4e94\u516d\u4e03\u516b\u4e5d\u5341]+\s*[\]\uff3d\u3011]"
    r"|\u5176[\u4e00\u4e8c\u4e09\u56db\u4e94\u516d\u4e03\u516b\u4e5d\u5341]+"
    r"|\u7b2c[\u4e00\u4e8c\u4e09\u56db\u4e94\u516d\u4e03\u516b\u4e5d\u5341]+\u9996"
    r").+"
)

NOTE_RE = re.compile(
    r"^\s*(?:\u3010\s*\u6ce8\s*\u3011|\[\s*\u6ce8\s*\]|\uff3b\s*\u6ce8\s*\uff3d|\u6ce8\s*[\uff1a:])"
)

HEADER_DATE_RE = re.compile(r"(?P<date>20\d{2}\s*[./\u5e74-]\s*\d{1,2}(?:\s*\u6708)?)")

SECTION_RE = re.compile(r"^[\u4e00\u4e8c\u4e09\u56db\u4e94\u516d\u4e03\u516b\u4e5d\u5341]+\u3001")

def _normalize_subsections(result: dict[str, Any]) -> None:
    normalized: list[dict[str, Any]] = []
    for block in result.get("blocks", []):
        if block.get("type") == "question":
            question = block.get("question", {})
            kept: list[list[dict[str, Any]]] = []
            lifted: list[dict[str, Any]] = []
            for segments in question.get("embedded_segments", []):
                text = "".join(str(item.get("text", "")) for item in segments).strip()
                match = SUBSECTION_RE.fullmatch(text)
                if match:
                    lifted.append(_subsection(match))
                else:
                    kept.append(segments)
            question["embedded_segments"] = kept
            normalized.append(block)
            normalized.extend(lifted)
            continue

        if block.get("type") != "material":
            normalized.append(block)
            continue

        paragraphs = [str(value).strip() for value in block.get("paragraphs", [])]
        if not any(SUBSECTION_RE.fullmatch(text) for text in paragraphs):
            normalized.append(block)
            continue

        current: list[str] = []
        for text in paragraphs:
            match = SUBSECTION_RE.fullmatch(text)
            if match:
                if current:
                    material = dict(block)
                    material["paragraphs"] = current
                    normalized.append(material)
                    current = []
                normalized.append(_subsection(match))
            else:
                current.append(text)
        if current:
            material = dict(block)
            material["paragraphs"] = current
            normalized.append(material)
    result["blocks"] = normalized

def _subsection(match: re.Match[str]) -> dict[str, str]:
    return {
        "type": "subsection",
        "name": match.group("name").strip(),
        "meta": match.group("meta").strip(),
    }

def _normalize_multi_poetry(result: dict[str, Any]) -> None:
    normalized: list[dict[str, Any]] = []
    previous_instruction = ""
    for block in result.get("blocks", []):
        if block.get("type") == "instruction":
            previous_instruction = str(block.get("text", ""))
            normalized.append(block)
            continue
        if block.get("type") != "material" or not any(
            cue in previous_instruction for cue in ("\u8bd7", "\u8bcd", "\u66f2")
        ):
            normalized.append(block)
            continue

        paragraphs = [str(value).strip() for value in block.get("paragraphs", [])]
        starts = [
            index
            for index in range(len(paragraphs) - 1)
            if POETRY_TITLE_RE.match(paragraphs[index])
            and AUTHOR_RE.fullmatch(paragraphs[index + 1])
        ]
        if not starts:
            normalized.append(block)
            continue

        prefix = paragraphs[: starts[0]]
        if prefix:
            normalized.append({"type": "material", "paragraphs": prefix})
        for position, start in enumerate(starts):
            end = starts[position + 1] if position + 1 < len(starts) else len(paragraphs)
            body = paragraphs[start + 2 : end]
            note = ""
            if body and NOTE_RE.match(body[-1]):
                note = body.pop()
            normalized.append(
                {
                    "type": "poetry",
                    "title": paragraphs[start],
                    "author": paragraphs[start + 1],
                    "paragraphs": body,
                    "note": note,
                }
            )
    result["blocks"] = normalized

def _normalize_composition_prompts(result: dict[str, Any]) -> None:
    cues = (
        "\u4ee5\u4e0a\u6750\u6599",
        "\u8bf7\u4f60\u4ee5",
        "\u8bf7\u4ee5",
        "\u8bf7\u5199\u4e00\u7bc7",
        "\u5199\u4e00\u7bc7",
        "\u53d1\u8868\u6f14\u8bb2",
        "\u611f\u609f\u4e0e\u601d\u8003",
        "\u8054\u60f3\u548c\u601d\u8003",
    )
    for block in result.get("blocks", []):
        question = block.get("question")
        if not question or int(question.get("number", 0)) != 23:
            continue
        material: list[str] = []
        prompt = [str(value) for value in question.get("composition_prompt", [])]
        for text in question.get("composition_material", []):
            value = str(text).strip()
            if any(cue in value for cue in cues):
                prompt.append(value)
            else:
                material.append(value)
        question["composition_material"] = material
        question["composition_prompt"] = prompt

def _normalize_header_metadata(source: Path, result: dict[str, Any]) -> None:
    document = Document(source)
    lines = [paragraph.text.strip() for paragraph in document.paragraphs]
    lines = [line for line in lines if line]
    first_section = next(
        (index for index, line in enumerate(lines) if SECTION_RE.match(line)),
        min(len(lines), 18),
    )
    header = lines[:first_section]
    metadata = result.setdefault("metadata", {})

    for text in header:
        date_match = HEADER_DATE_RE.search(text)
        if not date_match:
            continue
        left = text[: date_match.start()].strip()
        if (
            left
            and left != metadata.get("exam_name")
            and "\u6ee1\u5206" not in left
            and "\u8bed\u6587" not in left
        ):
            metadata["institution_text"] = left
            metadata["exam_date"] = date_match.group("date").strip()
            metadata["meta_text"] = (
                f"{metadata['institution_text']}\t{metadata['exam_date']}"
            )
            break

    exam_info = next(
        (
            text
            for text in header
            if "\u6ee1\u5206" in text
            and any(cue in text for cue in ("\u672c\u5377\u5171", "\u672c\u8bd5\u5377\u5171", "\u8003\u8bd5", "\u7528\u65f6"))
        ),
        "",
    )
    if exam_info:
        metadata["exam_info_text"] = exam_info


# merged rule group: notices and semantic materials


MAJOR_SECTION_RE = re.compile(
    r"^\s*[\u4e00\u4e8c\u4e09\u56db\u4e94\u516d\u4e03\u516b\u4e5d\u5341]+\s*[\u3001\uff0c,]\s*(?P<body>.+?)\s*$"
)

NOTICE_TITLE_RE = re.compile(r"^\s*(\u6ce8\u610f\u4e8b\u9879|\u8003\u751f\u987b\u77e5)\s*[\uff1a:]?\s*$")

NOTICE_ITEM_RE = re.compile(r"^\s*\d{1,2}\s*[\uff0e.\u3001]\s*(?P<body>.+?)\s*$")

LABEL_RE = re.compile(r"^\s*(?:\u6750\u6599|\u6587\u672c)(?:[\u4e00\u4e8c\u4e09\u56db\u4e94\u516d\u4e03\u516b\u4e5d\u5341]+|\d+)\s*[\uff1a:]")

NOTE_RE = re.compile(
    r"^\s*(?:\u3010\s*\u6ce8\s*\u3011|\[\s*\u6ce8\s*\]|\uff3b\s*\u6ce8\s*\uff3d|\u6ce8\s*[\uff1a:])"
)

SOURCE_RE = re.compile(
    r"^\s*[\uff08(]\s*(?:"
    r"(?:\u6458\u81ea|\u6458\u7f16\u81ea|\u6458\u9009\u81ea|\u9009\u81ea|\u8282\u9009\u81ea|\u6539\u7f16\u81ea|\u636e).+|"
    r"\u672c\u6587.+|\u6709\u5220\u6539|\u6709\u6539\u52a8"
    r")[\uff09)]\s*$"
)

AUTHOR_RE = re.compile(
    r"^\s*(?:(?:\u3010[^\u3011]{1,8}\u3011|\[[^\]]{1,8}\]|\uff3b[^\uff3d]{1,8}\uff3d|"
    r"\uff08[^\uff09]{1,8}\uff09|\([^)]{1,8}\))\s*)?"
    r"[\u3400-\u9fff\u00b7\u30fb\u3001\uff0c,\s]{2,30}\s*$"
)

BODY_LEAD_RE = re.compile(r"^(?:[\u2460-\u2473]|\d|[\u201c\u2018\uff08(\u300a]|\u3400-\u9fff)")

SECTION_CUES = ("\u9605\u8bfb", "\u8bed\u8a00\u6587\u5b57\u8fd0\u7528", "\u5199\u4f5c")

POETRY_CUES = ("\u8bd7", "\u8bcd", "\u66f2")

def _normalize_header_notices(source: Path, result: dict[str, Any]) -> None:
    """Treat ``\u8003\u751f\u987b\u77e5`` and ``\u6ce8\u610f\u4e8b\u9879`` as equivalent optional notice headings."""

    document = Document(source)
    lines = [paragraph.text.strip() for paragraph in document.paragraphs]
    lines = [line for line in lines if line]
    first_section = next(
        (
            index
            for index, line in enumerate(lines)
            if _is_major_section(line)
        ),
        min(len(lines), 18),
    )
    header = lines[:first_section]
    notice_index = next(
        (
            index
            for index, line in enumerate(header)
            if NOTICE_TITLE_RE.fullmatch(line)
        ),
        None,
    )
    if notice_index is None:
        return

    match = NOTICE_TITLE_RE.fullmatch(header[notice_index])
    assert match is not None
    raw_notice_lines = header[notice_index + 1 :]
    notices: list[str] = []
    for line in raw_notice_lines:
        item = NOTICE_ITEM_RE.fullmatch(line)
        notices.append(item.group("body").strip() if item else line.strip())

    metadata = result.setdefault("metadata", {})
    metadata["notice_title"] = f"{match.group(1)}\uff1a"
    metadata["notices"] = notices
    metadata["meta_text"] = _header_meta_before_notice(
        header[:notice_index],
        metadata,
    )
    exam_info = str(metadata.get("exam_info_text", "")).strip()
    if exam_info and exam_info in raw_notice_lines:
        metadata.pop("exam_info_text", None)

def _header_meta_before_notice(
    lines: list[str],
    metadata: dict[str, Any],
) -> str:
    excluded = {
        str(metadata.get("exam_name", "")).strip(),
        str(metadata.get("subject_name", "")).strip(),
        str(metadata.get("confidentiality_text", "")).strip(),
    }
    candidates = [
        line
        for line in lines
        if line not in excluded
        and ("\u6ee1\u5206" in line or "\u8003\u8bd5\u65f6\u95f4" in line or "\u8003\u8bd5\u7528\u65f6" in line)
    ]
    if candidates:
        return candidates[0]
    existing = str(metadata.get("meta_text", ""))
    if existing.rstrip("\uff1a:") in {"\u8003\u751f\u987b\u77e5", "\u6ce8\u610f\u4e8b\u9879"}:
        return ""
    return existing

def _lift_embedded_sections(result: dict[str, Any]) -> None:
    """Move major section headings out of question or material payloads."""

    normalized: list[dict[str, Any]] = []
    for block in result.get("blocks", []):
        lifted: list[dict[str, str]] = []
        question = block.get("question")
        if question:
            kept: list[list[dict[str, Any]]] = []
            for segments in question.get("embedded_segments", []):
                text = "".join(
                    str(segment.get("text", "")) for segment in segments
                ).strip()
                if _is_major_section(text):
                    lifted.append({"type": "section_title", "text": text})
                else:
                    kept.append(segments)
            question["embedded_segments"] = kept
            normalized.append(block)
            normalized.extend(lifted)
            continue

        if block.get("type") == "material":
            paragraphs = [str(value).strip() for value in block.get("paragraphs", [])]
            if any(_is_major_section(text) for text in paragraphs):
                current: list[str] = []
                for text in paragraphs:
                    if _is_major_section(text):
                        if current:
                            item = dict(block)
                            item["paragraphs"] = current
                            normalized.append(item)
                            current = []
                        normalized.append({"type": "section_title", "text": text})
                    else:
                        current.append(text)
                if current:
                    item = dict(block)
                    item["paragraphs"] = current
                    normalized.append(item)
                continue
        normalized.append(block)
    result["blocks"] = normalized

def _is_major_section(text: str) -> bool:
    match = MAJOR_SECTION_RE.fullmatch(text)
    return bool(match and any(cue in match.group("body") for cue in SECTION_CUES))

def _normalize_single_poetry(result: dict[str, Any]) -> None:
    """Promote one-poem material blocks with dynasty-prefixed authors."""

    normalized: list[dict[str, Any]] = []
    previous_instruction = ""
    for block in result.get("blocks", []):
        if block.get("type") == "instruction":
            previous_instruction = str(block.get("text", ""))
            normalized.append(block)
            continue
        if block.get("type") != "material" or not any(
            cue in previous_instruction for cue in POETRY_CUES
        ):
            normalized.append(block)
            continue
        paragraphs = [str(value).strip() for value in block.get("paragraphs", [])]
        if len(paragraphs) < 3 or not _is_author_line(paragraphs[1]):
            normalized.append(block)
            continue
        body = paragraphs[2:]
        note = str(block.get("note", ""))
        if not note and body and NOTE_RE.match(body[-1]):
            note = body.pop()
        normalized.append(
            {
                "type": "poetry",
                "title": paragraphs[0],
                "author": paragraphs[1],
                "paragraphs": body,
                "note": note,
            }
        )
    result["blocks"] = normalized

def _normalize_prose_title_authors(result: dict[str, Any]) -> None:
    """Promote prose title and author lines, including multiple authors."""

    previous_instruction = ""
    for block in result.get("blocks", []):
        if block.get("type") == "instruction":
            previous_instruction = str(block.get("text", ""))
            continue
        if block.get("type") != "material":
            continue
        if block.get("title") or block.get("author"):
            _promote_trailing_source(block)
            continue
        if any(cue in previous_instruction for cue in POETRY_CUES):
            continue
        paragraphs = [str(value).strip() for value in block.get("paragraphs", [])]
        if len(paragraphs) < 3:
            continue
        title, author, first_body = paragraphs[:3]
        if (
            not title
            or len(title) > 40
            or LABEL_RE.match(title)
            or not _is_author_line(author)
            or not BODY_LEAD_RE.match(first_body)
        ):
            continue
        block["title"] = title
        block["author"] = author
        block["paragraphs"] = paragraphs[2:]
        _promote_trailing_source(block)

def _is_author_line(text: str) -> bool:
    value = text.strip()
    if not AUTHOR_RE.fullmatch(value):
        return False
    han_count = len(re.findall(r"[\u3400-\u9fff]", value))
    return 2 <= han_count <= 20 and not any(
        mark in value for mark in ("\u3002", "\uff01", "\uff1f", "\uff1b", "\uff1a", "\u201c", "\u201d")
    )

def _promote_trailing_source(block: dict[str, Any]) -> None:
    paragraphs = [str(value).strip() for value in block.get("paragraphs", [])]
    if not block.get("source") and paragraphs and SOURCE_RE.fullmatch(paragraphs[-1]):
        block["source"] = paragraphs.pop()
        block["paragraphs"] = paragraphs

def _annotate_material_roles(
    source: Path | None,
    result: dict[str, Any],
) -> None:
    """Record label, source and centered subheading roles for later rendering."""

    centered_bold: set[str] = set()
    if source is not None:
        for paragraph in Document(source).paragraphs:
            text = paragraph.text.strip()
            visible_runs = [run for run in paragraph.runs if run.text.strip()]
            if (
                text
                and len(text) <= 60
                and paragraph.alignment == WD_ALIGN_PARAGRAPH.CENTER
                and any(run.bold for run in visible_runs)
            ):
                centered_bold.add(text)

    for block in result.get("blocks", []):
        if block.get("type") != "material":
            continue
        paragraphs = [str(value).strip() for value in block.get("paragraphs", [])]
        roles: list[str] = []
        formats = list(block.get("paragraph_formats", []))
        offset = int(bool(block.get("title"))) + int(bool(block.get("author")))
        for index, text in enumerate(paragraphs):
            role = "body"
            spec: dict[str, Any] | None = None
            if LABEL_RE.match(text):
                role = "label"
                spec = _format_spec("\u9ed1\u4f53", "\u5de6\u5bf9\u9f50")
            elif SOURCE_RE.fullmatch(text):
                role = "source"
                spec = _format_spec("\u4eff\u5b8b", "\u53f3\u5bf9\u9f50")
            elif text in centered_bold:
                role = "subheading"
                spec = _format_spec("\u9ed1\u4f53", "\u5c45\u4e2d")
            roles.append(role)
            if spec is not None:
                spec["target_index"] = offset + index
                spec["semantic_role"] = role
                formats.append(spec)
        block["paragraph_roles"] = roles
        if formats:
            block["paragraph_formats"] = formats

def _format_spec(font: str, alignment: str) -> dict[str, Any]:
    return {
        "font": font,
        "size_pt": 10.5,
        "bold": False,
        "alignment": alignment,
        "left_indent_chars": 0,
        "special_indent": "\u65e0",
        "special_indent_chars": 0,
        "line_spacing": 1.25,
        "space_before_pt": 0,
        "space_after_pt": 0,
    }


# merged rule group: answers and title-author enhancement


AUTHOR_RE = re.compile(
    r"^\s*(?:(?:\u3010[^\u3011]{1,8}\u3011|\[[^\]]{1,8}\]|"
    r"\uff08[^\uff09]{1,8}\uff09|\([^)]{1,8}\))\s*)?"
    r"[\u3400-\u9fff\u00b7\u3001\uff0c,\s]{2,30}"
    r"(?:\s*(?:\u3010\s*\u6ce8\s*\u3011|\[\s*\u6ce8\s*\]|\uff3b\s*\u6ce8\s*\uff3d))?\s*$"
)

SOURCE_RE = re.compile(
    r"^\s*[\uff08(]\s*(?:\u6458\u81ea|\u6458\u7f16\u81ea|\u6458\u9009\u81ea|\u9009\u81ea|\u8282\u9009\u81ea|"
    r"\u6539\u7f16\u81ea|\u636e.+|\u672c\u6587.+|\u6709\u5220\u6539|\u6709\u6539\u52a8).+[\uff09)]\s*$"
)

def _enhance_material_titles_authors(
    source: Path,
    result: dict[str, Any],
) -> None:
    document = Document(source)
    centered: set[str] = set()
    centered_bold: set[str] = set()
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text or len(text) > 60:
            continue
        if paragraph.alignment == WD_ALIGN_PARAGRAPH.CENTER:
            centered.add(text)
            if any(run.bold for run in paragraph.runs if run.text.strip()):
                centered_bold.add(text)

    for block in result.get("blocks", []):
        if block.get("type") != "material":
            continue
        paragraphs = [str(value).strip() for value in block.get("paragraphs", [])]
        if not block.get("title") and paragraphs:
            first = paragraphs[0]
            if first in centered_bold and not _is_author(first):
                block["title"] = first
                paragraphs.pop(0)
                block["paragraphs"] = paragraphs
        roles = [str(value) for value in block.get("paragraph_roles", [])]
        if len(roles) != len(paragraphs):
            roles = ["body"] * len(paragraphs)
        formats = list(block.get("paragraph_formats", []))
        for index, text in enumerate(paragraphs):
            if SOURCE_RE.fullmatch(text):
                roles[index] = "source"
                formats.append(_format(index, "\u4eff\u5b8b", "\u53f3\u5bf9\u9f50", "source"))
                continue
            if text not in centered:
                continue
            previous_centered = index > 0 and paragraphs[index - 1] in centered
            next_is_author = (
                index + 1 < len(paragraphs)
                and paragraphs[index + 1] in centered
                and _is_author(paragraphs[index + 1])
            )
            if _is_author(text) and previous_centered:
                roles[index] = "author"
                formats.append(_format(index, "\u4eff\u5b8b", "\u5c45\u4e2d", "author"))
            elif text in centered_bold or next_is_author:
                roles[index] = "subheading"
                formats.append(_format(index, "\u9ed1\u4f53", "\u5c45\u4e2d", "subheading"))
        block["paragraph_roles"] = roles
        if formats:
            block["paragraph_formats"] = _deduplicate_formats(formats)

def _trim_answer_from_exam(result: dict[str, Any]) -> None:
    for block in result.get("blocks", []):
        question = block.get("question")
        if not question or int(question.get("number", 0)) != 23:
            continue
        for key in (
            "composition_material",
            "composition_prompt",
            "composition_requirements",
        ):
            kept: list[str] = []
            for value in question.get(key, []):
                text = str(value).strip()
                if ANSWER_MARKER_RE.fullmatch(text):
                    break
                kept.append(str(value))
            question[key] = kept
        question["composition_prompt"] = [
            text
            for text in question.get("composition_prompt", [])
            if not ANSWER_MARKER_RE.fullmatch(str(text).strip())
        ]
        question["composition_requirements"] = [
            text
            for text in question.get("composition_requirements", [])
            if not ANSWER_MARKER_RE.fullmatch(str(text).strip())
        ]

def _is_author(text: str) -> bool:
    value = text.strip()
    if not AUTHOR_RE.fullmatch(value):
        return False
    han = len(re.findall(r"[\u3400-\u9fff]", value))
    return 2 <= han <= 20 and not any(
        mark in value for mark in ("\u3002", "\uff01", "\uff1f", "\uff1b", "\uff1a", "\u201c", "\u201d")
    )

def _format(
    index: int,
    font: str,
    alignment: str,
    role: str,
) -> dict[str, Any]:
    return {
        "target_index": index,
        "semantic_role": role,
        "font": font,
        "size_pt": 10.5,
        "bold": False,
        "alignment": alignment,
        "left_indent_chars": 0,
        "special_indent": "\u65e0",
        "special_indent_chars": 0,
        "line_spacing": 1.25,
        "space_before_pt": 0,
        "space_after_pt": 0,
    }

def _deduplicate_formats(
    formats: list[dict[str, Any]],
) -> list[dict[str, Any]]:
    result: list[dict[str, Any]] = []
    seen: set[tuple[int, str]] = set()
    for spec in reversed(formats):
        key = (
            int(spec.get("target_index", -1)),
            str(spec.get("semantic_role", "")),
        )
        if key in seen:
            continue
        seen.add(key)
        result.append(spec)
    result.reverse()
    return result


# merged rule group: publication-note semantics


PUBLICATION_NOTE_RE = re.compile(
    r"^\s*[\uff08(]"
    r"(?=[^\uff09)]*(?:\u8bd1|\u53d1\u8868\u4e8e|\u520a\u4e8e|\u8f7d\u4e8e|\u539f\u8f7d|\u521d\u520a|\u9996\u520a|\u6709\u5220\u6539|\u6709\u6539\u52a8))"
    r"(?!\s*(?:\u6458\u81ea|\u6458\u7f16\u81ea|\u6458\u9009\u81ea|\u9009\u81ea|\u8282\u9009\u81ea|\u6539\u7f16\u81ea|\u636e|\u6765\u6e90))"
    r"[^\uff09)]+[\uff09)]\s*$"
)

def mark_publication_notes(result: dict[str, Any]) -> None:
    """Mark publication or translation notes without changing source lines."""

    for block in result.get("blocks", []):
        if block.get("type") != "material":
            continue
        paragraphs = [str(value).strip() for value in block.get("paragraphs", [])]
        roles = [str(value) for value in block.get("paragraph_roles", [])]
        if len(roles) != len(paragraphs):
            roles = ["body"] * len(paragraphs)
        formats = list(block.get("paragraph_formats", []))
        changed = False
        for index, text in enumerate(paragraphs):
            if roles[index] == "source" or not PUBLICATION_NOTE_RE.fullmatch(text):
                continue
            roles[index] = "publication_note"
            formats = [
                spec
                for spec in formats
                if int(spec.get("target_index", -1)) != index
            ]
            formats.append(_publication_note_format(index))
            changed = True
        if changed:
            block["paragraph_roles"] = roles
            block["paragraph_formats"] = formats

def _publication_note_format(index: int) -> dict[str, Any]:
    return {
        "target_index": index,
        "semantic_role": "publication_note",
        "font": "\u5b8b\u4f53",
        "size_pt": 10.5,
        "bold": False,
        "alignment": "\u53f3\u5bf9\u9f50",
        "left_indent_chars": 0,
        "special_indent": "\u65e0",
        "special_indent_chars": 0,
        "line_spacing": 1.25,
        "space_before_pt": 0,
        "space_after_pt": 0,
    }


# merged rule group: automatic numbering


@dataclass
class NumberingLevel:
    num_fmt: str
    level_text: str
    start: int

def materialize_automatic_numbering(document: Any) -> int:
    """Prefix paragraph text with the list label Word normally paints on screen."""

    definitions = _numbering_definitions(document)
    counters: dict[tuple[int, int], int] = {}
    inserted = 0
    for paragraph in document.paragraphs:
        key = _paragraph_numbering(paragraph)
        if key is None or key not in definitions:
            continue
        num_id, level = key
        definition = definitions[key]
        counter_key = (num_id, level)
        value = counters.get(counter_key, definition.start - 1) + 1
        counters[counter_key] = value
        prefix = _format_prefix(definition, value)
        if not prefix or _already_has_prefix(paragraph.text, prefix):
            continue
        _prepend_text(paragraph, prefix)
        inserted += 1
    return inserted

def _numbering_definitions(document: Any) -> dict[tuple[int, int], NumberingLevel]:
    try:
        root = document.part.numbering_part.element
    except (AttributeError, KeyError, NotImplementedError):
        # Some valid DOCX files have no numbering relationship at all.  They
        # still contain ordinary paragraphs and should continue through the
        # normal parser without requiring a synthetic numbering part.
        return {}
    abstract: dict[int, dict[int, NumberingLevel]] = {}
    for item in root.findall(qn("w:abstractNum")):
        abstract_id = int(item.get(qn("w:abstractNumId")))
        levels: dict[int, NumberingLevel] = {}
        for level in item.findall(qn("w:lvl")):
            level_id = int(level.get(qn("w:ilvl"), "0"))
            fmt = level.find(qn("w:numFmt"))
            text = level.find(qn("w:lvlText"))
            start = level.find(qn("w:start"))
            levels[level_id] = NumberingLevel(
                str(fmt.get(qn("w:val"), "decimal")) if fmt is not None else "decimal",
                str(text.get(qn("w:val"), "%1.")) if text is not None else "%1.",
                int(start.get(qn("w:val"), "1")) if start is not None else 1,
            )
        abstract[abstract_id] = levels
    result: dict[tuple[int, int], NumberingLevel] = {}
    for item in root.findall(qn("w:num")):
        num_id = int(item.get(qn("w:numId")))
        abstract_ref = item.find(qn("w:abstractNumId"))
        if abstract_ref is None:
            continue
        abstract_id = int(abstract_ref.get(qn("w:val")))
        levels = dict(abstract.get(abstract_id, {}))
        for override in item.findall(qn("w:lvlOverride")):
            level_id = int(override.get(qn("w:ilvl"), "0"))
            current = levels.get(level_id, NumberingLevel("decimal", "%1.", 1))
            start_override = override.find(qn("w:startOverride"))
            if start_override is not None:
                current = NumberingLevel(
                    current.num_fmt,
                    current.level_text,
                    int(start_override.get(qn("w:val"), "1")),
                )
            levels[level_id] = current
        for level_id, definition in levels.items():
            result[(num_id, level_id)] = definition
    return result

def _paragraph_numbering(paragraph: Any) -> tuple[int, int] | None:
    num_pr = None
    if paragraph._p.pPr is not None:
        num_pr = paragraph._p.pPr.numPr
    if num_pr is None and paragraph.style is not None:
        style_ppr = paragraph.style.element.pPr
        if style_ppr is not None:
            num_pr = style_ppr.numPr
    if num_pr is None or num_pr.numId is None:
        return None
    num_id = int(num_pr.numId.val)
    level = int(num_pr.ilvl.val) if num_pr.ilvl is not None else 0
    return num_id, level

def _format_prefix(definition: NumberingLevel, value: int) -> str:
    if definition.num_fmt == "upperLetter":
        label = _letters(value)
    elif definition.num_fmt == "lowerLetter":
        label = _letters(value).lower()
    elif definition.num_fmt == "decimal":
        label = str(value)
    else:
        return ""
    prefix = definition.level_text.replace("%1", label)
    if re.fullmatch(r"[A-Za-z0-9]+\.", prefix):
        prefix = prefix[:-1] + "\uff0e"
    elif re.fullmatch(r"\([A-Za-z0-9]+\)", prefix):
        prefix = "\uff08" + prefix[1:-1] + "\uff09"
    return prefix

def _letters(value: int) -> str:
    result = ""
    current = max(1, int(value))
    while current:
        current, remainder = divmod(current - 1, 26)
        result = chr(ord("A") + remainder) + result
    return result

def _already_has_prefix(text: str, prefix: str) -> bool:
    stripped = str(text).lstrip()
    plain = prefix.replace("\uff0e", ".").replace("\uff08", "(").replace("\uff09", ")")
    return stripped.startswith(prefix) or stripped.startswith(plain)

def _prepend_text(paragraph: Any, prefix: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = prefix + paragraph.runs[0].text
        return
    paragraph.add_run(prefix)


# merged from current_importer.py


_AUTHOR_RE = re.compile(r"^[\u3400-\u9fff\u00b7\u3001\uff0c,\s]{2,30}$")

_TITLE_AUTHOR_RE = re.compile(
    r"^(?P<title>\u300a[^\u300b]+\u300b|[^|\uff5c\t]{2,40})"
    r"(?:\s{2,}|\u3000{1,}|[|\uff5c\t]+)"
    r"(?P<author>[\u3400-\u9fff\u00b7\u3001\uff0c,\s]{2,30})$"
)

_SOURCE_RE = re.compile(
    r"^\s*[\uff08(].*(?:\u6458\u81ea|\u6458\u7f16\u81ea|\u9009\u81ea|\u8282\u9009\u81ea|\u8bd1|\u53d1\u8868\u4e8e|\u520a\u4e8e|\u8f7d\u4e8e|\u6709\u5220\u6539|\u6709\u6539\u52a8).*[\uff09)]\s*$"
)

def _normalize_same_line_title_authors(
    result: dict[str, Any],
    diagnostics: list[dict[str, Any]],
) -> None:
    """Split centered-style ``title  author`` lines when evidence is strong."""

    for index, block in enumerate(result.get("blocks", [])):
        if not isinstance(block, dict) or block.get("type") not in {"material", "poetry"}:
            continue
        if block.get("title") or block.get("author"):
            continue
        paragraphs = [str(value).strip() for value in block.get("paragraphs", [])]
        if not paragraphs:
            continue
        match = _TITLE_AUTHOR_RE.fullmatch(paragraphs[0])
        if not match or not _is_author(match.group("author")):
            continue
        title = match.group("title").strip()
        author = match.group("author").strip()
        if _looks_like_label(title) or _looks_like_body(title):
            continue
        block["title"] = title
        block["author"] = author
        block["paragraphs"] = paragraphs[1:]
        roles = list(block.get("paragraph_roles", []))
        if len(roles) != len(paragraphs):
            roles = ["body"] * len(paragraphs)
        block["paragraph_roles"] = roles[1:]
        formats = list(block.get("paragraph_formats", []))
        formats.append(_role_format("author", 0, "\u4eff\u5b8b", "\u5c45\u4e2d"))
        block["paragraph_formats"] = formats
        diagnostics.append(
            {
                "code": "same-line-title-author",
                "block": block.get("id", f"block-{index + 1}"),
                "message": "\u5df2\u6839\u636e\u5206\u9694\u7b26\u548c\u4f5c\u8005\u6587\u5b57\u7279\u5f81\u62c6\u5206\u540c\u4e00\u884c\u6807\u9898\u4e0e\u4f5c\u8005",
            }
        )

def _normalize_publication_notes(result: dict[str, Any]) -> None:
    """Ensure article-end parenthetical notes use the source-note role."""

    for block in result.get("blocks", []):
        if not isinstance(block, dict) or block.get("type") not in {"material", "poetry"}:
            continue
        paragraphs = [str(value).strip() for value in block.get("paragraphs", [])]
        roles = list(block.get("paragraph_roles", []))
        if len(roles) != len(paragraphs):
            roles = ["body"] * len(paragraphs)
        formats = list(block.get("paragraph_formats", []))
        for index, text in enumerate(paragraphs):
            if _SOURCE_RE.fullmatch(text):
                roles[index] = "source"
                formats = [
                    item
                    for item in formats
                    if int(item.get("target_index", -1)) != index
                ]
                formats.append(_role_format("source", index, "\u4eff\u5b8b", "\u53f3\u5bf9\u9f50"))
        block["paragraph_roles"] = roles
        if formats:
            block["paragraph_formats"] = formats

def _role_format(
    role: str,
    index: int,
    font: str,
    alignment: str,
) -> dict[str, Any]:
    return {
        "target_index": index,
        "semantic_role": role,
        "font": font,
        "size_pt": 10.5,
        "bold": False,
        "alignment": alignment,
        "left_indent_chars": 0,
        "special_indent": "\u65e0",
        "special_indent_chars": 0,
        "line_spacing": 1.25,
        "space_before_pt": 0,
        "space_after_pt": 0,
    }

def _is_author(value: str) -> bool:
    text = value.strip()
    if not _AUTHOR_RE.fullmatch(text):
        return False
    return 2 <= len(re.findall(r"[\u3400-\u9fff]", text)) <= 20

def _looks_like_label(value: str) -> bool:
    return bool(re.match(r"^(?:\u6750\u6599|\u6587\u672c|\u6587\u6bb5|\u6ce8\u91ca|\u8bf4\u660e)[\u4e00\u4e8c\u4e09\u56db\u4e94\u516d\u4e03\u516b\u4e5d\u53410-9\uff1a:]", value))

def _looks_like_body(value: str) -> bool:
    return len(value) > 24 or value.endswith(("\u3002", "\uff01", "\uff1f", "\uff1b"))


def _parse_docx_payload(source: Path, original: Path) -> dict[str, Any]:
    """Extract a DOCX in document order while retaining native objects."""

    document = Document(source)
    lines, native_objects = _docx_lines_with_native(document)
    result = parse_plain_lines(lines, original.stem)
    metadata = result.setdefault("metadata", {})
    metadata["source_docx_path"] = str(original.resolve())
    metadata["native_objects"] = native_objects
    return result


def _run_docx_normalizers(source: Path, result: dict[str, Any]) -> dict[str, Any]:
    """Apply the verified normalization passes in their historical order."""

    _normalize_composition(result)
    _normalize_title_authors(result)
    _normalize_question_details(result)
    _normalize_docx_header(source, result)
    _normalize_poetry_blocks(result)
    _normalize_notes(result)
    result.setdefault("metadata", {})["source_decorations"] = (
        _collect_source_decorations(Document(source))
    )
    _refine_exam_name(source, result)
    _normalize_separate_title_authors(result)
    _normalize_spaced_scores(result)
    _record_segmentation_markers(result)
    _normalize_subsections(result)
    _normalize_multi_poetry(result)
    _normalize_composition_prompts(result)
    _normalize_header_metadata(source, result)
    _normalize_header_notices(source, result)
    _lift_embedded_sections(result)
    _normalize_single_poetry(result)
    _normalize_prose_title_authors(result)
    _annotate_material_roles(source, result)
    return result


def import_exam(path: str | Path) -> dict[str, Any]:
    """Import one exam or answer document through the canonical pipeline.

    The flexible layer keeps text and JSON compatibility for saved drafts. The
    desktop facade accepts only DOCX or DOC and performs legacy DOC conversion.
    PDF input is rejected consistently because PDF export/import was removed
    from the current product scope.
    """

    source = Path(path)
    suffix = source.suffix.lower()
    if suffix == ".json":
        result = json.loads(source.read_text(encoding="utf-8"))
        mark_publication_notes(result)
        return result
    if suffix in {".txt", ".md", ".markdown"}:
        text = source.read_text(encoding="utf-8-sig")
        result = parse_plain_lines(
            [line.strip() for line in text.splitlines() if line.strip()],
            source.stem,
        )
        _normalize_composition(result)
        _normalize_question_details(result)
        _normalize_subsections(result)
        _normalize_composition_prompts(result)
        mark_publication_notes(result)
        return result
    if suffix == ".pdf":
        raise ValueError("\u5f53\u524d\u7248\u672c\u4e0d\u518d\u63d0\u4f9b PDF \u5bfc\u5165\uff0c\u8bf7\u4f7f\u7528 DOCX \u6216 DOC \u6587\u4ef6\u3002")
    if suffix != ".docx":
        raise ValueError(
            "\u652f\u6301\u7684\u8bd5\u9898\u683c\u5f0f\u4e3a DOCX\u3001DOC\u3001JSON\u3001TXT \u548c Markdown\u3002"
        )
    if is_standalone_answer_docx(source):
        return standalone_answer_model(parse_answer_docx(source))

    document = Document(source)
    inserted = materialize_automatic_numbering(document)
    if inserted:
        with tempfile.TemporaryDirectory(prefix="exam-numbering-") as folder:
            prepared = Path(folder) / source.name
            document.save(prepared)
            result = _parse_docx_payload(prepared, source)
            result = _run_docx_normalizers(source, result)
    else:
        result = _parse_docx_payload(source, source)
        result = _run_docx_normalizers(source, result)

    _enhance_material_titles_authors(source, result)
    answer_start = find_answer_start(source)
    if answer_start is not None:
        _trim_answer_from_exam(result)
        answer = parse_answer_docx(
            source,
            start_paragraph=answer_start,
            fallback_title=str(result.get("metadata", {}).get("exam_name", "")),
        )
        attach_answer_blocks(result, answer)
    mark_publication_notes(result)
    ensure_block_ids(result)
    result.setdefault("metadata", {})["source_format"] = "docx"
    return result


__all__ = [
    "import_exam",
    "parse_plain_lines",
    "save_exam",
    "materialize_automatic_numbering",
    "mark_publication_notes",
]
