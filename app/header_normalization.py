"""Idempotent header cleanup and centered exam-information formatting."""

from __future__ import annotations

from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt

from .publication_note_formatting import apply_publication_note_formatting


def apply_header_normalization(
    docx_path: str | Path,
    raw_exam: dict[str, Any],
) -> None:
    """Apply v10, remove duplicate regional metadata, and center exam info."""

    target = Path(docx_path)
    apply_publication_note_formatting(target, raw_exam)
    normalize_header(target, raw_exam)


def normalize_header(
    docx_path: str | Path,
    raw_exam: dict[str, Any],
) -> None:
    """Normalize the rendered header without re-running earlier formatters."""

    target = Path(docx_path)
    document = Document(target)
    metadata = raw_exam.get("metadata", {})
    _deduplicate_institution_date(document, metadata)
    _deduplicate_and_center_exam_info(document, metadata)
    document.save(target)


def _deduplicate_institution_date(
    document: Any,
    metadata: dict[str, Any],
) -> None:
    institution = str(metadata.get("institution_text", "")).strip()
    exam_date = str(metadata.get("exam_date", "")).strip()
    if not institution or not exam_date:
        return
    matches = [
        paragraph
        for paragraph in document.paragraphs[:24]
        if institution in paragraph.text and exam_date in paragraph.text
    ]
    for paragraph in matches[1:]:
        _remove_paragraph(paragraph)


def _deduplicate_and_center_exam_info(
    document: Any,
    metadata: dict[str, Any],
) -> None:
    preferred = str(metadata.get("exam_info_text", "")).strip()
    if not preferred:
        candidate = str(metadata.get("meta_text", "")).strip()
        if _is_exam_info(candidate):
            preferred = candidate

    matches = [
        paragraph
        for paragraph in document.paragraphs[:24]
        if _is_exam_info(paragraph.text)
        and (not preferred or _compact(paragraph.text) == _compact(preferred))
    ]
    if not matches and preferred:
        matches = [
            paragraph
            for paragraph in document.paragraphs[:24]
            if _compact(paragraph.text) == _compact(preferred)
        ]
    for paragraph in matches[1:]:
        _remove_paragraph(paragraph)
    if not matches:
        return
    paragraph = matches[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.left_indent = Pt(0)
    paragraph.paragraph_format.right_indent = Pt(0)
    paragraph.paragraph_format.first_line_indent = Pt(0)
    for run in paragraph.runs:
        run.font.name = "SimSun"
        run.font.size = Pt(10.5)
        run.bold = False
        fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
        for key in ("ascii", "hAnsi", "eastAsia", "cs"):
            fonts.set(qn(f"w:{key}"), "SimSun")


def _is_exam_info(text: str) -> bool:
    value = _compact(text)
    if not value or "本题共" in value:
        return False
    has_scope = any(cue in value for cue in ("本试卷", "本卷", "全卷", "全试卷"))
    has_count = "题" in value
    has_score = "满分" in value or ("分" in value and "分钟" not in value)
    has_time = any(cue in value for cue in ("考试用时", "考试时间", "用时", "分钟"))
    return has_scope and has_count and has_score and has_time


def _compact(text: str) -> str:
    return "".join(str(text).split())


def _remove_paragraph(paragraph: Any) -> None:
    parent = paragraph._element.getparent()
    if parent is not None:
        parent.remove(paragraph._element)


__all__ = [
    "apply_header_normalization",
    "normalize_header",
]
