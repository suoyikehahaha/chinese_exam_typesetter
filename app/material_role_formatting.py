"""Formatting pass for inline material titles, authors, sources and labels."""

from __future__ import annotations

from pathlib import Path
from typing import Any, Iterator

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt

from .notice_material_formatting import apply_notice_material_formatting


def apply_material_role_formatting(
    docx_path: str | Path,
    raw_exam: dict[str, Any],
) -> None:
    """Apply v7 and all material paragraph roles recorded by importer v11."""

    target = Path(docx_path)
    apply_notice_material_formatting(target, raw_exam)
    document = Document(target)
    role_texts: dict[str, set[str]] = {
        "subheading": set(),
        "author": set(),
        "source": set(),
        "label": set(),
    }
    for block in raw_exam.get("blocks", []):
        if block.get("type") != "material":
            continue
        paragraphs = [str(value).strip() for value in block.get("paragraphs", [])]
        roles = [str(value) for value in block.get("paragraph_roles", [])]
        for index, text in enumerate(paragraphs):
            if index < len(roles) and roles[index] in role_texts:
                role_texts[roles[index]].add(text)

    for paragraph in _all_paragraphs(document):
        text = paragraph.text.strip()
        if text in role_texts["subheading"]:
            _paragraph_style(paragraph, "SimHei", 10.5, "center")
        elif text in role_texts["author"]:
            _paragraph_style(paragraph, "FangSong", 10.5, "center")
        elif text in role_texts["source"]:
            _paragraph_style(paragraph, "FangSong", 10.5, "right")
        elif text in role_texts["label"]:
            _paragraph_style(paragraph, "SimHei", 10.5, "left")
    document.save(target)


def _paragraph_style(
    paragraph: Any,
    font_name: str,
    size: float,
    alignment: str,
) -> None:
    paragraph.alignment = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
    }[alignment]
    paragraph.paragraph_format.left_indent = Pt(0)
    paragraph.paragraph_format.first_line_indent = Pt(0)
    for run in paragraph.runs:
        run.font.name = font_name
        run.font.size = Pt(size)
        fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
        for key in ("ascii", "hAnsi", "eastAsia", "cs"):
            fonts.set(qn(f"w:{key}"), font_name)


def _all_paragraphs(document: Any) -> Iterator[Any]:
    yield from document.paragraphs
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs


__all__ = ["apply_material_role_formatting"]
