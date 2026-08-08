"""DOCX importer support for Word automatic numbering."""

from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
import re
import tempfile
from typing import Any

from docx import Document
from docx.oxml.ns import qn

from .flexible_importers_v12 import import_exam as import_exam_v12
from .flexible_importers_v12 import parse_plain_lines, save_exam


@dataclass(frozen=True, slots=True)
class NumberingLevel:
    num_fmt: str
    level_text: str
    start: int


def import_exam(path: str | Path) -> dict[str, Any]:
    """Import DOCX while materializing Word list numbers for semantic parsing."""

    source = Path(path)
    if source.suffix.lower() != ".docx":
        return import_exam_v12(source)
    document = Document(source)
    inserted = materialize_automatic_numbering(document)
    if not inserted:
        return import_exam_v12(source)
    with tempfile.TemporaryDirectory(prefix="exam-numbering-") as folder:
        prepared = Path(folder) / source.name
        document.save(prepared)
        result = import_exam_v12(prepared)
    metadata = result.setdefault("metadata", {})
    metadata["source_docx_path"] = str(source.resolve())
    result.setdefault("diagnostics", []).append(
        {
            "code": "word-automatic-numbering",
            "message": f"已读取 Word 自动编号段落 {inserted} 处",
        }
    )
    return result


def materialize_automatic_numbering(document: Any) -> int:
    """Prefix paragraph text with the list label Word normally paints on screen."""

    definitions = _numbering_definitions(document)
    counters: dict[tuple[int, int], int] = {}
    inserted = 0
    for paragraph in document.paragraphs:
        key = _paragraph_numbering(paragraph)
        if key is None or key not in definitions:
            continue
        num_id, level = key
        definition = definitions[key]
        counter_key = (num_id, level)
        value = counters.get(counter_key, definition.start - 1) + 1
        counters[counter_key] = value
        prefix = _format_prefix(definition, value)
        if not prefix or _already_has_prefix(paragraph.text, prefix):
            continue
        _prepend_text(paragraph, prefix)
        inserted += 1
    return inserted


def _numbering_definitions(document: Any) -> dict[tuple[int, int], NumberingLevel]:
    root = document.part.numbering_part.element
    abstract: dict[int, dict[int, NumberingLevel]] = {}
    for item in root.findall(qn("w:abstractNum")):
        abstract_id = int(item.get(qn("w:abstractNumId")))
        levels: dict[int, NumberingLevel] = {}
        for level in item.findall(qn("w:lvl")):
            level_id = int(level.get(qn("w:ilvl"), "0"))
            fmt = level.find(qn("w:numFmt"))
            text = level.find(qn("w:lvlText"))
            start = level.find(qn("w:start"))
            levels[level_id] = NumberingLevel(
                str(fmt.get(qn("w:val"), "decimal")) if fmt is not None else "decimal",
                str(text.get(qn("w:val"), "%1.")) if text is not None else "%1.",
                int(start.get(qn("w:val"), "1")) if start is not None else 1,
            )
        abstract[abstract_id] = levels
    result: dict[tuple[int, int], NumberingLevel] = {}
    for item in root.findall(qn("w:num")):
        num_id = int(item.get(qn("w:numId")))
        abstract_ref = item.find(qn("w:abstractNumId"))
        if abstract_ref is None:
            continue
        abstract_id = int(abstract_ref.get(qn("w:val")))
        levels = dict(abstract.get(abstract_id, {}))
        for override in item.findall(qn("w:lvlOverride")):
            level_id = int(override.get(qn("w:ilvl"), "0"))
            current = levels.get(level_id, NumberingLevel("decimal", "%1.", 1))
            start_override = override.find(qn("w:startOverride"))
            if start_override is not None:
                current = NumberingLevel(
                    current.num_fmt,
                    current.level_text,
                    int(start_override.get(qn("w:val"), "1")),
                )
            levels[level_id] = current
        for level_id, definition in levels.items():
            result[(num_id, level_id)] = definition
    return result


def _paragraph_numbering(paragraph: Any) -> tuple[int, int] | None:
    num_pr = None
    if paragraph._p.pPr is not None:
        num_pr = paragraph._p.pPr.numPr
    if num_pr is None and paragraph.style is not None:
        style_ppr = paragraph.style.element.pPr
        if style_ppr is not None:
            num_pr = style_ppr.numPr
    if num_pr is None or num_pr.numId is None:
        return None
    num_id = int(num_pr.numId.val)
    level = int(num_pr.ilvl.val) if num_pr.ilvl is not None else 0
    return num_id, level


def _format_prefix(definition: NumberingLevel, value: int) -> str:
    if definition.num_fmt == "upperLetter":
        label = _letters(value)
    elif definition.num_fmt == "lowerLetter":
        label = _letters(value).lower()
    elif definition.num_fmt == "decimal":
        label = str(value)
    else:
        return ""
    prefix = definition.level_text.replace("%1", label)
    if re.fullmatch(r"[A-Za-z0-9]+\.", prefix):
        prefix = prefix[:-1] + "．"
    elif re.fullmatch(r"\([A-Za-z0-9]+\)", prefix):
        prefix = "（" + prefix[1:-1] + "）"
    return prefix


def _letters(value: int) -> str:
    result = ""
    current = max(1, int(value))
    while current:
        current, remainder = divmod(current - 1, 26)
        result = chr(ord("A") + remainder) + result
    return result


def _already_has_prefix(text: str, prefix: str) -> bool:
    stripped = str(text).lstrip()
    plain = prefix.replace("．", ".").replace("（", "(").replace("）", ")")
    return stripped.startswith(prefix) or stripped.startswith(plain)


def _prepend_text(paragraph: Any, prefix: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = prefix + paragraph.runs[0].text
        return
    paragraph.add_run(prefix)


__all__ = [
    "NumberingLevel",
    "import_exam",
    "materialize_automatic_numbering",
    "parse_plain_lines",
    "save_exam",
]
