from __future__ import annotations

from pathlib import Path
import tempfile
import unittest

from docx import Document
from docx.oxml.ns import qn

from app.chinese_typography_v2 import enable_chinese_typography_v2
from app.exam_format_rules_v2 import apply_exam_format_rules_v2
from app.flexible_importers_v4 import _normalize_question_details, import_exam
from app.semantic_formatting_v3 import apply_semantic_formatting_v3


class V070RuleTests(unittest.TestCase):
    def test_segmentation_and_subquestions_are_promoted(self) -> None:
        data = {
            "blocks": [
                {
                    "type": "question",
                    "question": {
                        "number": 10,
                        "embedded_segments": [
                            [{"text": "甲A乙B丙C丁D戊E己F庚G辛H壬", "role": "body"}]
                        ],
                    },
                },
                {
                    "type": "question",
                    "question": {
                        "number": 13,
                        "embedded_segments": [
                            [{"text": "（1）第一句", "role": "body"}],
                            [{"text": "（2）第二句", "role": "body"}],
                        ],
                    },
                },
            ]
        }
        _normalize_question_details(data)
        self.assertEqual(
            data["blocks"][0]["question"]["segmentation_text"],
            "甲A乙B丙C丁D戊E己F庚G辛H壬",
        )
        self.assertEqual(
            data["blocks"][1]["question"]["subquestions"],
            ["第一句", "第二句"],
        )

    def test_fixed_indents_and_fonts(self) -> None:
        with tempfile.TemporaryDirectory() as folder:
            path = Path(folder) / "rules.docx"
            document = Document()
            question = document.add_paragraph("10．题干内容")
            option = document.add_paragraph("A．选择项内容")
            sub = document.add_paragraph("（1）翻译句子")
            sub.style = document.styles["Normal"]
            sub.style.name
            sub._p.get_or_add_pPr()
            document.save(path)
            result = Document(path)
            result.paragraphs[2].style = result.styles.add_style(
                "Exam_subquestion",
                1,
            )
            result.save(path)
            apply_exam_format_rules_v2(path)
            result = Document(path)
            self.assertAlmostEqual(
                result.paragraphs[0].paragraph_format.left_indent.pt,
                15.75,
            )
            self.assertAlmostEqual(
                result.paragraphs[0].paragraph_format.first_line_indent.pt,
                -15.75,
            )
            self.assertAlmostEqual(
                result.paragraphs[1].paragraph_format.first_line_indent.pt,
                -17.85,
            )
            self.assertAlmostEqual(
                result.paragraphs[2].paragraph_format.first_line_indent.pt,
                21,
            )

    def test_broad_source_and_paragraph_kinsoku(self) -> None:
        with tempfile.TemporaryDirectory() as folder:
            path = Path(folder) / "chinese.docx"
            document = Document()
            document.add_paragraph("（本文写于1959年，有删改）")
            document.add_paragraph("正文，标点不得出现在下一行开头。")
            document.save(path)
            apply_semantic_formatting_v3(path)
            enable_chinese_typography_v2(path)
            result = Document(path)
            source = result.paragraphs[0]
            self.assertEqual(source.alignment, 2)
            self.assertEqual(
                source.runs[0]._element.rPr.rFonts.get(qn("w:eastAsia")),
                "FangSong",
            )
            for paragraph in result.paragraphs:
                ppr = paragraph._p.pPr
                self.assertEqual(ppr.find(qn("w:kinsoku")).get(qn("w:val")), "1")
                self.assertEqual(
                    ppr.find(qn("w:overflowPunct")).get(qn("w:val")),
                    "1",
                )

    def test_pdf_import_is_rejected(self) -> None:
        with tempfile.TemporaryDirectory() as folder:
            path = Path(folder) / "sample.pdf"
            path.write_bytes(b"%PDF-1.4")
            with self.assertRaisesRegex(ValueError, "不再提供 PDF 导入"):
                import_exam(path)


if __name__ == "__main__":
    unittest.main()
