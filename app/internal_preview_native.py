"""Complete internal preview for native tables and appended answer blocks."""

from __future__ import annotations

from copy import deepcopy
from pathlib import Path
import re
from threading import Lock
from typing import Any

from docx import Document
from PIL import Image, ImageDraw

from . import internal_preview_core as preview_core


_ORIGINAL_VISUAL_LOADER = preview_core._load_source_visuals
_NATIVE_OBJECT_RE = re.compile(r"^\[\[NATIVE_(?:DRAWING|TABLE):\d+\]\]$")
_PATCH_LOCK = Lock()


def _table_rows(table: Any) -> list[list[str]]:
    """Return visible cell text while suppressing duplicate merged cells."""

    rows: list[list[str]] = []
    for row in table.rows:
        seen: set[int] = set()
        values: list[str] = []
        for cell in row.cells:
            identity = id(cell._tc)
            if identity in seen:
                values.append("")
            else:
                seen.add(identity)
                values.append("\n".join(part.strip() for part in cell.text.splitlines() if part.strip()))
        rows.append(values)
    return rows


def _render_table_image(table: Any, *, max_width: int = 620) -> Image.Image:
    """Render table borders and every cell string into a preview image."""

    rows = _table_rows(table)
    columns = max((len(row) for row in rows), default=1)
    width = max(320, int(max_width))
    column_width = max(56, width // max(1, columns))
    width = column_width * columns
    padding = 6
    font = preview_core._font("SimSun", 10.5)
    line_height = max(17, int(getattr(font, "size", 14) * 1.3))
    scratch = ImageDraw.Draw(Image.new("RGB", (1, 1), "white"))
    wrapped_rows: list[list[list[str]]] = []
    row_heights: list[int] = []
    for row in rows or [[""]]:
        wrapped: list[list[str]] = []
        maximum = 1
        for index in range(columns):
            value = row[index] if index < len(row) else ""
            lines: list[str] = []
            for paragraph in value.splitlines() or [""]:
                lines.extend(
                    preview_core._wrap_text(
                        scratch,
                        paragraph,
                        font,
                        max(20, column_width - 2 * padding),
                    )
                )
            lines = lines or [""]
            maximum = max(maximum, len(lines))
            wrapped.append(lines)
        wrapped_rows.append(wrapped)
        row_heights.append(maximum * line_height + 2 * padding)
    height = max(24, sum(row_heights) + 1)
    image = Image.new("RGBA", (width + 1, height), "white")
    draw = ImageDraw.Draw(image)
    y = 0
    for row_index, wrapped in enumerate(wrapped_rows):
        row_height = row_heights[row_index]
        x = 0
        for column_index, lines in enumerate(wrapped):
            draw.rectangle(
                (x, y, x + column_width, y + row_height),
                outline="#4F4F4F",
                width=1,
            )
            text_y = y + padding
            for line in lines:
                draw.text((x + padding, text_y), line, fill="#111111", font=font)
                text_y += line_height
            x += column_width
        y += row_height
    return image


def _load_source_visuals(raw_exam: dict[str, Any]) -> Any:
    visuals = _ORIGINAL_VISUAL_LOADER(raw_exam)
    metadata = raw_exam.get("metadata", {})
    source = Path(str(metadata.get("source_docx_path", "")))
    if not source.is_file():
        return visuals
    try:
        document = Document(source)
    except (OSError, ValueError, RuntimeError):
        return visuals
    for item in metadata.get("native_objects", []):
        if not isinstance(item, dict) or item.get("kind") != "table":
            continue
        marker = str(item.get("marker", "")).strip()
        try:
            table_index = int(item.get("source_table_index", -1))
        except (TypeError, ValueError):
            continue
        if not marker or not (0 <= table_index < len(document.tables)):
            continue
        visuals.images.setdefault(marker, []).append(
            [_render_table_image(document.tables[table_index])]
        )
    return visuals


def _answer_paragraph_text(entry: Any) -> str:
    if isinstance(entry, dict):
        return str(entry.get("text", ""))
    return str(entry)


def _answer_block_for_preview(block: dict[str, Any]) -> dict[str, Any]:
    """Map answer model blocks to semantic material blocks without losing text."""

    kind = str(block.get("type", ""))
    if kind == "answer_section":
        return {"type": "section_title", "text": str(block.get("text", block.get("header", "")))}
    if kind == "answer_subsection":
        return {
            "type": "subsection",
            "name": str(block.get("name", block.get("header", ""))),
            "meta": str(block.get("meta", "")),
        }
    if kind == "answer_table":
        marker = f"[[NATIVE_TABLE:{int(block.get('source_table_index', 0))}]]"
        return {"type": "material", "paragraphs": [marker], "paragraph_roles": ["body"]}
    if kind not in {"answer_question", "answer_text"}:
        return block
    header = str(block.get("header", "")).strip()
    paragraphs = [_answer_paragraph_text(value) for value in block.get("paragraphs", [])]
    values = ([header] if header else []) + [value for value in paragraphs if value.strip()]
    formats: list[dict[str, Any]] = []
    roles = ["body"] * len(values)
    if header:
        formats.append(
            {
                "target_index": 0,
                "font": "SimSun",
                "size_pt": 10.5,
                "bold": True,
                "alignment": "left",
                "first_line_indent_chars": 0,
                "line_spacing": 1.25,
                "space_after_pt": 1,
            }
        )
    return {
        "type": "material",
        "paragraphs": values,
        "paragraph_roles": roles,
        "paragraph_formats": formats,
    }


def preview_exam(raw_exam: dict[str, Any]) -> dict[str, Any]:
    """Return a preview-only copy whose every supported block carries visible text."""

    prepared = deepcopy(raw_exam)
    prepared["blocks"] = [
        _answer_block_for_preview(block)
        for block in prepared.get("blocks", [])
    ]
    return prepared


def render_internal_preview(
    raw_exam: dict[str, Any],
    layout_path: str | Path,
    output_dir: str | Path,
) -> Any:
    """Render tables, drawings, questions, materials and appended answers."""

    prepared = preview_exam(raw_exam)
    with _PATCH_LOCK:
        previous_loader = preview_core._load_source_visuals
        previous_marker = preview_core._NATIVE_DRAWING_RE
        preview_core._load_source_visuals = _load_source_visuals
        preview_core._NATIVE_DRAWING_RE = _NATIVE_OBJECT_RE
        try:
            return preview_core.render_internal_preview(prepared, layout_path, output_dir)
        finally:
            preview_core._load_source_visuals = previous_loader
            preview_core._NATIVE_DRAWING_RE = previous_marker


__all__ = [
    "_answer_block_for_preview",
    "_render_table_image",
    "_table_rows",
    "preview_exam",
    "render_internal_preview",
]
