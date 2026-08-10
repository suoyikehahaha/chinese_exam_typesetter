"""Context-aware formatting for header dates, sources and memorization text."""

from __future__ import annotations

from pathlib import Path
import re
from typing import Any, Iterator

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt


DATE_RE = re.compile(r"^20\d{2}\s*[./年-]\s*\d{1,2}(?:\s*月)?$")
SOURCE_RE = re.compile(
    r"^\s*[（(]\s*(?:"
    r"(?:摘自|摘编自|选自|节选自|改编自|据).+|"
    r"本文.+|有删改|有改动"
    r")[）)]\s*$"
)


def apply_exam_context_formatting(
    docx_path: str | Path,
    raw_exam: dict[str, Any],
) -> None:
    """Apply confirmed styles whose meaning depends on surrounding structure."""

    target = Path(docx_path)
    document = Document(target)
    memorization_texts = _memorization_texts(raw_exam)
    meta_text = str(raw_exam.get("metadata", {}).get("meta_text", "")).strip()
    for paragraph in _all_paragraphs(document):
        text = paragraph.text.strip()
        if text and (text == meta_text or DATE_RE.fullmatch(text)):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            paragraph.paragraph_format.left_indent = Pt(0)
            paragraph.paragraph_format.first_line_indent = Pt(0)
            for run in paragraph.runs:
                _font(run, "SimSun", 10.5)
        if SOURCE_RE.fullmatch(text):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            paragraph.paragraph_format.left_indent = Pt(0)
            paragraph.paragraph_format.first_line_indent = Pt(0)
            for run in paragraph.runs:
                _font(run, "FangSong", 10.5)
        if text in memorization_texts:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.left_indent = Pt(0)
            paragraph.paragraph_format.first_line_indent = Pt(10.5 * 2)
            for run in paragraph.runs:
                _font(run, "SimSun", 10.5)
    document.save(target)


def _memorization_texts(raw_exam: dict[str, Any]) -> set[str]:
    result: set[str] = set()
    in_memorization = False
    for block in raw_exam.get("blocks", []):
        if block.get("type") == "subsection":
            label = f"{block.get('name', '')}{block.get('meta', '')}"
            in_memorization = "名篇名句默写" in label or "默写" in label
            continue
        question = block.get("question")
        if not question:
            continue
        stem = str(question.get("stem", ""))
        is_memorization = in_memorization or (
            "补写出下列句子中的空缺部分" in stem
            or "名篇名句默写" in stem
        )
        if not is_memorization:
            continue
        for segments in question.get("embedded_segments", []):
            text = "".join(str(segment.get("text", "")) for segment in segments).strip()
            if text:
                result.add(text)
        for text in question.get("subquestions", []):
            if str(text).strip():
                result.add(str(text).strip())
    return result


def _font(run: Any, name: str, size: float) -> None:
    run.font.name = name
    run.font.size = Pt(size)
    rfonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    for key in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn(f"w:{key}"), name)


def _all_paragraphs(document: Any) -> Iterator[Any]:
    yield from document.paragraphs
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs
