from pathlib import Path
import tempfile
import unittest

from docx import Document

from app.config import load_layout
from app.models import load_exam
from app.pagination import apply_pagination_guards
from app.renderers import DocxRenderer


ROOT = Path(__file__).resolve().parents[1]


class PaginationTests(unittest.TestCase):
    def test_poetry_and_composition_guards(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            target = Path(directory) / "sample.docx"
            DocxRenderer(load_layout(ROOT / "templates" / "layout.yaml")).render(
                load_exam(ROOT / "samples" / "exam.json"), target
            )
            apply_pagination_guards(target)
            document = Document(target)
            composition = next(p for p in document.paragraphs if p.text.startswith("23．"))
            self.assertTrue(composition.paragraph_format.keep_with_next)
            poetry = [p for p in document.paragraphs if p.style.name == "Exam_poetry"]
            self.assertTrue(poetry[0].paragraph_format.keep_with_next)
            self.assertFalse(poetry[3].paragraph_format.keep_with_next)


if __name__ == "__main__":
    unittest.main()
