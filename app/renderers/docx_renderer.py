"""使用 python-docx 渲染结构化试卷。"""

from __future__ import annotations

from copy import deepcopy
from pathlib import Path
import re
from typing import Any

from docx import Document
from docx.document import Document as DocumentObject
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt

from app.models import Block, ExamDocument, InlineSegment, Question


ALIGNMENTS = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}


class DocxRenderer:
    """把 ExamDocument 渲染为 DOCX。"""

    def __init__(self, layout: dict[str, Any], template_path: str | Path | None = None):
        self.layout = layout
        self.template_path = Path(template_path) if template_path else None
        self.doc: DocumentObject | None = None

    def render(self, exam: ExamDocument, output_path: str | Path) -> Path:
        """生成 DOCX 并返回输出路径。"""

        if self.template_path:
            if not self.template_path.exists():
                raise FileNotFoundError(f"母版不存在：{self.template_path}")
            self.doc = Document(str(self.template_path))
            self._clear_body(self.doc)
        else:
            self.doc = Document()

        self._configure_page()
        self._create_styles()
        self._render_header(exam)
        for block in exam.blocks:
            self._render_block(block)
        self._add_page_number()

        target = Path(output_path)
        target.parent.mkdir(parents=True, exist_ok=True)
        self.doc.save(target)
        return target

    @property
    def document(self) -> DocumentObject:
        if self.doc is None:
            raise RuntimeError("渲染器尚未初始化文档")
        return self.doc

    @staticmethod
    def _clear_body(doc: DocumentObject) -> None:
        body = doc._element.body
        for child in list(body):
            if child.tag != qn("w:sectPr"):
                body.remove(child)

    def _configure_page(self) -> None:
        page = self.layout["page"]
        for section in self.document.sections:
            section.page_width = Mm(page["width_mm"])
            section.page_height = Mm(page["height_mm"])
            section.top_margin = Mm(page["margin_top_mm"])
            section.bottom_margin = Mm(page["margin_bottom_mm"])
            section.left_margin = Mm(page["margin_left_mm"])
            section.right_margin = Mm(page["margin_right_mm"])
            section.header_distance = Mm(page["header_distance_mm"])
            section.footer_distance = Mm(page["footer_distance_mm"])

    def _create_styles(self) -> None:
        for role, spec in self.layout["styles"].items():
            name = f"Exam_{role}"
            styles = self.document.styles
            style = styles[name] if name in styles else styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
            style.font.name = spec["font"]
            style.font.size = Pt(spec["size_pt"])
            style.font.bold = bool(spec.get("bold", False))
            rpr = style.element.get_or_add_rPr()
            rfonts = rpr.get_or_add_rFonts()
            for key in ("ascii", "hAnsi", "eastAsia", "cs"):
                rfonts.set(qn(f"w:{key}"), spec["font"])
            self._apply_paragraph_format(style.paragraph_format, spec)

    def _apply_paragraph_format(self, fmt: Any, spec: dict[str, Any]) -> None:
        size = float(spec["size_pt"])
        if "alignment" in spec:
            fmt.alignment = ALIGNMENTS[spec["alignment"]]
        if "first_line_indent_chars" in spec:
            fmt.first_line_indent = Pt(size * float(spec["first_line_indent_chars"]))
        if "left_indent_chars" in spec:
            fmt.left_indent = Pt(size * float(spec["left_indent_chars"]))
        if "hanging_indent_chars" in spec:
            fmt.first_line_indent = Pt(-size * float(spec["hanging_indent_chars"]))
        fmt.space_before = Pt(float(spec.get("space_before_pt", 0)))
        fmt.space_after = Pt(float(spec.get("space_after_pt", self.layout["defaults"]["paragraph_after_pt"])))
        fmt.line_spacing = float(spec.get("line_spacing", self.layout["defaults"]["body_line_spacing"]))
        fmt.keep_with_next = bool(spec.get("keep_with_next", False))
        fmt.widow_control = True

    def _render_header(self, exam: ExamDocument) -> None:
        metadata = exam.metadata
        self._paragraph(metadata.exam_name, "exam_name")
        self._paragraph(metadata.subject_name, "subject_name")
        self._paragraph(metadata.meta_text, "exam_meta")
        if metadata.notices:
            self._paragraph("注意事项：", "notice_title")
            for index, notice in enumerate(metadata.notices, start=1):
                self._paragraph(f"{index}．{notice}", "notice_body")

    def _render_block(self, block: Block) -> None:
        handlers = {
            "section_title": lambda: self._paragraph(block.text, "section_title"),
            "subsection": lambda: self._subsection(block.name, block.meta),
            "instruction": lambda: self._paragraph(block.text, "instruction"),
            "material": lambda: self._material(block),
            "poetry": lambda: self._poetry(block),
            "question": lambda: self._question(block.question),
            "page_break": lambda: self.document.add_page_break(),
        }
        if block.type not in handlers:
            raise ValueError(f"未知内容块类型：{block.type}")
        handlers[block.type]()

    def _subsection(self, name: str, meta: str) -> None:
        paragraph = self.document.add_paragraph(style="Exam_subsection")
        self._run(paragraph, name, "subsection", bold=True)
        if meta:
            self._run(paragraph, meta, "subsection", bold=False)

    def _material(self, block: Block) -> None:
        if block.title:
            self._paragraph(block.title, "material_title")
        if block.author:
            self._paragraph(block.author, "material_author")
        for text in block.paragraphs:
            self._paragraph(text, "material_body")
        if block.note:
            self._paragraph(block.note, "material_note")
        if block.source:
            self._paragraph(block.source, "material_source")

    def _poetry(self, block: Block) -> None:
        if block.title:
            self._paragraph(block.title, "material_title")
        if block.author:
            self._paragraph(block.author, "material_author")
        for line in block.paragraphs:
            self._paragraph(line, "poetry")
        if block.note:
            self._paragraph(block.note, "material_note")

    def _question(self, question: Question | None) -> None:
        if question is None:
            raise ValueError("question 内容块缺少 question 数据")
        role = "objective_stem" if question.kind == "objective" else "subjective_stem"
        score = f"（{self._format_score(question.score)}分）" if question.score is not None else ""
        self._paragraph(f"{question.number}．{question.stem}{score}", role)

        if question.options:
            if question.option_layout == "two_column":
                self._two_column_options(question.options)
            else:
                for option in question.options:
                    self._paragraph(option, "choice_option")

        for segments in question.embedded_segments:
            self._mixed_paragraph(segments)

        if question.segmentation_text:
            self._segmentation(question.segmentation_text)

        for index, text in enumerate(question.subquestions, start=1):
            paragraph = self.document.add_paragraph(style="Exam_subquestion")
            self._run(paragraph, f"（{index}）", "subquestion", bold=True)
            self._run(paragraph, text, "subquestion")

        for text in question.composition_material:
            self._paragraph(text, "composition_material")
        for text in question.composition_prompt:
            self._paragraph(text, "composition_prompt")
        for text in question.composition_requirements:
            self._paragraph(text, "composition_requirements")

    @staticmethod
    def _format_score(score: float) -> str:
        return str(int(score)) if score.is_integer() else str(score)

    def _mixed_paragraph(self, segments: list[InlineSegment]) -> None:
        paragraph = self.document.add_paragraph(style="Exam_embedded_body")
        for segment in segments:
            role = "embedded_label" if segment.role == "label" else "embedded_body"
            self._run(paragraph, segment.text, role)

    def _segmentation(self, text: str) -> None:
        paragraph = self.document.add_paragraph(style="Exam_segmentation_text")
        for token in re.split(r"([A-H])", text):
            if not token:
                continue
            role = "segmentation_marker" if len(token) == 1 and token in "ABCDEFGH" else "segmentation_text"
            self._run(paragraph, token, role, boxed=(role == "segmentation_marker"))

    def _two_column_options(self, options: list[str]) -> None:
        if len(options) != 4:
            raise ValueError("双列选择项必须正好包含四项")
        table = self.document.add_table(rows=2, cols=2)
        table.autofit = False
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        usable_mm = (
            self.layout["page"]["width_mm"]
            - self.layout["page"]["margin_left_mm"]
            - self.layout["page"]["margin_right_mm"]
        )
        for index, text in enumerate(options):
            cell = table.cell(index // 2, index % 2)
            cell.width = Mm(usable_mm / 2)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            paragraph = cell.paragraphs[0]
            paragraph.style = "Exam_choice_option"
            self._run(paragraph, text, "choice_option")
        self._remove_table_borders(table)

    @staticmethod
    def _remove_table_borders(table: Any) -> None:
        tbl_pr = table._tbl.tblPr
        borders = tbl_pr.first_child_found_in("w:tblBorders")
        if borders is None:
            borders = OxmlElement("w:tblBorders")
            tbl_pr.append(borders)
        for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
            element = borders.find(qn(f"w:{edge}"))
            if element is None:
                element = OxmlElement(f"w:{edge}")
                borders.append(element)
            element.set(qn("w:val"), "nil")

    def _paragraph(self, text: str, role: str) -> Any:
        paragraph = self.document.add_paragraph(style=f"Exam_{role}")
        self._run(paragraph, text, role)
        return paragraph

    def _run(
        self,
        paragraph: Any,
        text: str,
        role: str,
        *,
        bold: bool | None = None,
        boxed: bool = False,
    ) -> Any:
        spec = self.layout["styles"][role]
        run = paragraph.add_run(text)
        run.font.name = spec["font"]
        run.font.size = Pt(spec["size_pt"])
        run.bold = bool(spec.get("bold", False)) if bold is None else bold
        rpr = run._element.get_or_add_rPr()
        rfonts = rpr.get_or_add_rFonts()
        for key in ("ascii", "hAnsi", "eastAsia", "cs"):
            rfonts.set(qn(f"w:{key}"), spec["font"])
        if boxed or spec.get("boxed"):
            border = OxmlElement("w:bdr")
            border.set(qn("w:val"), "single")
            border.set(qn("w:sz"), "4")
            border.set(qn("w:space"), "1")
            border.set(qn("w:color"), "auto")
            rpr.append(border)
        return run

    def _add_page_number(self) -> None:
        for section in self.document.sections:
            paragraph = section.footer.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run()
            begin = OxmlElement("w:fldChar")
            begin.set(qn("w:fldCharType"), "begin")
            instruction = OxmlElement("w:instrText")
            instruction.set(qn("xml:space"), "preserve")
            instruction.text = " PAGE "
            separate = OxmlElement("w:fldChar")
            separate.set(qn("w:fldCharType"), "separate")
            text = OxmlElement("w:t")
            text.text = "1"
            end = OxmlElement("w:fldChar")
            end.set(qn("w:fldCharType"), "end")
            for node in (begin, instruction, separate, text, end):
                run._r.append(node)
