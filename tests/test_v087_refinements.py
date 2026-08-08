from __future__ import annotations

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

from app.contextual_formatting_v10 import apply_contextual_formatting_v10
from app.flexible_importers_v12 import (
    PUBLICATION_NOTE_RE,
    mark_publication_notes,
)
from app.windows_activation_v1 import ACTIVE, INACTIVE


def test_publication_note_pattern_excludes_source_provenance() -> None:
    assert PUBLICATION_NOTE_RE.fullmatch("（马小弥译，有删改）")
    assert PUBLICATION_NOTE_RE.fullmatch("（发表于1938年10月，有删改）")
    assert PUBLICATION_NOTE_RE.fullmatch("（原载《某刊》，有改动）")
    assert not PUBLICATION_NOTE_RE.fullmatch("（摘编自刘妍《某文》）")
    assert not PUBLICATION_NOTE_RE.fullmatch("（节选自某书，有删改）")


def test_mark_publication_notes_keeps_source_role() -> None:
    raw = {
        "blocks": [
            {
                "type": "material",
                "paragraphs": [
                    "正文",
                    "（马小弥译，有删改）",
                    "（摘编自刘妍《某文》）",
                ],
                "paragraph_roles": ["body", "body", "source"],
            }
        ]
    }
    mark_publication_notes(raw)
    block = raw["blocks"][0]
    assert block["paragraph_roles"] == ["body", "publication_note", "source"]
    spec = next(
        value
        for value in block["paragraph_formats"]
        if value["semantic_role"] == "publication_note"
    )
    assert spec["font"] == "宋体"
    assert spec["size_pt"] == 10.5
    assert spec["alignment"] == "右对齐"


def test_contextual_formatting_applies_simsun_right_alignment(tmp_path) -> None:
    target = tmp_path / "publication-note.docx"
    document = Document()
    document.add_paragraph("正文")
    document.add_paragraph("（发表于1938年10月，有删改）")
    document.save(target)
    raw = {
        "blocks": [
            {
                "type": "material",
                "paragraphs": ["正文", "（发表于1938年10月，有删改）"],
                "paragraph_roles": ["body", "publication_note"],
            }
        ]
    }

    apply_contextual_formatting_v10(target, raw)

    result = Document(target)
    paragraph = result.paragraphs[1]
    assert paragraph.alignment == WD_ALIGN_PARAGRAPH.RIGHT
    assert paragraph.paragraph_format.first_line_indent.pt == 0
    assert all(run.font.size.pt == 10.5 for run in paragraph.runs)
    assert all(
        run._element.get_or_add_rPr()
        .get_or_add_rFonts()
        .get(qn("w:eastAsia"))
        == "SimSun"
        for run in paragraph.runs
    )


def test_active_and_inactive_palettes_are_visibly_distinct() -> None:
    assert ACTIVE["toolbar"] != INACTIVE["toolbar"]
    assert ACTIVE["status"] != INACTIVE["status"]
    assert ACTIVE["selection"] != INACTIVE["selection"]
