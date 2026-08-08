from pathlib import Path
import tempfile
import unittest

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

from app.semantic_formatting import apply_semantic_formatting


class SemanticFormattingTests(unittest.TestCase):
    def test_material_label_and_source_styles(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            path = Path(directory) / "semantic.docx"
            document = Document()
            document.add_paragraph("材料一：")
            document.add_paragraph("材料二：正文从同一行开始。")
            document.add_paragraph("（摘编自张新科《史记体现的大一统观念》）")
            document.save(path)
            apply_semantic_formatting(path)
            result = Document(path)
            first, second, source = result.paragraphs
            self.assertEqual("SimHei", first.runs[0]._element.rPr.rFonts.get(qn("w:eastAsia")))
            self.assertEqual("SimHei", second.runs[0].font.name)
            self.assertEqual("KaiTi", second.runs[1].font.name)
            self.assertEqual("FangSong", source.runs[0].font.name)
            self.assertEqual(WD_ALIGN_PARAGRAPH.RIGHT, source.alignment)


if __name__ == "__main__":
    unittest.main()
