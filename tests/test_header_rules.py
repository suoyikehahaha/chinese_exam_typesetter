from __future__ import annotations

import tempfile
import unittest
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

from app.header_normalization import normalize_header


class HeaderRuleTests(unittest.TestCase):
    def test_regional_institution_date_is_idempotent(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            target = Path(directory) / "header.docx"
            document = Document()
            for _ in range(3):
                document.add_paragraph("武汉市教育科学研究院命制\t2026.4")
            document.add_paragraph(
                "本卷共10页，23题。全卷满分150分。用时150分钟。"
            )
            document.save(target)
            raw = {
                "metadata": {
                    "institution_text": "武汉市教育科学研究院命制",
                    "exam_date": "2026.4",
                    "exam_info_text": (
                        "本卷共10页，23题。全卷满分150分。用时150分钟。"
                    ),
                }
            }

            normalize_header(target, raw)

            result = Document(target)
            institution_lines = [
                paragraph
                for paragraph in result.paragraphs
                if "武汉市教育科学研究院命制" in paragraph.text
                and "2026.4" in paragraph.text
            ]
            self.assertEqual(len(institution_lines), 1)

    def test_exam_information_is_centered_simsun_five(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            target = Path(directory) / "meta.docx"
            text = "本试卷共23小题，满分150分。考试用时150分钟。"
            document = Document()
            document.add_paragraph(text)
            document.save(target)

            normalize_header(
                target,
                {"metadata": {"meta_text": text}},
            )

            paragraph = Document(target).paragraphs[0]
            self.assertEqual(paragraph.alignment, WD_ALIGN_PARAGRAPH.CENTER)
            self.assertEqual(paragraph.paragraph_format.first_line_indent.pt, 0)
            self.assertEqual(paragraph.runs[0].font.size.pt, 10.5)
            fonts = (
                paragraph.runs[0]
                ._element.get_or_add_rPr()
                .get_or_add_rFonts()
            )
            self.assertEqual(fonts.get(qn("w:eastAsia")), "SimSun")


if __name__ == "__main__":
    unittest.main()
