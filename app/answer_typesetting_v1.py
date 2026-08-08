"""Guangzhou-style answer recognition and DOCX rendering."""

from __future__ import annotations

from copy import deepcopy
from pathlib import Path
import re
from typing import Any, Iterable

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt
from docx.table import Table
from docx.text.paragraph import Paragraph

from .native_docx_objects import _remap_image_relationships


ANSWER_MARKER_RE = re.compile(
    r"^\s*(?:参考答案|答案与解析|参考答案与评分建议|"
    r"语文试题参考答案及评分建议|试题答案)\s*[：:]?\s*$"
)
MAJOR_SECTION_RE = re.compile(r"^\s*[一二三四五六七八九十]+、.+")
SUBSECTION_RE = re.compile(
    r"^\s*(?P<name>[（(][一二三四五六七八九十]+[）)]\s*.+?)"
    r"(?P<meta>[（(]\s*\d+\s*分\s*[）)])\s*$"
)
QUESTION_RE = re.compile(
    r"^\s*(?P<number>\d{1,2})\s*[．.]\s*(?P<rest>.*)$"
)
SCORE_PREFIX_RE = re.compile(
    r"^\s*(?P<prefix>[（(]\s*\d+\s*[）)]\s*"
    r"(?:[（(]\s*\d+\s*分\s*[）)])?)"
)
LABELED_ANSWER_RE = re.compile(
    r"^\s*(?P<prefix>(?:[①-⑳]处|材料[一二三四五六七八九十]+)"
    r"\s*[：:])"
)
EXAMPLE_LABEL_RE = re.compile(r"^\s*示例[一二三四五六七八九十\d]+\s*[：:]")


def is_standalone_answer_docx(path: str | Path) -> bool:
    """Return whether the opening area identifies a standalone answer document."""

    source = Path(path)
    if source.suffix.lower() != ".docx":
        return False
    document = Document(source)
    opening = [
        paragraph.text.strip()
        for paragraph in document.paragraphs[:12]
        if paragraph.text.strip()
    ]
    return any(
        "参考答案" in text or "评分建议" in text or "答案与解析" in text
        for text in opening[:5]
    )


def find_answer_start(path: str | Path) -> int | None:
    """Find the source paragraph index where an appended answer starts."""

    document = Document(path)
    for index, paragraph in enumerate(document.paragraphs):
        if ANSWER_MARKER_RE.fullmatch(paragraph.text.strip()):
            return index
    return None


def parse_answer_docx(
    path: str | Path,
    *,
    start_paragraph: int | None = None,
    fallback_title: str = "",
) -> dict[str, Any]:
    """Parse an answer DOCX into editable semantic blocks."""

    source = Path(path)
    document = Document(source)
    title_lines: list[str] = []
    blocks: list[dict[str, Any]] = []
    current: dict[str, Any] | None = None
    state = "answer"
    question_number: int | None = None
    table_index = 0
    paragraph_counter = -1

    def flush() -> None:
        nonlocal current
        if current is not None:
            if current.get("header") or current.get("paragraphs"):
                blocks.append(current)
            current = None

    started = start_paragraph is None
    skipped_marker = False
    for child in document.element.body.iterchildren():
        if child.tag == qn("w:p"):
            paragraph_counter += 1
            paragraph = Paragraph(child, document)
            source_index = paragraph_counter
            if not started:
                if source_index is not None and source_index >= int(start_paragraph):
                    started = True
                else:
                    continue
            text = paragraph.text.strip()
            if not text:
                continue
            if (
                start_paragraph is not None
                and not skipped_marker
                and ANSWER_MARKER_RE.fullmatch(text)
            ):
                skipped_marker = True
                continue
            if start_paragraph is None and not blocks and current is None:
                if _looks_like_title(text, paragraph, len(title_lines)):
                    title_lines.append(text)
                    continue

            subsection = SUBSECTION_RE.fullmatch(text)
            question = QUESTION_RE.fullmatch(text)
            if MAJOR_SECTION_RE.fullmatch(text):
                flush()
                blocks.append({"type": "answer_section", "text": text})
                state = "answer"
                question_number = None
            elif subsection:
                flush()
                blocks.append(
                    {
                        "type": "answer_subsection",
                        "name": subsection.group("name").strip(),
                        "meta": subsection.group("meta").strip(),
                    }
                )
                state = "answer"
                question_number = None
            elif question:
                flush()
                question_number = int(question.group("number"))
                current = {
                    "type": "answer_question",
                    "number": question_number,
                    "header": text,
                    "paragraphs": [],
                }
                state = "composition" if question_number == 23 else "answer"
            else:
                if current is None:
                    current = {
                        "type": "answer_text",
                        "header": "",
                        "paragraphs": [],
                    }
                role, state = _answer_role(text, state, question_number)
                current["paragraphs"].append(
                    {
                        "text": text,
                        "role": role,
                        "runs": _source_run_decorations(paragraph),
                    }
                )
        elif child.tag == qn("w:tbl"):
            if not started:
                table_index += 1
                continue
            flush()
            blocks.append(
                {
                    "type": "answer_table",
                    "source_table_index": table_index,
                }
            )
            table_index += 1

    flush()
    if not title_lines:
        title_lines = [
            fallback_title or source.stem,
            "语文试题参考答案及评分建议",
        ]
    elif len(title_lines) == 1:
        title_lines.append("语文试题参考答案及评分建议")

    return {
        "title_lines": title_lines[:2],
        "source_docx_path": str(source.resolve()),
        "source_start_paragraph": start_paragraph,
        "blocks": blocks,
    }


def standalone_answer_model(answer: dict[str, Any]) -> dict[str, Any]:
    """Expose answer blocks through the existing editable workbench model."""

    title_lines = list(answer.get("title_lines", []))
    while len(title_lines) < 2:
        title_lines.append("语文试题参考答案及评分建议")
    return {
        "document_kind": "answer",
        "metadata": {
            "exam_name": title_lines[0],
            "subject_name": title_lines[1],
            "meta_text": "",
            "notices": [],
            "total_score": 150,
            "answer_title_lines": title_lines[:2],
            "answer_source_docx_path": answer.get("source_docx_path", ""),
        },
        "blocks": deepcopy(answer.get("blocks", [])),
    }


def attach_answer_blocks(raw_exam: dict[str, Any], answer: dict[str, Any]) -> None:
    """Attach prefixed answer blocks without exposing them to exam validation."""

    metadata = raw_exam.setdefault("metadata", {})
    title_lines = list(answer.get("title_lines", []))
    if title_lines and ANSWER_MARKER_RE.fullmatch(title_lines[0]):
        title_lines[0] = str(metadata.get("exam_name", "")).strip()
    if len(title_lines) < 2:
        title_lines.append("语文试题参考答案及评分建议")
    metadata["answer_title_lines"] = title_lines[:2]
    metadata["answer_source_docx_path"] = answer.get("source_docx_path", "")
    raw_exam["document_kind"] = "exam_with_answer"
    raw_exam.setdefault("blocks", []).extend(deepcopy(answer.get("blocks", [])))


def answer_blocks(raw_exam: dict[str, Any]) -> list[dict[str, Any]]:
    """Return answer-only blocks from a standalone or combined model."""

    return [
        block
        for block in raw_exam.get("blocks", [])
        if str(block.get("type", "")).startswith("answer_")
    ]


def exam_blocks(raw_exam: dict[str, Any]) -> list[dict[str, Any]]:
    """Return blocks understood by the ordinary exam renderer."""

    return [
        block
        for block in raw_exam.get("blocks", [])
        if not str(block.get("type", "")).startswith("answer_")
    ]


def render_answer_docx(
    raw_exam: dict[str, Any],
    output_path: str | Path,
) -> Path:
    """Render a standalone answer in the Guangzhou reference layout."""

    document = Document()
    _configure_answer_page(document)
    _render_answer_into(document, raw_exam, add_page_break=False)
    target = Path(output_path)
    target.parent.mkdir(parents=True, exist_ok=True)
    document.save(target)
    return target


def append_answer_to_docx(
    docx_path: str | Path,
    raw_exam: dict[str, Any],
) -> None:
    """Append a Guangzhou-style answer section to an exam DOCX."""

    if not answer_blocks(raw_exam):
        return
    target = Path(docx_path)
    document = Document(target)
    _render_answer_into(document, raw_exam, add_page_break=True)
    document.save(target)


def _render_answer_into(
    document: Any,
    raw_exam: dict[str, Any],
    *,
    add_page_break: bool,
) -> None:
    if add_page_break and document.paragraphs:
        document.add_page_break()
    metadata = raw_exam.get("metadata", {})
    titles = list(metadata.get("answer_title_lines", []))
    if not titles:
        titles = [
            str(metadata.get("exam_name", "")).strip(),
            "语文试题参考答案及评分建议",
        ]
    for title in titles[:2]:
        paragraph = _new_paragraph(document, 0, "center", 0, 0, Pt(23.5))
        _add_run(paragraph, str(title), "SimSun", 14, bold=True)

    source_path = Path(str(metadata.get("answer_source_docx_path", "")))
    source = Document(source_path) if source_path.exists() else None
    for block in answer_blocks(raw_exam):
        kind = block.get("type")
        if kind == "answer_section":
            paragraph = _new_paragraph(document, 0, "left", 6.3, 0, Pt(17))
            _add_run(paragraph, str(block.get("text", "")), "SimHei", 12)
        elif kind == "answer_subsection":
            paragraph = _new_paragraph(document, 21, "left", 6.3, 0, Pt(17))
            _add_run(paragraph, str(block.get("name", "")), "SimSun", 10.5, bold=True)
            _add_run(paragraph, str(block.get("meta", "")), "SimSun", 10.5)
        elif kind in {"answer_question", "answer_text"}:
            header = str(block.get("header", "")).strip()
            if header:
                paragraph = _new_paragraph(document, 0, "left", 0, 0, Pt(17))
                _add_run(paragraph, header, "SimSun", 10.5)
            for entry in block.get("paragraphs", []):
                _render_answer_paragraph(document, entry)
        elif kind == "answer_table" and source is not None:
            index = int(block.get("source_table_index", -1))
            if 0 <= index < len(source.tables):
                copied = deepcopy(source.tables[index]._element)
                _remap_image_relationships(copied, source.part, document.part)
                document.element.body.insert(-1, copied)
    _add_page_numbers(document)


def _answer_role(
    text: str,
    state: str,
    question_number: int | None,
) -> tuple[str, str]:
    if text.startswith("评分参考"):
        return "scoring_label", "scoring"
    if text.startswith("答案示例"):
        return "answer_label", "answer"
    if EXAMPLE_LABEL_RE.match(text):
        return "example_label", "answer"
    if "参考译文" in text:
        return "translation_label", "translation"
    if state == "translation":
        return "translation_body", state
    if state == "scoring":
        return "scoring_rule", state
    if question_number == 23 or state == "composition":
        return "composition", "composition"
    if re.match(r"^\s*[A-D]\s*[（(]", text):
        return "objective_answer", "answer"
    if SCORE_PREFIX_RE.match(text) or LABELED_ANSWER_RE.match(text):
        return "mixed_answer", "answer"
    return "subjective_answer", "answer"


def _render_answer_paragraph(document: Any, entry: dict[str, Any]) -> None:
    role = str(entry.get("role", "subjective_answer"))
    text = str(entry.get("text", ""))
    style = {
        "objective_answer": ("SimSun", 10.5, False),
        "subjective_answer": ("KaiTi", 10.5, False),
        "mixed_answer": ("KaiTi", 10.5, False),
        "answer_label": ("SimSun", 10.5, False),
        "example_label": ("SimHei", 10.5, False),
        "scoring_label": ("SimHei", 10.5, False),
        "scoring_rule": ("SimSun", 10.5, False),
        "translation_label": ("SimHei", 10.5, False),
        "translation_body": ("SimSun", 10.5, False),
        "composition": ("SimSun", 10.5, False),
    }.get(role, ("KaiTi", 10.5, False))
    paragraph = _new_paragraph(document, 21, "left", 0, 0, Pt(17))
    prefix_end = 0
    if role == "mixed_answer":
        match = SCORE_PREFIX_RE.match(text) or LABELED_ANSWER_RE.match(text)
        prefix_end = match.end("prefix") if match else 0
    _render_with_decorations(
        paragraph,
        text,
        entry.get("runs", []),
        style[0],
        style[1],
        prefix_end=prefix_end,
    )


def _render_with_decorations(
    paragraph: Any,
    text: str,
    source_runs: Iterable[dict[str, Any]],
    font_name: str,
    size: float,
    *,
    prefix_end: int = 0,
) -> None:
    runs = list(source_runs)
    if not runs:
        if prefix_end:
            _add_run(paragraph, text[:prefix_end], "SimSun", size)
            _add_run(paragraph, text[prefix_end:], font_name, size)
        else:
            _add_run(paragraph, text, font_name, size)
        return
    boundaries = {0, len(text)}
    cursor = 0
    normalized: list[tuple[int, int, dict[str, Any]]] = []
    for spec in runs:
        value = str(spec.get("text", ""))
        start, end = cursor, min(len(text), cursor + len(value))
        boundaries.update((start, end))
        normalized.append((start, end, spec))
        cursor = end
    if prefix_end:
        boundaries.add(prefix_end)
    points = sorted(boundaries)
    for start, end in zip(points, points[1:]):
        if start >= end:
            continue
        source = next(
            (spec for left, right, spec in normalized if left <= start < right),
            {},
        )
        fragment_font = "SimSun" if prefix_end and start < prefix_end else font_name
        run = _add_run(paragraph, text[start:end], fragment_font, size)
        if source.get("underline"):
            run.underline = True
        emphasis = str(source.get("emphasis", ""))
        if emphasis:
            element = OxmlElement("w:em")
            element.set(qn("w:val"), emphasis)
            run._element.get_or_add_rPr().append(element)


def _source_run_decorations(paragraph: Any) -> list[dict[str, Any]]:
    result: list[dict[str, Any]] = []
    for run in paragraph.runs:
        if not run.text:
            continue
        emphasis = run._r.xpath("./w:rPr/w:em")
        result.append(
            {
                "text": run.text,
                "underline": bool(run.underline),
                "emphasis": emphasis[0].get(qn("w:val")) if emphasis else "",
            }
        )
    return result


def _new_paragraph(
    document: Any,
    first_indent_pt: float,
    alignment: str,
    before_pt: float,
    after_pt: float,
    line_spacing: Any,
) -> Any:
    paragraph = document.add_paragraph()
    paragraph.alignment = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
    }[alignment]
    fmt = paragraph.paragraph_format
    fmt.first_line_indent = Pt(first_indent_pt)
    fmt.space_before = Pt(before_pt)
    fmt.space_after = Pt(after_pt)
    fmt.line_spacing = line_spacing
    fmt.widow_control = True
    return paragraph


def _add_run(
    paragraph: Any,
    text: str,
    font_name: str,
    size: float,
    *,
    bold: bool = False,
) -> Any:
    run = paragraph.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(size)
    run.bold = bold
    fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    for key in ("ascii", "hAnsi", "eastAsia", "cs"):
        fonts.set(qn(f"w:{key}"), font_name)
    return run


def _configure_answer_page(document: Any) -> None:
    for section in document.sections:
        section.page_width = Mm(210)
        section.page_height = Mm(297)
        section.top_margin = Mm(25.4)
        section.bottom_margin = Mm(25.4)
        section.left_margin = Mm(31.75)
        section.right_margin = Mm(31.75)
        section.header_distance = Mm(15.01)
        section.footer_distance = Mm(17.50)


def _add_page_numbers(document: Any) -> None:
    for section in document.sections:
        paragraph = section.footer.paragraphs[0]
        if paragraph.text.strip():
            continue
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run()
        for kind, value in (
            ("fldChar", "begin"),
            ("instrText", " PAGE "),
            ("fldChar", "separate"),
            ("t", "1"),
            ("fldChar", "end"),
        ):
            node = OxmlElement(f"w:{kind}")
            if kind == "fldChar":
                node.set(qn("w:fldCharType"), value)
            elif kind == "instrText":
                node.set(qn("xml:space"), "preserve")
                node.text = value
            else:
                node.text = value
            run._r.append(node)


def _looks_like_title(text: str, paragraph: Any, count: int) -> bool:
    if count >= 2 or MAJOR_SECTION_RE.fullmatch(text):
        return False
    if count == 0:
        return paragraph.alignment == WD_ALIGN_PARAGRAPH.CENTER or (
            "参考答案" not in text and "评分建议" not in text
        )
    return "参考答案" in text or "评分建议" in text or "答案与解析" in text


def _paragraph_index_by_element(document: Any, element: Any) -> int | None:
    for index, paragraph in enumerate(document.paragraphs):
        if paragraph._p is element:
            return index
    return None


__all__ = [
    "ANSWER_MARKER_RE",
    "answer_blocks",
    "append_answer_to_docx",
    "attach_answer_blocks",
    "exam_blocks",
    "find_answer_start",
    "is_standalone_answer_docx",
    "parse_answer_docx",
    "render_answer_docx",
    "standalone_answer_model",
]
