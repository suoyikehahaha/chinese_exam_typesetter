"""Final contextual formatting for regional headers and semantic text roles."""

from __future__ import annotations

from pathlib import Path
import re
from typing import Any, Iterator

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.text.paragraph import Paragraph


SUBSECTION_META_RE = re.compile(r"[（(]\s*本题共")


def apply_contextual_formatting_v6(
    docx_path: str | Path,
    raw_exam: dict[str, Any],
) -> None:
    """Apply deterministic header, subsection and semantic role formatting."""

    target = Path(docx_path)
    document = Document(target)
    metadata = raw_exam.get("metadata", {})
    _format_header(document, metadata)
    _format_semantic_paragraphs(document)
    document.save(target)


def _format_header(document: Any, metadata: dict[str, Any]) -> None:
    meta_text = str(metadata.get("meta_text", "")).strip()
    meta_paragraph = next(
        (paragraph for paragraph in document.paragraphs if paragraph.text.strip() == meta_text),
        None,
    )
    institution = str(metadata.get("institution_text", "")).strip()
    exam_date = str(metadata.get("exam_date", "")).strip()
    if meta_paragraph is not None and institution and exam_date:
        meta_paragraph.clear()
        meta_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        meta_paragraph.paragraph_format.left_indent = Pt(0)
        meta_paragraph.paragraph_format.first_line_indent = Pt(0)
        usable_width = (
            document.sections[0].page_width
            - document.sections[0].left_margin
            - document.sections[0].right_margin
        )
        meta_paragraph.paragraph_format.tab_stops.add_tab_stop(
            usable_width,
            WD_TAB_ALIGNMENT.RIGHT,
        )
        _font(meta_paragraph.add_run(institution), "SimSun", 10.5)
        meta_paragraph.add_run("\t")
        _font(meta_paragraph.add_run(exam_date), "SimSun", 10.5)
        _bottom_rule(meta_paragraph)

        info = str(metadata.get("exam_info_text", "")).strip()
        if info and not any(
            paragraph.text.strip() == info for paragraph in document.paragraphs
        ):
            element = OxmlElement("w:p")
            meta_paragraph._p.addnext(element)
            paragraph = Paragraph(element, meta_paragraph._parent)
            paragraph.style = "Exam_exam_meta"
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.left_indent = Pt(0)
            paragraph.paragraph_format.first_line_indent = Pt(10.5 * 2)
            _font(paragraph.add_run(info), "SimSun", 10.5)


def _format_semantic_paragraphs(document: Any) -> None:
    for paragraph in _all_paragraphs(document):
        style_name = paragraph.style.name
        if style_name == "Exam_subsection":
            _format_subsection(paragraph)
        elif style_name == "Exam_material_title":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _font_all(paragraph, "SimHei", 10.5)
        elif style_name == "Exam_material_author":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _font_all(paragraph, "FangSong", 10.5)
        elif style_name == "Exam_poetry":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _font_all(paragraph, "KaiTi", 10.5)
        elif style_name == "Exam_composition_prompt":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.left_indent = Pt(0)
            paragraph.paragraph_format.first_line_indent = Pt(10.5 * 2)
            _font_all(paragraph, "SimSun", 10.5)
        elif style_name == "Exam_notice_title":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.left_indent = Pt(0)
            paragraph.paragraph_format.first_line_indent = Pt(0)
        elif style_name == "Exam_notice_body":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.left_indent = Pt(0)
            paragraph.paragraph_format.first_line_indent = Pt(10.5 * 2)
            _font_all(paragraph, "SimSun", 10.5)


def _format_subsection(paragraph: Any) -> None:
    text = paragraph.text
    match = SUBSECTION_META_RE.search(text)
    name = text[: match.start()] if match else text
    meta = text[match.start() :] if match else ""
    paragraph.clear()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.left_indent = Pt(0)
    paragraph.paragraph_format.first_line_indent = Pt(0)
    run = paragraph.add_run(name)
    _font(run, "SimSun", 10.5)
    run.bold = True
    if meta:
        run = paragraph.add_run(meta)
        _font(run, "SimSun", 10.5)
        run.bold = False


def _bottom_rule(paragraph: Any) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    borders = ppr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        ppr.append(borders)
    bottom = borders.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        borders.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "808080")


def _font_all(paragraph: Any, name: str, size: float) -> None:
    for run in paragraph.runs:
        _font(run, name, size)


def _font(run: Any, name: str, size: float) -> None:
    run.font.name = name
    run.font.size = Pt(size)
    rfonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    for key in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn(f"w:{key}"), name)


def _all_paragraphs(document: Any) -> Iterator[Any]:
    yield from document.paragraphs
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs


__all__ = ["apply_contextual_formatting_v6"]
