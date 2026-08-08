"""题干、选项、题内小题和注意事项的最终固定规则。"""

from __future__ import annotations

from pathlib import Path
import re
from typing import Any, Iterator

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt


QUESTION_RE = re.compile(r"^\d{1,2}．")
OPTION_RE = re.compile(r"^[A-D]．")


def apply_exam_format_rules_v2(docx_path: str | Path) -> None:
    """应用用户确认的题号和段落缩进规则。"""

    target = Path(docx_path)
    document = Document(target)
    size = 10.5
    for paragraph in _all_paragraphs(document):
        text = paragraph.text.strip()
        if QUESTION_RE.match(text):
            paragraph.paragraph_format.left_indent = Pt(size * 1.5)
            paragraph.paragraph_format.first_line_indent = Pt(-size * 1.5)
        if OPTION_RE.match(text):
            paragraph.paragraph_format.left_indent = Pt(size * 1.5)
            paragraph.paragraph_format.first_line_indent = Pt(-size * 1.7)
        if paragraph.style.name == "Exam_subquestion":
            paragraph.paragraph_format.left_indent = Pt(0)
            paragraph.paragraph_format.first_line_indent = Pt(size * 2)
            for run in paragraph.runs:
                _font(run, "SimSun", size)
        if paragraph.style.name == "Exam_segmentation_text":
            paragraph.paragraph_format.left_indent = Pt(0)
            paragraph.paragraph_format.first_line_indent = Pt(size * 2)
        if paragraph.style.name == "Exam_notice_body":
            paragraph.paragraph_format.left_indent = Pt(0)
            paragraph.paragraph_format.first_line_indent = Pt(size * 2)
    document.save(target)


def _font(run: Any, name: str, size: float) -> None:
    run.font.name = name
    run.font.size = Pt(size)
    rfonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    for key in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn(f"w:{key}"), name)


def _all_paragraphs(document: Any) -> Iterator[Any]:
    yield from document.paragraphs
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs
