from __future__ import annotations

import tempfile
import unittest
from pathlib import Path

from docx import Document
from PIL import Image

from app.flexible_importers_v6 import _collect_source_decorations
from app.flexible_importers_v7 import _refine_exam_name
from app.internal_preview_v02 import _material, render_internal_preview


class V089PreviewRegressionTests(unittest.TestCase):
    def test_notice_item_is_not_used_as_exam_title(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            source = Path(directory) / "header.docx"
            document = Document()
            for value in (
                "高三语文学科练习",
                "考试时间：150分钟 试卷满分：150分",
                "注意事项：",
                "1.考试结束后，将答题卡交回。",
                "语文",
                "一、阅读（72分）",
            ):
                document.add_paragraph(value)
            document.save(source)

            result: dict[str, object] = {}
            _refine_exam_name(source, result)

        self.assertEqual(result["metadata"]["exam_name"], "高三语文学科练习")


    def test_source_collector_keeps_bold_and_underline(self) -> None:
        document = Document()
        paragraph = document.add_paragraph()
        run = paragraph.add_run("正文标记")
        run.bold = True
        run.underline = True

        entries = _collect_source_decorations(document)

        self.assertEqual(len(entries), 1)
        self.assertTrue(entries[0]["ranges"][0]["bold"])
        self.assertEqual(entries[0]["ranges"][0]["underline"], "single")

    def test_internal_preview_renders_native_drawing_marker(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            folder = Path(directory)
            image_path = folder / "source.png"
            Image.new("RGB", (80, 40), (220, 30, 30)).save(image_path)
            source = folder / "source.docx"
            document = Document()
            paragraph = document.add_paragraph()
            paragraph.add_run().add_picture(str(image_path))
            document.save(source)

            raw = {
                "metadata": {
                    "source_docx_path": str(source),
                    "native_objects": [
                        {
                            "kind": "drawing",
                            "source_paragraph_index": 0,
                            "marker": "[[NATIVE_DRAWING:0]]",
                        }
                    ],
                },
                "blocks": [
                    {
                        "type": "material",
                        "paragraphs": ["[[NATIVE_DRAWING:0]]"],
                    }
                ],
            }
            output = folder / "preview"
            result = render_internal_preview(raw, Path("templates/layout.yaml"), output)

            self.assertEqual(result.actual_pages, 1)
            with Image.open(result.pages[0]).convert("RGB") as rendered:
                red_pixels = sum(
                    1
                    for red, green, blue in rendered.get_flattened_data()
                    if red > 150 and green < 80 and blue < 80
                )
            self.assertGreater(red_pixels, 100)

    def test_material_label_is_left_aligned_in_internal_preview(self) -> None:
        calls: list[tuple[tuple[object, ...], dict[str, object]]] = []

        def capture(*args: object, **kwargs: object) -> None:
            calls.append((args, kwargs))

        _material(
            {"paragraphs": ["材料一：", "正文内容。"]},
            0,
            capture,
        )

        self.assertEqual(calls[0][0][1], "材料一：")
        self.assertEqual(calls[0][1]["align"], "left")


if __name__ == "__main__":
    unittest.main()
