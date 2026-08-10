"""Run-level protection regression test."""

from __future__ import annotations

from pathlib import Path
import tempfile
import unittest

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from app.run_formatting import set_run_font_preserving_properties


class RunFormattingTests(unittest.TestCase):
    def test_font_change_keeps_underline_emphasis_and_border(self) -> None:
        document = Document()
        paragraph = document.add_paragraph()
        run = paragraph.add_run("A")
        run.underline = True
        rpr = run._element.get_or_add_rPr()
        emphasis = OxmlElement("w:em")
        emphasis.set(qn("w:val"), "underDot")
        rpr.append(emphasis)
        border = OxmlElement("w:bdr")
        border.set(qn("w:val"), "single")
        rpr.append(border)
        set_run_font_preserving_properties(run, "SimSun", 10.5)
        with tempfile.TemporaryDirectory() as folder:
            path = Path(folder) / "run.docx"
            document.save(path)
            result = Document(path)
            target = result.paragraphs[0].runs[0]._element
            self.assertIsNotNone(target.xpath("./w:rPr/w:u"))
            self.assertIsNotNone(target.xpath("./w:rPr/w:em"))
            self.assertIsNotNone(target.xpath("./w:rPr/w:bdr"))


if __name__ == "__main__":
    unittest.main()
