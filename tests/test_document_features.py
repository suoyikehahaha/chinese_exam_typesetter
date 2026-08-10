from __future__ import annotations

from pathlib import Path
import tempfile
import unittest

from docx import Document
from docx.oxml.ns import qn
from PIL import Image

from app.chinese_typography import enable_chinese_typography
from app.flexible_importers import (
    _normalize_composition,
    _normalize_title_authors,
)
from app.inline_formatting import apply_inline_formats
from app.native_docx_objects import restore_native_objects
from app.semantic_formatting import apply_semantic_formatting


class DocumentFeatureTests(unittest.TestCase):
    def test_composition_is_split_into_three_roles(self) -> None:
        data = {
            "blocks": [
                {
                    "type": "question",
                    "question": {
                        "number": 23,
                        "stem": "阅读下面的材料，根据要求写作。",
                        "embedded_segments": [
                            [{"text": "核心材料。", "role": "body"}],
                            [{"text": "以上材料引发你怎样的联想和思考？", "role": "body"}],
                            [{"text": "要求：不少于800字。", "role": "body"}],
                        ],
                    },
                }
            ]
        }
        _normalize_composition(data)
        question = data["blocks"][0]["question"]
        self.assertEqual(question["composition_material"], ["核心材料。"])
        self.assertEqual(
            question["composition_prompt"],
            ["以上材料引发你怎样的联想和思考？"],
        )
        self.assertEqual(
            question["composition_requirements"],
            ["要求：不少于800字。"],
        )
        self.assertEqual(question["embedded_segments"], [])

    def test_same_line_title_and_author_are_recognized(self) -> None:
        data = {
            "blocks": [
                {"type": "instruction", "text": "阅读下面的文字，完成6～9题。"},
                {
                    "type": "material",
                    "paragraphs": ["卢沟桥之夜  林斤澜", "正文第一段。"],
                },
            ]
        }
        _normalize_title_authors(data)
        material = data["blocks"][1]
        self.assertEqual(material["title"], "卢沟桥之夜")
        self.assertEqual(material["author"], "林斤澜")
        self.assertEqual(material["paragraphs"], ["正文第一段。"])

    def test_text_label_source_and_chinese_typography(self) -> None:
        with tempfile.TemporaryDirectory() as folder:
            path = Path(folder) / "sample.docx"
            document = Document()
            document.add_paragraph("文本一：正文")
            document.add_paragraph("（摘编自某文）")
            document.save(path)
            apply_semantic_formatting(path)
            enable_chinese_typography(path)
            result = Document(path)
            label = result.paragraphs[0]
            self.assertEqual(label.runs[0].text, "文本一：")
            self.assertEqual(
                label.runs[0]._element.rPr.rFonts.get(qn("w:eastAsia")),
                "SimHei",
            )
            self.assertEqual(
                result.paragraphs[1]._element.pPr.jc.get(qn("w:val")),
                "right",
            )
            compat = result.settings.element.find(qn("w:compat"))
            self.assertIsNotNone(compat.find(qn("w:overflowPunct")))
            self.assertIsNotNone(compat.find(qn("w:kinsoku")))

    def test_inline_format_only_changes_selected_range(self) -> None:
        with tempfile.TemporaryDirectory() as folder:
            path = Path(folder) / "inline.docx"
            document = Document()
            document.add_paragraph("1．甲乙丙丁")
            document.save(path)
            raw = {
                "blocks": [
                    {
                        "type": "question",
                        "question": {
                            "number": 1,
                            "inline_formats": [
                                {
                                    "target": "stem",
                                    "target_index": 0,
                                    "start": 1,
                                    "end": 3,
                                    "font": "黑体",
                                    "size_pt": 10.5,
                                }
                            ],
                        },
                    }
                ]
            }
            apply_inline_formats(path, raw)
            paragraph = Document(path).paragraphs[0]
            fonts = [
                (
                    run.text,
                    run._element.rPr.rFonts.get(qn("w:eastAsia"))
                    if run._element.rPr is not None
                    and run._element.rPr.rFonts is not None
                    else None,
                )
                for run in paragraph.runs
            ]
            self.assertIn(("乙丙", "SimHei"), fonts)
            self.assertNotEqual(fonts[0][1], "SimHei")

    def test_native_table_and_image_are_restored(self) -> None:
        with tempfile.TemporaryDirectory() as folder:
            root = Path(folder)
            image_path = root / "pixel.png"
            Image.new("RGB", (40, 30), "red").save(image_path)
            source_path = root / "source.docx"
            source = Document()
            paragraph = source.add_paragraph("试卷标题")
            paragraph.add_run().add_picture(str(image_path))
            table = source.add_table(rows=2, cols=2)
            table.cell(0, 0).text = "甲"
            table.cell(0, 1).text = "乙"
            table.cell(1, 0).text = "丙"
            table.cell(1, 1).text = "丁"
            source.save(source_path)

            target_path = root / "target.docx"
            target = Document()
            target.add_paragraph("试卷标题")
            target.add_paragraph("[[NATIVE_TABLE:0]]")
            target.save(target_path)
            raw = {
                "metadata": {
                    "source_docx_path": str(source_path),
                    "native_objects": [
                        {
                            "kind": "drawing",
                            "source_paragraph_index": 0,
                            "target_text": "试卷标题",
                        },
                        {
                            "kind": "table",
                            "source_table_index": 0,
                            "marker": "[[NATIVE_TABLE:0]]",
                        },
                    ],
                }
            }
            restore_native_objects(target_path, raw)
            result = Document(target_path)
            self.assertEqual(len(result.tables), 1)
            self.assertEqual(result.tables[0].cell(1, 1).text, "丁")
            self.assertTrue(result.paragraphs[0]._p.xpath(".//w:drawing"))


if __name__ == "__main__":
    unittest.main()
