"""从导入的 DOCX 中原样复制表格和图片对象。"""

from __future__ import annotations

from copy import deepcopy
from io import BytesIO
from pathlib import Path
from typing import Any

from docx import Document
from docx.oxml.ns import qn


def restore_native_objects(docx_path: str | Path, raw_exam: dict[str, Any]) -> None:
    """恢复原文件中的表格 XML、图片数据、尺寸与定位参数。"""

    metadata = raw_exam.get("metadata", {})
    source_path = metadata.get("source_docx_path")
    objects = metadata.get("native_objects", [])
    if not source_path or not objects or not Path(source_path).exists():
        return
    source = Document(source_path)
    target_path = Path(docx_path)
    target = Document(target_path)
    for spec in objects:
        if spec.get("kind") == "table":
            _restore_table(source, target, spec)
        elif spec.get("kind") == "drawing":
            _restore_drawing(source, target, spec)
    target.save(target_path)


def _restore_table(source: Any, target: Any, spec: dict[str, Any]) -> None:
    index = int(spec["source_table_index"])
    if index >= len(source.tables):
        return
    paragraph = _find_paragraph(target, str(spec.get("marker", "")))
    if paragraph is None:
        return
    table_element = deepcopy(source.tables[index]._element)
    _remap_image_relationships(table_element, source.part, target.part)
    paragraph._p.addnext(table_element)
    paragraph._element.getparent().remove(paragraph._element)


def _restore_drawing(source: Any, target: Any, spec: dict[str, Any]) -> None:
    index = int(spec["source_paragraph_index"])
    if index >= len(source.paragraphs):
        return
    source_paragraph = source.paragraphs[index]
    target_paragraph = _find_paragraph(
        target,
        str(spec.get("marker") or spec.get("target_text") or ""),
    )
    if target_paragraph is None:
        return
    if spec.get("marker"):
        for run in list(target_paragraph.runs):
            target_paragraph._p.remove(run._r)
    for drawing in source_paragraph._p.xpath("./w:r/w:drawing | ./w:r/w:pict"):
        copied = deepcopy(drawing)
        _remap_image_relationships(copied, source.part, target.part)
        run = target_paragraph.add_run()
        run._r.append(copied)


def _remap_image_relationships(element: Any, source_part: Any, target_part: Any) -> None:
    for blip in element.xpath(".//a:blip"):
        old_id = blip.get(qn("r:embed"))
        if not old_id or old_id not in source_part.related_parts:
            continue
        image_part = source_part.related_parts[old_id]
        new_id, _ = target_part.get_or_add_image(BytesIO(image_part.blob))
        blip.set(qn("r:embed"), new_id)


def _find_paragraph(document: Any, text: str) -> Any | None:
    for paragraph in document.paragraphs:
        if paragraph.text.strip() == text.strip():
            return paragraph
    return None
