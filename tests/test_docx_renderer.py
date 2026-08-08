from pathlib import Path
import tempfile
import unittest

from docx import Document
from docx.oxml.ns import qn

from app.config import load_layout
from app.models import load_exam
from app.renderers import DocxRenderer


ROOT = Path(__file__).resolve().parents[1]


class DocxRendererTests(unittest.TestCase):
    def setUp(self) -> None:
        self.exam = load_exam(ROOT / "samples" / "exam.json")
        self.layout = load_layout(ROOT / "templates" / "layout.yaml")

    def _render(self, directory: str) -> Path:
        target = Path(directory) / "sample.docx"
        DocxRenderer(self.layout).render(self.exam, target)
        return target

    def test_confirmed_styles_are_encoded(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            doc = Document(self._render(temp_dir))
            objective = doc.styles["Exam_objective_stem"].paragraph_format
            subjective = doc.styles["Exam_subjective_stem"].paragraph_format
            author = doc.styles["Exam_material_author"]
            self.assertEqual(0, objective.first_line_indent.pt)
            self.assertAlmostEqual(15.75, subjective.first_line_indent.pt, places=1)
            self.assertEqual("FangSong", author.font.name)

    def test_east_asia_fonts_and_boxed_markers_exist(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            doc = Document(self._render(temp_dir))
            all_runs = [run for paragraph in doc.paragraphs for run in paragraph.runs]
            marker_runs = [run for run in all_runs if run.text in set("ABCDEFGH")]
            self.assertEqual(8, len(marker_runs))
            for run in marker_runs:
                rfonts = run._element.rPr.rFonts
                self.assertEqual("SimSun", rfonts.get(qn("w:eastAsia")))
                self.assertIsNotNone(run._element.rPr.find(qn("w:bdr")))

    def test_composition_uses_three_roles(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            doc = Document(self._render(temp_dir))
            styles = {paragraph.style.name for paragraph in doc.paragraphs}
            self.assertIn("Exam_composition_material", styles)
            self.assertIn("Exam_composition_prompt", styles)
            self.assertIn("Exam_composition_requirements", styles)


if __name__ == "__main__":
    unittest.main()
