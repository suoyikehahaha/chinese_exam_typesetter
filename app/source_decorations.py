"""Restore source underline and emphasis marks after typesetting."""

from __future__ import annotations

from copy import deepcopy
from pathlib import Path
from typing import Any, Iterator

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def restore_source_decorations(
    docx_path: str | Path,
    raw_exam: dict[str, Any],
) -> None:
    """Restore recorded character marks on matching output paragraphs."""

    entries = raw_exam.get("metadata", {}).get("source_decorations", [])
    if not entries:
        return
    target = Path(docx_path)
    document = Document(target)
    paragraphs = list(_all_paragraphs(document))
    for entry in entries:
        matches = [
            paragraph
            for paragraph in paragraphs
            if paragraph.text.strip() == str(entry.get("text", "")).strip()
        ]
        occurrence = int(entry.get("occurrence", 0))
        if occurrence >= len(matches):
            continue
        paragraph = matches[occurrence]
        for mark in entry.get("ranges", []):
            _decorate_range(
                paragraph,
                int(mark.get("start", 0)),
                int(mark.get("end", 0)),
                str(mark.get("underline", "")),
                str(mark.get("emphasis", "")),
                bool(mark.get("bold", False)),
            )
    document.save(target)


def _decorate_range(
    paragraph: Any,
    start: int,
    end: int,
    underline: str,
    emphasis: str,
    bold: bool = False,
) -> None:
    if start >= end:
        return
    old_runs = list(paragraph.runs)
    fragments: list[tuple[str, Any, bool]] = []
    cursor = 0
    for run in old_runs:
        text = run.text
        run_start = cursor
        run_end = cursor + len(text)
        points = {run_start, run_end}
        if run_start < start < run_end:
            points.add(start)
        if run_start < end < run_end:
            points.add(end)
        boundaries = sorted(points)
        for left, right in zip(boundaries, boundaries[1:]):
            fragments.append(
                (
                    text[left - run_start : right - run_start],
                    run,
                    start <= left and right <= end,
                )
            )
        cursor = run_end
    for run in old_runs:
        paragraph._p.remove(run._r)
    for text, source, selected in fragments:
        run = paragraph.add_run(text)
        if source._r.rPr is not None:
            run._r.insert(0, deepcopy(source._r.rPr))
        if selected:
            rpr = run._element.get_or_add_rPr()
            if underline:
                old = rpr.find(qn("w:u"))
                if old is not None:
                    rpr.remove(old)
                element = OxmlElement("w:u")
                element.set(qn("w:val"), underline)
                rpr.append(element)
            if emphasis:
                old = rpr.find(qn("w:em"))
                if old is not None:
                    rpr.remove(old)
                element = OxmlElement("w:em")
                element.set(qn("w:val"), emphasis)
                rpr.append(element)
            if bold:
                old = rpr.find(qn("w:b"))
                if old is None:
                    rpr.append(OxmlElement("w:b"))


def _all_paragraphs(document: Any) -> Iterator[Any]:
    yield from document.paragraphs
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs
